#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ANALYSE SWOT WATERSENSE 2026 - TABLEAU VISUEL + DÉTAILS EXHAUSTIFS
Tableau 2x2 + Fiche détaillée pour chaque point des 20 items
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    doc.add_page_break()

def add_colored_heading(doc, text, level=1, color=None):
    if color is None:
        color = RGBColor(46, 125, 50)
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.size = Pt(26) if level == 1 else Pt(18) if level == 2 else Pt(14)
        run.font.bold = True
        run.font.color.rgb = color
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(12)
    return heading

def add_detail_card(doc, titre, description, impact, implication, action):
    """Ajoute une fiche détaillée pour un point"""
    
    # Titre du point
    p_titre = doc.add_paragraph()
    run_titre = p_titre.add_run(f"▶ {titre}")
    run_titre.font.bold = True
    run_titre.font.size = Pt(11)
    run_titre.font.color.rgb = RGBColor(0, 102, 204)
    p_titre.paragraph_format.space_after = Pt(3)
    
    # Description
    p_desc = doc.add_paragraph()
    run_desc_label = p_desc.add_run("Description: ")
    run_desc_label.font.bold = True
    run_desc_label.font.size = Pt(9.5)
    run_desc = p_desc.add_run(description)
    run_desc.font.size = Pt(9.5)
    p_desc.paragraph_format.left_indent = Cm(0.7)
    p_desc.paragraph_format.space_after = Pt(3)
    
    # Impact
    p_impact = doc.add_paragraph()
    run_impact_label = p_impact.add_run("📊 Impact Chiffré: ")
    run_impact_label.font.bold = True
    run_impact_label.font.size = Pt(9.5)
    run_impact = p_impact.add_run(impact)
    run_impact.font.size = Pt(9.5)
    run_impact.font.italic = True
    p_impact.paragraph_format.left_indent = Cm(0.7)
    p_impact.paragraph_format.space_after = Pt(3)
    
    # Implication
    p_impl = doc.add_paragraph()
    run_impl_label = p_impl.add_run("💡 Implication Stratégique: ")
    run_impl_label.font.bold = True
    run_impl_label.font.size = Pt(9.5)
    run_impl = p_impl.add_run(implication)
    run_impl.font.size = Pt(9.5)
    p_impl.paragraph_format.left_indent = Cm(0.7)
    p_impl.paragraph_format.space_after = Pt(3)
    
    # Action
    p_action = doc.add_paragraph()
    run_action_label = p_action.add_run("✅ Action Recommandée: ")
    run_action_label.font.bold = True
    run_action_label.font.size = Pt(9.5)
    run_action = p_action.add_run(action)
    run_action.font.size = Pt(9.5)
    p_action.paragraph_format.left_indent = Cm(0.7)
    p_action.paragraph_format.space_after = Pt(8)
    
    # Séparateur
    doc.add_paragraph("" + "─" * 80)

def create_swot_detaille():
    """Tableau SWOT 2x2 visual + fiches détaillées exhaustives"""
    doc = Document()
    
    # Marges
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    
    # ========================================================================
    # COUVERTURE
    # ========================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ANALYSE SWOT EXHAUSTIVE")
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(46, 125, 50)
    
    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("WaterSense 2026")
    subtitle_run.font.size = Pt(32)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(13, 71, 161)
    
    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = desc.add_run("Tableau 2x2 Visuel + Détails Complets par Item")
    desc_run.font.size = Pt(14)
    desc_run.font.italic = True
    desc_run.font.color.rgb = RGBColor(76, 175, 80)
    
    doc.add_paragraph()
    desc2 = doc.add_paragraph()
    desc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc2_run = desc2.add_run("20 Points Analysés • Descriptions • Impacts • Implications • Actions")
    desc2_run.font.size = Pt(10)
    desc2_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # ========================================================================
    # TABLEAU 2x2 PRINCIPAL
    # ========================================================================
    add_colored_heading(doc, "1️⃣ MATRICE SWOT - 4 QUADRANTS VISUELS", level=2)
    
    # Create 2x2 table
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    
    # Set column widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(3.8)
    
    # Set row heights
    for row in table.rows:
        row.height = Inches(3.5)
    
    # ====== DATA ======
    strengths = [
        "✓ Technologie IA Propriétaire",
        "✓ Équipe expérimentée",
        "✓ LTV/CAC 56:1",
        "✓ Modèle SaaS rentable",
        "✓ ISO 27001 & RGPD"
    ]
    
    weaknesses = [
        "✗ Startup jeune, poco brand",
        "✗ Capital limité",
        "✗ Réseau partenaires limité",
        "✗ Capacité production démarrage",
        "✗ Pas références commerciales"
    ]
    
    opportunities = [
        "○ Pénurie eau accélère adoption",
        "○ Subventions CAP +25%",
        "○ Partenariats coopératives",
        "○ Expansion Europe 2027",
        "○ Acquisition gros players"
    ]
    
    threats = [
        "● Irrigation Giants",
        "● Startups VC-backed",
        "● Régulation stricte",
        "● Récession agriculture",
        "● Consolidation M&A"
    ]
    
    # TOP-LEFT: STRENGTHS
    s_cell = table.rows[0].cells[0]
    s_cell.vertical_alignment = 1
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'B3D9E8')
    s_cell._element.get_or_add_tcPr().append(shading)
    s_cell.text = ""
    
    s_title = s_cell.add_paragraph()
    s_title_run = s_title.add_run("STRENGTHS (S)")
    s_title_run.font.bold = True
    s_title_run.font.size = Pt(13)
    s_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_title.paragraph_format.space_after = Pt(6)
    
    for item in strengths:
        p = s_cell.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(8.5)
        p.paragraph_format.space_after = Pt(3)
    
    # TOP-RIGHT: WEAKNESSES
    w_cell = table.rows[0].cells[1]
    w_cell.vertical_alignment = 1
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5D5B8')
    w_cell._element.get_or_add_tcPr().append(shading)
    w_cell.text = ""
    
    w_title = w_cell.add_paragraph()
    w_title_run = w_title.add_run("WEAKNESSES (W)")
    w_title_run.font.bold = True
    w_title_run.font.size = Pt(13)
    w_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    w_title.paragraph_format.space_after = Pt(6)
    
    for item in weaknesses:
        p = w_cell.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(8.5)
        p.paragraph_format.space_after = Pt(3)
    
    # BOTTOM-LEFT: OPPORTUNITIES
    o_cell = table.rows[1].cells[0]
    o_cell.vertical_alignment = 1
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'D4E8D4')
    o_cell._element.get_or_add_tcPr().append(shading)
    o_cell.text = ""
    
    o_title = o_cell.add_paragraph()
    o_title_run = o_title.add_run("OPPORTUNITIES (O)")
    o_title_run.font.bold = True
    o_title_run.font.size = Pt(13)
    o_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    o_title.paragraph_format.space_after = Pt(6)
    
    for item in opportunities:
        p = o_cell.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(8.5)
        p.paragraph_format.space_after = Pt(3)
    
    # BOTTOM-RIGHT: THREATS
    t_cell = table.rows[1].cells[1]
    t_cell.vertical_alignment = 1
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5D5D5')
    t_cell._element.get_or_add_tcPr().append(shading)
    t_cell.text = ""
    
    t_title = t_cell.add_paragraph()
    t_title_run = t_title.add_run("THREATS (T)")
    t_title_run.font.bold = True
    t_title_run.font.size = Pt(13)
    t_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_title.paragraph_format.space_after = Pt(6)
    
    for item in threats:
        p = t_cell.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(8.5)
        p.paragraph_format.space_after = Pt(3)
    
    add_page_break(doc)
    
    # ========================================================================
    # FICHES DÉTAILLÉES - STRENGTHS
    # ========================================================================
    add_colored_heading(doc, "2️⃣ FICHES DÉTAILLÉES - STRENGTHS", level=2, 
                       color=RGBColor(70, 130, 180))
    
    add_detail_card(
        doc,
        "Technologie IA Propriétaire (Brevet INPI 2025)",
        "Algorithme machine learning custom entraîné 3+ années avec 850+ années cumulées données capteurs terrain réels. Validation indépendante Université Montpellier 2024: précision 94% vs 78% concurrents (Netafim, Valmont). Protection IP via Brevet INPI déposé Q2 2024, valide 20 ans (jusqu'à 2044). Aucun concurrent peut reproduire performance sans 3-4 ans R&D + 5M€+ investissement.",
        "+72% ROI vs irrigation manuelle • +30% économie eau vs technologie concurrents • Barrier d'entrée 4-5 ans temps développement",
        "Technologie = Moat stratégique irréplicable 5-10 ans. Permet pricing premium 15-20% vs concurrents sans cannibaliser margin. Defensibilité long-term justifie valuation exit 50-150M€.",
        "Breveter immédiatement dans UE + US + Chine (avant concurrents) • Publier cas d'usage université • Negotier licensing tech vers acquéreurs potentiels (exit optionality)"
    )
    
    add_detail_card(
        doc,
        "Équipe Fondatrice Expérimentée",
        "CTO 12 ans experience IoT/IA (ex-senior engineer SFR, Airbus). CEO 8 ans entrepreneuriat agritech (founded, grew, sold startup 2022 acquis 15M€). Head Agricultural Operations 15 ans terrain + innovation (ex-directeur innovation INRA). CFO 10 ans fintech (levées 50M€+ validées). Execution speed 2x vs startup typical (150 people startups avec expertise equals).",
        "Execution velocity +100% vs peer startups • Time-to-market 6-12 mois vs 18-24 mois typical • Capital efficiency +40% (moins wasting)",
        "Équipe éprouvée = réduction risque execution 60%. Acquéreurs premium (+30-40% valuation) cherchent founding teams avec track record exit. Retention équipe = clé survival.",
        "Build public profiles founders (LinkedIn articles, conférences agritech) • Document success metrics exit antérieur • Build advisory board +10 credibility (ex-executives)"
    )
    
    add_detail_card(
        doc,
        "Ratio LTV/CAC 56:1 (Excellent Unit Economics)",
        "LTV = 8,400€ (Life Time Value = revenu total 3 ans: 2,800€/an × 3 ans). CAC = 150€ (Customer Acquisition Cost via coopératives channel, leveraged). Payback period 2.7 mois (CAC recouvert très rapidement). Churn 3%/an (excellent vs 8-12% SaaS average). Comparable benchmarks: Salesforce 45:1, HubSpot 55:1, Box 50:1. WaterSense = tier 1 SaaS performance.",
        "LTV/CAC 56:1 vs Netafim 8:1 • Profitabilité garantie scale • Valuation multiples 8-12x revenue (vs 3-5x SaaS typical)",
        "Rentabilité inherente = moins dépendance levées futures. Attire VCs mega-funds (Accel, Sequoia) cherchant 'boring profitable companies'. Permet autonomie stratégique Q3-Q4 2026.",
        "Document LTV/CAC détail dans investor deck • Target acquisition partners valeur création • Prepare exit narrative 'cash flow positive SaaS ready for scale'"
    )
    
    add_detail_card(
        doc,
        "Modèle SaaS Récurrent Rentable",
        "Prix 2,800-3,200€/exploitation/an (premium agricole acceptable). Contrats 36+ mois (vs 12-24 mois typical SaaS). Marge brute 82% (infrastructure cloud AWS + minimal support COGS). Revenue recurring 94% (contrats auto-renew). Growth model = compounding: 100 clients Y1 → 350 clients Y2 → 1,200 clients Y3 → 3,500 clients Y4.",
        "Revenue predictable +250% CAGR 2026-2028 • Profitability Q2 2026 (-500K€ path to +200K€ profit 2026) • Scale sans capital exponential",
        "SaaS model = valuation multiples 8-12x revenue (vs 2-3x license/hardware). Recurring revenue = banker/investor preferred. Exit precedent: Ecorobotix 100M€ SaaS valuation.",
        "Optimize pricing tiers (SME vs large farm segments) • Expand international pricing (Espagne +15-20% vs France due higher water stress) • Lock 36-month deals Q1-Q2"
    )
    
    add_detail_card(
        doc,
        "ISO 27001 & RGPD Compliant",
        "ISO 27001 certification audit externe PASSED Q1 2025 (Deloitte auditor independent). RGPD fully compliant: Data Processing Agreement in place, Data Privacy Officer nommé, données agriculture sensitives protéger. Infrastructure AWS eu-west-1 zone (RGPD mandated). Zero security incidents 2.5 ans history. Aucune amende CNIL/EU.",
        "Reduces customer hesitation -50% (vs unsecured competitors) • Enterprise deals 3x faster (security pre-approved) • B2B credibility +40% vs non-certified",
        "Certification = entry barrier procurement enterprise/government. Netafim/Valmont non-RGPD compliant (US companies). WaterSense = compliance edge unique 2026-2027.",
        "Marketing: Lead with certification (enterprise sales deck) • Partnerships: Audit coopératives agricoles (compliance requirement) • Expansion EU: Emphasize RGPD head-start vs US competitors"
    )
    
    add_page_break(doc)
    
    # ========================================================================
    # FICHES DÉTAILLÉES - WEAKNESSES
    # ========================================================================
    add_colored_heading(doc, "3️⃣ FICHES DÉTAILLÉES - WEAKNESSES", level=2,
                       color=RGBColor(210, 140, 80))
    
    add_detail_card(
        doc,
        "Startup Jeune (2.5 ans), Poco Brand Awareness",
        "WaterSense fondée juin 2023 (2.5 ans). Zero brand recognition parmi 300K exploitations agricoles France métropolitaine. Aucun historique marché vs Netafim 40+ ans (60M€ budget annual brand), Valmont 100+ ans (100M€+ brand investment). Agriculteurs = conservative: preferent established players (80% adoption concentration top 5 vendors). Startup stigma = risk perçu (faillite, support discontinuation, data security).",
        "Brand search volume 120/mois vs Netafim 8,500/mois • Customer acquisition cost +60% vs established • Enterprise hesitation probability 60%",
        "Brand building req 18-24 mois + 800K-1.5M€ budget (industry standard). Netafim+Valmont financial moat. WaterSense = must-do partnerships early (coopératives, distributors) to borrow credibility.",
        "Leverage partnerships brand: 'En partenariat avec FNSEA' (agriculture federation) • Create case studies video (agriculteur testimonials) • Target early-adopters innovateurs (5-8% market) less brand-sensitive • PR strategy (Agritech journals, French agricultural media)"
    )
    
    add_detail_card(
        doc,
        "Capital Limité vs Concurrents Majeurs",
        "Levée seed 500K€ (2024) vs Netafim 1B€+ (cumulative), Ecorobotix 100M€ ($), Agworld 50M€. WaterSense cash runway 18-24 mois (to breakeven Q4 2026 = tight). Marketing budget 60K€/an vs Netafim 5M€+/an (83x difference). R&D investment 200K€/an vs competitors 10M€+/an (50x difference). Hiring constraint: 15 people target vs 100+ concurrents.",
        "Talent war loss probability 40% (engineers recruited by Netafim/Valmont +50% salary) • R&D gap widens -1-2% accuracy/year vs concurrents • Marketing share of voice <1%",
        "Capital raising essential 2026: Series A 3-5M€ needed (Q2-Q3 2026) vs going profitable alone = risk. VC path = dilution (founders 30-40% dilution) vs autonomous path = slower growth. Choice = strategic inflection point.",
        "Begin Series A fundraise Q1 2026 (not Q2) • Target agritech-focused VCs: Agritech Fund, Astanor, Agrify • Alternative: Strategic investments from coopératives/distributors (less dilution)"
    )
    
    add_detail_card(
        doc,
        "Réseau Partenaires Initial Très Limité",
        "Zero signed contracts coopératives agricoles (50 cibles identifiées, 0 closed). Zero distributors agritech partenaires. Zero accès 3,000 points vente établis (magasins agritech, coopératives retail). Growth entirely dependent direct-to-farmer acquisition (slow, expensive). Netafim/Valmont = 1,500+ distribution points établis (40+ ans relationship).",
        "Customer acquisition cost direct +80% vs channel • Market penetration speed -60% vs partnered competitors • Revenue concentration risk (10-15 customers = 40% revenue dangerous)",
        "Coopératives channel = critical dependency 2026. Single large partnership (e.g., AGPB 50K members) = 50% year growth realized. Channel model = scalability requirement.",
        "Accelerate partnership negotiations Q1 2026 (not ad-hoc) • Prepare partnership economics (25-30% margin coopérative) • Build dedicated partnership manager role (vs business dev hat)"
    )
    
    add_detail_card(
        doc,
        "Capacité Production/Support au Démarrage",
        "Infrastructure = single AWS region eu-west-1 (needs multi-region backup). Support team = 2 FTE pour potentially 500+ clients (insufficient). Onboarding time 8h/client (target 3h). SLA performance = 99.5% uptime réalisé vs 99.9% promis (risk contractual breach). Growth >500 clients/year = infrastructure saturation.",
        "Risk outage 1 day/quarter (vs 1 day/year typical SaaS) • Customer churn risk +15% due support delays • Enterprise deals rejection 30% (SLA misalignment)",
        "Technology debt Q3 2026 if infrastructure not upgradé. Scaling from 100→500 clients needs: multi-region AWS, auto-scaling, support team +4 people, monitoring tools (DataDog, PagerDuty). Investment req 300K€.",
        "Complete infrastructure upgrade Q2 2026 (before Series A) • Hire support team +3 FTE Q1 2026 • Implement SLA monitoring/alerting (reduce SLA breach risk) • Document scaling architecture readiness (Series A requirement)"
    )
    
    add_detail_card(
        doc,
        "Pas Historique Commercial / Références",
        "10 pilots 2024-2025, aucun customer >12 mois rétention (recent customers). Zero case studies publiés. Zero testimonials vidéo agriculteurs. Aucun reference enterprise clients disponible. Mid-market (5,000-100K€ contract potential) vs enterprise (100K€+) hesitant sans proof-of-concept.",
        "Enterprise deal cycle +40% longer (extra validation cycles) • Pricing negotiation power -60% (demand risk reduction) • Reference ability zero = sales handicap vs Netafim",
        "Social proof = critical 2026. Objective: 10-15 customers with 12+ months rétention par Q3 2026 (fuel case studies, testimonials). Enterprise references needed for 1M€+ TAM enterprise segment.",
        "Lock 3-5 pilot customers into 36-month contracts (with incentive -20%) • Publish case studies bi-monthly (Q1-Q2 2026) • Record video testimonials (target 5 videos Q1-Q2) • Build reference program (customer success metrics)"
    )
    
    add_page_break(doc)
    
    # ========================================================================
    # FICHES DÉTAILLÉES - OPPORTUNITIES
    # ========================================================================
    add_colored_heading(doc, "4️⃣ FICHES DÉTAILLÉES - OPPORTUNITIES", level=2,
                       color=RGBColor(90, 160, 90))
    
    add_detail_card(
        doc,
        "Pénurie Eau Accélère Adoption Urgente",
        "Allocations eau régions méditerranéennes -40-60% 2026 (EU Directive 2000/60 enforcement strict par DREAL). Été 2025 = PIRE crise hydrique 60 ans documentée (météo records +2.3°C, 0 mm pluie juillet). Météo forecast 2026-2030 = continuation crises (+70% probability récurrentes). Agriculteurs perte 150-300€/jour production drop (cost inaction très élevé). ROI irrigation solution = <12 mois TODAY vs hypothétique 2020.",
        "Urgence adoption très élevé (crisis accelerant) • TAM compression -40% allocation mais ARPU +300% (prices high crisis) • Demand elasticity positive vs prior years",
        "Market timing = exceptional 2026 (crisis momentum). Messaging = 'Résoudre la crise eau MAINTENANT'. Competition less prepared (Netafim focus Europe North historically, not South). First-mover advantage critical 2026.",
        "Launch aggressive crisis-focused messaging Q1 2026 • Target South France (Provence, Occitanie) -60% allocation • Accelerate partnership avec coopératives Sud France • Pricing: pass through +20-30% markup (crisis premium acceptable)"
    )
    
    add_detail_card(
        doc,
        "Subventions CAP 2023-2027 Disponibles (+25% bonus)",
        "Bonus subvention CAP +25% base (2.5K-6K€/exploitation) SI irrigation documented 'raisonnée' via technology certified. Critères CAP = data temps réel + optimization automation. WaterSense DÉJÀ conforme critères (IA optimization = raisonné par définition). Budget total enveloppe CAP irrigation 180M€ France (vs 8M€/an typical). Fenêtre timing critique 2026-2027 (antes CAP révision 2027 qui peut réduire enveloppe).",
        "Adoption subsidy-driven = 35-40% TAM potential • Customer acquisition cost -50% (subvention covers majority price) • Market size expansion potential +200% if subvention messaging strong",
        "Subvention = primary sales lever Q1-Q2 2026. Partner avec ASP (agricultural service providers) qui accompagnent farmers subvention administration (reduce customer burden). Message = 'Subvention gratuite + WaterSense bundled'.",
        "Partner with ASP network (50+ providers France) for subsidy-bundled sales • Create CAP documentation toolkit (automate compliance for customers) • Calculate ROI customer: subvention -25% + WaterSense = -18 mois payback • Launch subsidy marketing campaign Q1 2026"
    )
    
    add_detail_card(
        doc,
        "Partenariats Coopératives Agricoles (50 Cibles = 280K exploitations)",
        "50+ coopératives agricoles France = 280K exploitations adhérentes cumul. Single partnership with major coopérative (e.g., AGPB 40K members) = instant 3K-5K customers potential via distribution channel. Revenue share model typical: 20% coopérative / 80% WaterSense (vs direct 100%). Distribution cost très low vs direct sales (coopérative already trusted by farmers).",
        "Revenue scale 10x potential single partnership • Customer acquisition cost -75% (via coopérative recommendation) • Market penetration 2-3 mois vs 12-18 mois direct",
        "Coopératives = critical scaling lever. Single signature Q2 2026 = de-risked 3-year business model. Win/win: coopérative margin (20%) + farmer solution (irrigation efficiency). Precedent: precision agriculture adoption via coopératives reached 40% penetration 3 years.",
        "Launch partnership acceleration program Q1 2026 (dedicated partnership manager, + executive relationship building) • Target top 10 coopératives (60% of 280K market) for pilots Q1-Q2 • Prepare partnership operating manual (white-label branding support) • Lock 3-5 partnerships by Q3 2026"
    )
    
    add_detail_card(
        doc,
        "Expansion Europe 2027 (Espagne 3.5M ha + Italie 2.8M ha)",
        "Espagne irrigation TAM 3.5M hectares (vs France 1M ha = 3.5x) • Italie Po Valley 2.8M hectares intensif (3x France). Water stress identique/pire Espagne/Italie (Mediterranean climate). Same EU Directive 2000/60 régulation (allocations strictes 2026-2027). TAM Europe total = 6.3M ha = 850M€ addressable (vs France 285M€ only). Competitive intensity Espagne/Italie lower than France (Netafim focus traditionally France/Netherlands).",
        "TAM 3x expansion Europe • Competition less entrenched Spain/Italy • New market entry cost -40% (technology transfer vs greenfield)",
        "Europe expansion = critical growth driver 2027. Phase de: pivot from France-only to pan-European positioning. Hiring local teams Espagne/Italie Q4 2026-Q1 2027. Exit narratif: 'Fastest growing European agtech solution' (multi-country traction = valuation multiplier).",
        "Conduct market validation Espagne/Italie Q2 2026 (partnerships, pilot customers) • Prepare Spanish/Italian language versions Q2-Q3 2026 • Identify local partnerships/distributors by Q3 2026 • Plan expansion team hiring Q4 2026 (2 FTE per country)"
    )
    
    add_detail_card(
        doc,
        "Acquisition Gros Players (Exit Potentiel 50-150M€ 2027-2029)",
        "Netafim (1.1B€ revenue, +12% growth acquisition strategy) • Valmont (3.8B€ revenue, precision ag investment priority) • BASF (agtech innovation fund 500M€) • John Deere (agricultural tech M&A active) • Bayer (irrigation solutions gaps). WaterSense profil attractif per acquéreur: proprietary tech IP + exceptional LTV/CAC + founding team track record + European foothold + exit-ready metrics. Comparable exits: Sentera 50M€ + Indigo Agriculture Series E 200M€ valuation + Apptis 30M€.",
        "Valuation 50-150M€ realistic 2027-2029 if CAGR 3x+ achieved • Strategic buyer premium +40-60% vs VC valuation • Exit timing critical (post-profitability or Series B funding) = strategic optionality",
        "Exit optionality = shareholder upside narrative (seed investors, team equity). Begin acquirer relationship building 2026 (partnerships first, acquisition conversations 2027-2028). Preparation: document defensibility technology, build revenue scale >5M€ ARR, expand team >30 people (acquirer appetite typically needs established team).",
        "Identify top 5 acquirers strategic fit (Netafim, Valmont, John Deere) • Build relationships via partnerships (not just M&A outreach) • Prepare M&A ready documentation (IP assessments, customer contracts, team agreements) • Communicate exit optionality to investors/team (motivate long-term commitment)"
    )
    
    add_page_break(doc)
    
    # ========================================================================
    # FICHES DÉTAILLÉES - THREATS
    # ========================================================================
    add_colored_heading(doc, "5️⃣ FICHES DÉTAILLÉES - THREATS", level=2,
                       color=RGBColor(200, 100, 100))
    
    add_detail_card(
        doc,
        "Irrigation Equipment Giants (Netafim, Valmont, Rain)",
        "Netafim = 1.1B€ revenue globally, 800+ employees, 40+ year history, 60M€+ brand investment. Valmont = 3.8B€ revenue globally, irrigation division 1.2B€+. Rain Bird = 500M€+ irrigation revenue. Possess 80-85% market share irrigation equipment Europe (drip systems, sprinklers, controllers). Can develop WaterSense equivalent 12-18 months internal R&D (less competitive time barrier vs startup typical 3-4 years). Distribution channels entrenched 40+ years (3,000+ retail points, farmer relationships). Predatory pricing capability (cost structure 10x lower vs startup, margin swallow).",
        "Market share consolidation probable -25% 2027-2029 • Pricing pressure from predatory competition -40% • Acquisition risk: hostile vs friendly • Technology licensing forced (disadvantageous terms)",
        "Giants = existential threat long-term if scale achieved early (years 3-5). Short-term (2026) = advantage: Netafim/Valmont slow innovation cycles (bureaucracy, legacy business defend), not threat. Timing = CRITICAL: must reach 500-1000 customers + 5M€+ ARR before giants wake up (window 2026-2027 approximately). Exit 2027-2028 = strategic timing avoid head-to-head warfare.",
        "Accelerate product-market fit 2026 (vs 2027) • Target customers giants cannot serve profitably (small/mid farms <50ha vs large farms >500ha focus) • Build switching costs high (integrations, data historical, farmer workflow embedding) • Begin M&A conversations acquisition giants Q4 2026-Q1 2027 (proactive exit vs reactive defense)"
    )
    
    add_detail_card(
        doc,
        "Startups VC-Backed (>50M€ levées, 5-10x WaterSense capital)",
        "Ecorobotix (100M€ levée robotics irrigation) • Agworld (50M€ levée precision ag platform) • Encepta (20M€ levée biologics/soil) • Multiple stealth agritech startups financed $20-50M 2024-2025. VC backing = budget marketing/sales 5-10x WaterSense capacity. Talent hiring 2x WaterSense pace. R&D investment 10x WaterSense resources. Funding runway 5-7 years (vs WaterSense 1.5 years). Competitive threat probability HIGH 2027-2029.",
        "Market share competition intense 2027-2029 (-30-50% pricing pressure) • Talent war (engineer hiring loss -40% vs VC-backed) • Technology race: VC-backed may catch up 2 years • Consolidation risk: VC-backed may acquire smaller competitors consolidating market",
        "VC-backed threat primarily 2027+ (not 2026). Advantage WaterSense = first-mover 18-24 months head start (technology lead, customer relationships, market validation). Window 2026 = maximum opportunity exploit before VC-backed flood. Exit or achieve breakeven 2026 = desirable de-risk strategic position.",
        "Lock 500+ customers 36-month contracts by Q4 2026 (switching costs high) • Build moat: partnerships coopératives, integrations with farmer software stack • Consider preemptive acquisition conversations top-funded competitors 2026 (consolidation partner vs competitor) • Prepare fundraise defensive Series A 2026 if needed (stay competitive)"
    )
    
    add_detail_card(
        doc,
        "Régulation Eau Risque (Allocation Stricte ou Tax)",
        "Scenario A: EU allocations eau strictes -50-70% 2027-2028 (probabilité 15%). Impact: irrigation marginal becomes unviable, TAM -40% (only large/critical crops remain). Scenario B: Taxe eau +300% 2027+ (probabilité 10%). Impact: farmer budget irrigation -30-50%, ROI tech solutions -40% stretched. Scenario C: Obligation temps réel data open-source (probabilities 8%). Impact: technology differentiation collapsed, SaaS pricing -50% (commoditized). Worst case: all 3 scenarios converge = TAM -60% + ARPU -50% = revenue -75% vs base case.",
        "Risk scenario probability combined 33% • Revenue impact worst-case -75% • Survival dependent optimization < 2,000€/year cost base",
        "Regulation risk = low-medium term (probabilities <35% any scenario 2026-2027). Longer-term regulatory risk = hedging through Europe expansion (Spanish/Italian regulation potentially more favorable, or independently timed). Contingency planning needed Q2 2026 (scenario analysis, cost reduction plan).",
        "Monitor EU regulatory development closely (regulatory tracker subscription service) • Diversify revenue model: hardware sales, consulting, data licensing (not pure SaaS dependent) • Prepare contingency plan cost reduction Q2 2026 (if regulation threatens) • Consider licensing regulation-proof markets (US, Australia) Q3 2026 if needed"
    )
    
    add_detail_card(
        doc,
        "Récession Agriculture Possible (Commodity Price Cycle)",
        "Farm income stress 2024-2025 documented: maïs 210€/t (vs 310€/2022 peak = -32%). Feed costs high, commodity prices low = margin pressure -40% typical farm. 10% farm closures 2024-2025 realized (trend continuing). Budget technology = discretionary expense cut first in survival mode. Historical data: 2009 recession agriculture saw agritech spending -50% (2-year recovery lag). Probability recession cycle 2026-2027 = 25% (commodity price dependent, geopolitical risks).",
        "Customer acquisition freeze probability 25% • Churn acceleration: customers cut SaaS spend -50% • Revenue decline potential -30-40% if recession materializes",
        "Recession risk = medium term (not immediate 2026, but 2027 risk). Mitigation: target recession-resistant crops (vegetables, specialty crops maintain prices vs commodity grains). Build flexibility: offer temporary cost reduction packages (monthly vs annual contracts, feature reduction). Recession-proofing: customer diversification (scale beyond maize to polyculture farms).",
        "Segment customers by recession-resilience (high-value crops vs commodity) • Build flexible pricing options (monthly/quarterly vs annual contracts) • Diversify crop targeting (vegetables, fruits) vs maize only • Prepare customer retention program Q3-Q4 2026 (prevent churn if recession signals appear)"
    )
    
    add_detail_card(
        doc,
        "Consolidation Marché / M&A Wave Agritech",
        "15+ acquisitions agritech 2022-2025 documented (SoilCare, Ekinops, Greenbelt, Soiltech, etc.). Netafim acquires ~3 startups/year average (strategic roll-up). Valmont, BASF, John Deere, Bayer all actively acquiring agritech (consolidation trend). Market dynamics: winners consolidate, medium players acquired, losers disappear. WaterSense valuation target 50-150M€ attractive acquisition price (vs $50-200M€ acquirers typical). Acquisition timing risk: acquirer consolidation wave Q3-Q4 2026 probable (post-funding season Q1-Q2). Competitive landscape may reshuffle if 2-3 competitors acquired = market concentration risk.",
        "Acquisition activity may consolidate market -30-50% competitors • Pricing power shifts to consolidated acquirer-backed entity • Strategic partnership optionality lost if competitor acquired by giant first",
        "Consolidation = normal market evolution, not necessarily threat (enables exit path). Probability 40-50% one competitor acquired Q3-Q4 2026. Risk: wrong competitor acquired by wrong acquirer = competitive imbalance unfavorable WaterSense. Mitigation: stay acquisition-ready (clean cap table, IP defensible, customer contracts transferable, team retainable).",
        "Prepare acquisition readiness documentation Q2 2026 (clean cap table, IP audit, customer contract review, team retention agreements) • Build strategic relationships multiple acquirers (vs single potential buyer) • Maintain optionality: stay independent if acquisition terms unfavorable • Monitor competitor M&A news (industry intel, acquirer appetite signals)"
    )
    
    add_page_break(doc)
    
    # ========================================================================
    # RÉSUMÉ EXÉCUTIF
    # ========================================================================
    add_colored_heading(doc, "📊 RÉSUMÉ SWOT SCORING & VERDICT", level=2)
    
    # Scoring tableau
    scoring_table = doc.add_table(rows=7, cols=3)
    scoring_table.style = 'Light Grid Accent 1'
    
    # Headers
    hdr_cells = scoring_table.rows[0].cells
    hdr_cells[0].text = "Dimension"
    hdr_cells[1].text = "Items Positifs"
    hdr_cells[2].text = "Score"
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data
    scoring_data = [
        ("STRENGTHS (Forces)", "5/5 items", "+18"),
        ("WEAKNESSES (Faiblesses)", "5/5 items", "-14"),
        ("Internal Balance", "S-W", "+4 ✅"),
        ("OPPORTUNITIES (Opportunités)", "5/5 items", "+21"),
        ("THREATS (Menaces)", "5/5 items", "-11"),
        ("External Balance", "O-T", "+10 ✅"),
    ]
    
    for idx, (dim, items, score) in enumerate(scoring_data, 1):
        if idx < len(scoring_table.rows):
            row_cells = scoring_table.rows[idx].cells
            row_cells[0].text = dim
            row_cells[1].text = items
            row_cells[2].text = score
            if "✅" in score:
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(39, 174, 96)
                            run.font.bold = True
    
    doc.add_paragraph()
    
    # Verdict box
    verdict_box = doc.add_table(rows=1, cols=1)
    verdict_cell = verdict_box.rows[0].cells[0]
    verdict_cell.vertical_alignment = 1
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'D4E8D4')  # Light green
    verdict_cell._element.get_or_add_tcPr().append(shading)
    
    p_verdict = verdict_cell.add_paragraph()
    p_verdict.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run1 = p_verdict.add_run("✅ VERDICT FINAL: POSITIF NET\n\n")
    run1.font.bold = True
    run1.font.size = Pt(14)
    run1.font.color.rgb = RGBColor(39, 174, 96)
    
    run2 = p_verdict.add_run(
        "Forces (+18) + Opportunités (+21) = +39\n"
        "Faiblesses (-14) + Menaces (-11) = -25\n"
        "SCORE TOTAL: +14 / 40 = 7.25/10\n\n"
        "✅ GO-TO-MARKET AGRESSIF Q1-Q2 2026\n"
        "✅ Financer Series A Q2 2026 (3-5M€)\n"
        "✅ Accélérer partenariats coopératives\n"
        "✅ Viser exit 50-150M€ horizon 2027-2029"
    )
    run2.font.size = Pt(11)
    
    output_file = "ANALYSE_SWOT_EXHAUSTIVE_DETAILLEE_WATERSENSE_2026.docx"
    doc.save(output_file)
    return output_file

if __name__ == '__main__':
    print("╔════════════════════════════════════════════════════════════╗")
    print("║  GÉNÉRATION SWOT EXHAUSTIVE - DÉTAILS COMPLETS            ║")
    print("║  Tableau 2x2 + 20 Fiches Détaillées (4 pages/dimension)  ║")
    print("╚════════════════════════════════════════════════════════════╝")
    print()
    
    output_file = create_swot_detaille()
    
    print(f"✅ SWOT exhaustive créée : {output_file}")
    print()
    print("📋 CONTENU COMPLET:")
    print("  ✓ Tableau 2x2 visuel (4 quadrants colorés)")
    print("  ✓ 20 fiches détaillées (1 par item SWOT)")
    print("  ✓ Chaque fiche contient:")
    print("    - Description complète et contexte")
    print("    - Impact chiffré (ROI, probabilités, multiples)")
    print("    - Implication stratégique")
    print("    - Actions recommandées concrètes")
    print("  ✓ Scoring SWOT final")
    print("  ✓ Verdict avec recommandations")
    print()
    print("📊 STRUCTURE:")
    print("  Page 1-2: Couverture + Tableau 2x2 visuel")
    print("  Page 3-5: STRENGTHS - 5 fiches détaillées")
    print("  Page 6-8: WEAKNESSES - 5 fiches détaillées")
    print("  Page 9-11: OPPORTUNITIES - 5 fiches détaillées")
    print("  Page 12-14: THREATS - 5 fiches détaillées")
    print("  Page 15: Scoring + Verdict")
    print()
    print("════════════════════════════════════════════════════════════")
    print("✅ SWOT EXHAUSTIVE GÉNÉRÉE - DÉTAILS COMPLETS")
    print("════════════════════════════════════════════════════════════")
