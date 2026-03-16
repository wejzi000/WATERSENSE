from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0, 102, 204)
        elif level == 2:
            run.font.color.rgb = RGBColor(0, 102, 204)
    return para

# Créer un nouveau document cohérent
doc = Document()

# Page de titre
title = doc.add_heading('WaterSense', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.font.bold = True

subtitle = doc.add_paragraph('Irrigation Intelligente pour Agriculteurs')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 170, 68)
    run.font.bold = True

tagline = doc.add_paragraph('Arrosez Moins, Gagnez Plus 💰')
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in tagline.runs:
    run.font.size = Pt(14)

doc.add_page_break()

# Sommaire
add_heading(doc, 'Sommaire', 1)
toc_items = [
    '1. Présentation WaterSense',
    '2. Comment ça marche',
    '3. Ce que vous gagnez',
    '4. Tarification',
    '5. Témoignages',
    '6. FAQ',
]
for item in toc_items:
    doc.add_paragraph(item)

doc.add_page_break()

# Section 1: Présentation
add_heading(doc, '1. Présentation WaterSense', 1)

add_heading(doc, 'Le Problème', 2)
doc.add_paragraph(
    'Les agriculteurs font face à deux défis majeurs:\n'
    '• Vous ne savez jamais précisément quand arroser\n'
    '• Vous gaspillez de l\'eau ou vous perdez des récoltes'
)

add_heading(doc, 'La Solution WaterSense', 2)
doc.add_paragraph(
    'WaterSense vous dit QUAND et COMBIEN arroser, '
    'en mesurant directement votre sol et en regardant la météo.\n'
    'Résultat: vous économisez 25% d\'eau ET gagnez 15% de récolte en plus.'
)

add_heading(doc, 'Les Chiffres', 2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
shade_cell(table.rows[0].cells[0], '0066CC')
shade_cell(table.rows[0].cells[1], '0066CC')
for cell in table.rows[0].cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True

data = [
    ('Économie d\'eau', '-25%'),
    ('Augmentation récolte', '+15%'),
    ('Gain par hectare/an', '1700-2300€'),
]
for i, (key, value) in enumerate(data, 1):
    table.rows[i].cells[0].text = key
    table.rows[i].cells[1].text = value

doc.add_page_break()

# Section 2: Comment ça marche
add_heading(doc, '2. Comment ça marche', 1)

steps = [
    ('Étape 1️⃣: Les capteurs mesurent', 
     'De petits boîtiers dans votre champ mesurent l\'humidité du sol en temps réel.'),
    ('Étape 2️⃣: WaterSense calcule', 
     'Le système regarde la météo, le sol, la culture et décide du meilleur moment d\'arroser.'),
    ('Étape 3️⃣: Vous recevez une alerte', 
     'SMS ou app: "Demain matin 5h, irriguer 45 minutes". Vous appuyez sur un bouton (ou auto).'),
    ('Résultat ✓: Vous gagnez plus', 
     'Pas de surcharge, pas de manque. Juste ce qu\'il faut. Chaque fois.'),
]

for title, desc in steps:
    add_heading(doc, title, 2)
    doc.add_paragraph(desc)

doc.add_page_break()

# Section 3: Ce que vous gagnez
add_heading(doc, '3. Ce que vous gagnez', 1)

benefits = [
    ('💧 Réduction d\'eau', 
     'Vous arrosez juste ce qu\'il faut, pas plus.',
     '200-300€/ha d\'économie'),
    ('+🌾 Plus de récolte', 
     'Meilleur timing = plus de kilos/ha.',
     '1500-2000€ de revenu supplémentaire'),
    ('✓ Moins de stress', 
     'Fini les décisions difficiles. WaterSense vous guide.',
     'Temps gagné + tranquillité'),
]

for title, desc, gain in benefits:
    add_heading(doc, title, 2)
    doc.add_paragraph(desc)
    p = doc.add_paragraph(gain)
    for run in p.runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 170, 68)

doc.add_page_break()

# Section 4: Tarification
add_heading(doc, '4. Tarification', 1)

table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
for cell in table.rows[0].cells:
    shade_cell(cell, '0066CC')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True

table.rows[0].cells[0].text = 'Plan'
table.rows[0].cells[1].text = 'Prix/an'
table.rows[0].cells[2].text = 'Pour qui'

pricing = [
    ('Essential', '2,900€', 'Petite exploitation (1-5ha)'),
    ('Standard', '4,200€', 'Exploitation moyenne (5-20ha) - POPULAIRE'),
    ('Premium', '6,500€', 'Grandes exploitations (20ha+)'),
]

for i, (plan, price, target) in enumerate(pricing, 1):
    table.rows[i].cells[0].text = plan
    table.rows[i].cells[1].text = price
    table.rows[i].cells[2].text = target

doc.add_paragraph()
p = doc.add_paragraph('💰 Rentabilité: Vous économisez en 2-3 mois!')
for run in p.runs:
    run.font.bold = True

doc.add_page_break()

# Section 5: Témoignages
add_heading(doc, '5. Témoignages', 1)

testimonials = [
    ('"Avant, je n\'étais jamais sûr de quand arroser. Maintenant, WaterSense me dit exactement quoi faire. '
     'J\'ai économisé 250€ d\'eau et gagné 1800€ en plus cette année. C\'est fou !"',
     'Jean-Pierre - Agriculteur Maïs, Nouvelle-Aquitaine'),
    
    ('"J\'utilise WaterSense sur mon verger depuis 2 ans. Les pommes sont plus grosses, '
     'j\'utilise 30% moins d\'eau. C\'est le meilleur investissement que j\'ai fait."',
     'Marie - Arboricultrice, Provence'),
]

for quote, author in testimonials:
    p = doc.add_paragraph(quote)
    p.paragraph_format.left_indent = Inches(0.5)
    p_author = doc.add_paragraph(author)
    for run in p_author.runs:
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph()

doc.add_page_break()

# Section 6: FAQ
add_heading(doc, '6. Questions Fréquentes', 1)

faqs = [
    ('Ça complique pas le travail ?',
     'Non. Au lieu de vous demander "j\'arrose demain ?", WaterSense vous dit "oui, à 5h du matin, 45 min". C\'est plus simple.'),
    
    ('Et si c\'est compliqué à installer ?',
     'Nos techniciens installent tout pour vous. Vous n\'avez rien à faire. On vous montre comment ça marche et c\'est bon.'),
    
    ('Et si j\'ai des problèmes ?',
     'Vous appelez. On répond. C\'est compris dans le tarif. On est là pour vous aider, pas pour vous vendre un truc et disparaître.'),
    
    ('Ça marche pour toutes les cultures ?',
     'Maïs, blé, pommes, melons, tomates... Oui. On a réglé WaterSense pour chaque culture. Vous nous dites ce que vous cultivez, c\'est fait.'),
    
    ('Essai gratuit: comment ça marche ?',
     '30 jours gratuits. Aucun engagement. Aucune surprise. Pas de carte bancaire demandée.'),
]

for q, a in faqs:
    add_heading(doc, q, 2)
    doc.add_paragraph(a)

doc.add_page_break()

# Conclusion
add_heading(doc, 'Conclusion', 1)

doc.add_paragraph(
    'WaterSense est conçu pour les agriculteurs qui veulent:'
)

conclusion_points = [
    'Économiser de l\'eau (25% en moyenne)',
    'Augmenter leurs rendements (+15%)',
    'Gagner du temps et de l\'argent (1700-2300€/ha/an)',
    'Avoir la tranquillité d\'esprit',
]

for point in conclusion_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('Prêt à commencer ? Essai gratuit 30 jours.')
for run in p.runs:
    run.font.bold = True
    run.font.size = Pt(12)

# Sauvegarder
output_path = r'c:\Users\wejde\Downloads\APP WATERSENSE\DOSSIER_MARKETING_WATERSENSE_COHERENT.docx'
doc.save(output_path)

print(f"✅ Nouveau document créé: DOSSIER_MARKETING_WATERSENSE_COHERENT.docx")
print(f"📄 Chemin: {output_path}")
print(f"\n✓ Document restructuré et cohérent avec le site agricole:")
print(f"  - 6 sections claires et structurées")
print(f"  - Langage agricole simple (pas de jargon technique)")
print(f"  - Tableaux et mise en forme professionnelle")
print(f"  - Cohérent avec le site WaterSense")
