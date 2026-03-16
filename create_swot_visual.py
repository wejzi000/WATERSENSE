#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ANALYSE SWOT WATERSENSE 2026 - TABLEAU 2x2 VISUAL AVEC DÉTAILS
Format comme l'image SWOT avec quadrants colorés remplis de contenu
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_colored_heading(doc, text, level=1, color=None):
    if color is None:
        color = RGBColor(46, 125, 50)
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.size = Pt(26) if level == 1 else Pt(18) if level == 2 else Pt(14)
        run.font.bold = True
        run.font.color.rgb = color
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(12)
    return heading

def create_swot_visual():
    """Tableau SWOT 2x2 visual avec détails"""
    doc = Document()
    
    # Marges
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    
    # ========================================================================
    # COUVERTURE
    # ========================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ANALYSE SWOT")
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(46, 125, 50)
    
    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("WaterSense 2026")
    subtitle_run.font.size = Pt(32)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(13, 71, 161)
    
    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = desc.add_run("Tableau SWOT Détaillé 2x2")
    desc_run.font.size = Pt(14)
    desc_run.font.italic = True
    desc_run.font.color.rgb = RGBColor(76, 175, 80)
    
    doc.add_paragraph()
    
    # ========================================================================
    # TABLEAU 2x2 PRINCIPAL
    # ========================================================================
    add_colored_heading(doc, "MATRICE SWOT - 4 QUADRANTS", level=2)
    
    # Create 2x2 table
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    
    # Set column widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(3.8)
    
    # Set row heights
    for row in table.rows:
        row.height = Inches(3.5)
    
    # Remove default borders and set spacing
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # ====== DATA ======
    strengths = [
        "✓ Technologie IA Propriétaire (Brevet INPI 2025)",
        "✓ Équipe expérimentée (CTO 12 ans IoT, CEO exit agritech 2022)",
        "✓ LTV/CAC 56:1 (excellent unit economics)",
        "✓ Modèle SaaS rentable (82% marge brute)",
        "✓ ISO 27001 & RGPD compliant"
    ]
    
    weaknesses = [
        "✗ Startup jeune (2.5 ans), poco brand awareness",
        "✗ Capital limité (500K€ vs Netafim 1B€+)",
        "✗ Réseau partenaires très limité",
        "✗ Capacité production démarrage (2 FTE support)",
        "✗ Pas historique commercial/références"
    ]
    
    opportunities = [
        "○ Pénurie eau accélère adoption (allocations -40-60%)",
        "○ Subventions CAP +25% disponibles 2026-2027",
        "○ Partenariats coopératives (280K exploitations)",
        "○ Expansion Europe 2027 (Espagne 3.5M ha + Italie)",
        "○ Acquisition gros players (50-150M€ exit 2027-2029)"
    ]
    
    threats = [
        "● Irrigation Giants (Netafim, Valmont, 80%+ marché)",
        "● Startups VC-backed ($100M+) entrent marché",
        "● Régulation eau stricte possible (15-20% prob)",
        "● Récession agriculture (-40% budget tech)",
        "● Consolidation marché (15+ acquisitions/an)"
    ]
    
    # TOP-LEFT: STRENGTHS (Blue)
    s_cell = table.rows[0].cells[0]
    s_cell.vertical_alignment = 1  # Top
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'B3D9E8')  # Light blue
    s_cell._element.get_or_add_tcPr().append(shading)
    
    # Clear cell
    s_cell.text = ""
    
    # Add title
    s_title = s_cell.add_paragraph()
    s_title_run = s_title.add_run("STRENGTHS (S)")
    s_title_run.font.bold = True
    s_title_run.font.size = Pt(13)
    s_title_run.font.color.rgb = RGBColor(0, 0, 0)
    s_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_title.paragraph_format.space_after = Pt(6)
    
    # Add items
    for item in strengths:
        p = s_cell.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(20, 20, 20)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.1
    
    # TOP-RIGHT: WEAKNESSES (Orange)
    w_cell = table.rows[0].cells[1]
    w_cell.vertical_alignment = 1
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5D5B8')  # Light orange
    w_cell._element.get_or_add_tcPr().append(shading)
    
    w_cell.text = ""
    
    w_title = w_cell.add_paragraph()
    w_title_run = w_title.add_run("WEAKNESSES (W)")
    w_title_run.font.bold = True
    w_title_run.font.size = Pt(13)
    w_title_run.font.color.rgb = RGBColor(0, 0, 0)
    w_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    w_title.paragraph_format.space_after = Pt(6)
    
    for item in weaknesses:
        p = w_cell.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(20, 20, 20)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.1
    
    # BOTTOM-LEFT: OPPORTUNITIES (Green)
    o_cell = table.rows[1].cells[0]
    o_cell.vertical_alignment = 1
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'D4E8D4')  # Light green
    o_cell._element.get_or_add_tcPr().append(shading)
    
    o_cell.text = ""
    
    o_title = o_cell.add_paragraph()
    o_title_run = o_title.add_run("OPPORTUNITIES (O)")
    o_title_run.font.bold = True
    o_title_run.font.size = Pt(13)
    o_title_run.font.color.rgb = RGBColor(0, 0, 0)
    o_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    o_title.paragraph_format.space_after = Pt(6)
    
    for item in opportunities:
        p = o_cell.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(20, 20, 20)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.1
    
    # BOTTOM-RIGHT: THREATS (Pink/Red)
    t_cell = table.rows[1].cells[1]
    t_cell.vertical_alignment = 1
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5D5D5')  # Light pink
    t_cell._element.get_or_add_tcPr().append(shading)
    
    t_cell.text = ""
    
    t_title = t_cell.add_paragraph()
    t_title_run = t_title.add_run("THREATS (T)")
    t_title_run.font.bold = True
    t_title_run.font.size = Pt(13)
    t_title_run.font.color.rgb = RGBColor(0, 0, 0)
    t_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_title.paragraph_format.space_after = Pt(6)
    
    for item in threats:
        p = t_cell.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(20, 20, 20)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.1
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ========================================================================
    # ANALYSE DÉTAILLÉE PAR QUADRANT
    # ========================================================================
    add_colored_heading(doc, "ANALYSE DÉTAILLÉE - STRENGTHS", level=2, 
                       color=RGBColor(70, 130, 180))
    
    strengths_detail = [
        ("Technologie IA Propriétaire (Brevet INPI 2025)", 
         "Algorithme custom entraîné 3+ années données terrain. Précision 94% vs 78% concurrents. Propriété IP protégée 20 ans. Barrière d'entrée TRÈS SIGNIFICANTE vs Netafim/Valmont."),
        
        ("Équipe Fondatrice Expérimentée",
         "CTO 12 ans IoT/IA • CEO 8 ans agritech startup (exit 2022 acquis) • Head Agriculture 15 ans terrain & innovation • CFO 10 ans fintech levées 50M€+. Exécution 2x vs startup typical."),
        
        ("LTV/CAC 56:1 (Excellent Unit Economics)",
         "LTV 8,400€ (3 ans @ 2,800€/an) • CAC 150€ (via coopératives) • Payback 2.7 mois • Churn 3%/an vs 8-12% SaaS average. Rentabilité garantie long-term."),
        
        ("Modèle SaaS Récurrent Rentable",
         "Prix 2,800-3,200€/exploitation/an • Marge brute 82% (cloud + support minimal COGS) • Contrats 36+ mois prévisibilité • Scaling sans capital lourd. Revenue prévisible investisseurs."),
        
        ("ISO 27001 & RGPD Compliant",
         "ISO 27001 certification audit externe PASSED 2025 • RGPD conforme complet • DPO nommé • Données stockage AWS eu-west-1 • Aucune amende/violation. Entrée marché without friction.")
    ]
    
    for titre, detail in strengths_detail:
        p = doc.add_paragraph()
        p_run = p.add_run(f"✓ {titre}")
        p_run.font.bold = True
        p_run.font.size = Pt(11)
        p_run.font.color.rgb = RGBColor(70, 130, 180)
        p.paragraph_format.space_after = Pt(2)
        
        p2 = doc.add_paragraph(detail)
        for run in p2.runs:
            run.font.size = Pt(9.5)
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(8)
    
    # ========================================================================
    # ANALYSE DÉTAILLÉE PAR QUADRANT - WEAKNESSES
    # ========================================================================
    add_colored_heading(doc, "ANALYSE DÉTAILLÉE - WEAKNESSES", level=2,
                       color=RGBColor(210, 140, 80))
    
    weaknesses_detail = [
        ("Startup Jeune, Poco Brand Awareness",
         "WaterSense fondée 2023 (2.5 ans existence). 0% brand recognition parmi 300K exploitations agricoles France. Aucun historique marché vs Netafim 40+ ans, Valmont 100+ ans. Nécessite gros budget marketing."),
        
        ("Capital Limité vs Concurrents Majeurs",
         "Levée 500K€ vs Netafim 1B€+, Ecorobotix $100M, Agworld $50M. Budget marketing 60K€/an vs 5M€+ concurrents. Impossible R&D race vs gros players. Dépendance financement externe."),
        
        ("Réseau Partenaires Initial Très Limité",
         "0 contrats coopératives agricoles signés (50 cibles identifiées). 0 distributeurs agritech partenaires. Pas d'accès 3,000 points vente établis. Croissance dépendante early-mover adoption directe, pas channel."),
        
        ("Capacité Production/Support au Démarrage",
         "Infrastructure AWS single region (besoin multi-region backup). Support 2 FTE pour 100+ clients (doit être 5+). Onboarding 8h/client (target 3h). SLA 99.5% réalisé vs 99.9% promis. Risk outage si croissance >500 clients/an."),
        
        ("Pas Historique Commercial/Références",
         "10 pilotes court terme 2024-2025, aucun client >12 mois rétention. Zéro case studies publiés. Aucun testimonials vidéo agriculteurs. Mid-market & enterprise hesitants sans references/social proof etablies.")
    ]
    
    for titre, detail in weaknesses_detail:
        p = doc.add_paragraph()
        p_run = p.add_run(f"✗ {titre}")
        p_run.font.bold = True
        p_run.font.size = Pt(11)
        p_run.font.color.rgb = RGBColor(210, 140, 80)
        p.paragraph_format.space_after = Pt(2)
        
        p2 = doc.add_paragraph(detail)
        for run in p2.runs:
            run.font.size = Pt(9.5)
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(8)
    
    doc.add_page_break()
    
    # ========================================================================
    # ANALYSE DÉTAILLÉE PAR QUADRANT - OPPORTUNITIES
    # ========================================================================
    add_colored_heading(doc, "ANALYSE DÉTAILLÉE - OPPORTUNITIES", level=2,
                       color=RGBColor(90, 160, 90))
    
    opportunities_detail = [
        ("Pénurie Eau Accélère Adoption Urgente",
         "Allocations eau -40-60% zones méditerranéennes 2026 (EU Directive 2000/60 enforcement). Été 2025 = PIRE crise hydrique 60 ans documentée. Probabilité crises récurrentes 70%+ 2026-2030. BESOIN #1 agriculteurs NOW vs hypothétique 2020."),
        
        ("Subventions CAP 2023-2027 Disponibles",
         "Bonus +25% subvention base CAP (2.5K-6K€/exploitation) SI irrigation 'raisonnée' documentée. Critères: données temps réel + technologie certifiée. WaterSense DÉJÀ conforme critères. Budget total 180M€ irrigation CAP. Fenêtre 2026-2027 avant CAP révision 2027."),
        
        ("Partenariats Coopératives Agricoles (50 Cibles)",
         "50+ coopératives agricoles France existantes = 280K exploitations adhérentes. Accord partenaire unique coopérative = 3K-5K clients instantanés via distribution channel established. Revenue share 20% coopérative / 80% WaterSense. Distribution scale 10x vs direct sales."),
        
        ("Expansion Europe 2027 (Espagne/Italie)",
         "Espagne: 3.5M ha irrigation (vs France 1M ha) • Italie: 2.8M ha irrigation Po Valley intensif • TAM Europe 850M€ vs France 285M€ = 3x croissance potential. Même réglementation UE (Directive Eau 2000/60). Crises eau identiques. Moins tech competition."),
        
        ("Acquisition Gros Players (Exit Potentiel)",
         "Netafim, Valmont, John Deere, BASF cherchent tech irrigation innovation. WaterSense profil attractif: tech IP propriétaire + LTV/CAC excellent + équipe. Comparable exits: Sentera $50M. Valuation 50-150M€ realistic 2027-2029 if 3x+ CAGR. Exit timeline attractif shareholders.")
    ]
    
    for titre, detail in opportunities_detail:
        p = doc.add_paragraph()
        p_run = p.add_run(f"○ {titre}")
        p_run.font.bold = True
        p_run.font.size = Pt(11)
        p_run.font.color.rgb = RGBColor(90, 160, 90)
        p.paragraph_format.space_after = Pt(2)
        
        p2 = doc.add_paragraph(detail)
        for run in p2.runs:
            run.font.size = Pt(9.5)
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(8)
    
    # ========================================================================
    # ANALYSE DÉTAILLÉE PAR QUADRANT - THREATS
    # ========================================================================
    add_colored_heading(doc, "ANALYSE DÉTAILLÉE - THREATS", level=2,
                       color=RGBColor(200, 100, 100))
    
    threats_detail = [
        ("Irrigation Equipment Giants (Netafim, Valmont, Rain)",
         "Netafim revenue 1.1B€ • Valmont revenue 3.8B€ • Possèdent 80%+ marché hardware irrigation existant. R&D 50M€+/an. Peuvent développer produit similaire 12-18 mois avec internal resources. Distribution channel entrenched 40+ ans. Risk predatory pricing ou acquisition hostile."),
        
        ("Startups Bien Financées (>50M€ levées, >10x WS)",
         "Ecorobotix $100M levée • Agworld $50M levée • Encepta $20M levée entrent marché irrigation 2026-2027. VC backing massive. Budget marketing/sales/R&D 5-10x vs WaterSense. Risk undercut pricing. Talent poaching ressources humaines limitées."),
        
        ("Réglementation Eau Plus Stricte (Improbable but possible)",
         "Scenario A: Taxes eau augmentent 3x 2027+ (budget farms réduit) • Scenario B: Restrictions irrigation strictes (ROI eau moins urgent) • Scenario C: Obligation temps réel data (adoption forçée mais friction). Probabilité 15-20% scenarios B/C. Impact: -30% TAM worst-case."),
        
        ("Récession Agriculture Possible (Cycle Commodity)",
         "10% farm closures 2024-2025 documentées. Budget technologie -40% si rentabilité margin <5%. Prix céréales <200€/t = faillites. Survival mode vs growth mode investissement tech. Cycle commodity-driven independent WaterSense (prix eau, climat, geopolitique)."),
        
        ("Consolidation Marché (M&A Wave Agritech)",
         "15+ acquisitions agritech 2022-2025 documentées. Netafim acquiert ~3 startups/an. Acquéreurs: Netafim, Valmont, BASF, Bayer, John Deere. Vague consolidation Q3-Q4 2026 probable. 2-3 competiteurs acquired = landscape reshuffles. Acquisition offer vs independence trade-off.")
    ]
    
    for titre, detail in threats_detail:
        p = doc.add_paragraph()
        p_run = p.add_run(f"● {titre}")
        p_run.font.bold = True
        p_run.font.size = Pt(11)
        p_run.font.color.rgb = RGBColor(200, 100, 100)
        p.paragraph_format.space_after = Pt(2)
        
        p2 = doc.add_paragraph(detail)
        for run in p2.runs:
            run.font.size = Pt(9.5)
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(8)
    
    doc.add_page_break()
    
    # ========================================================================
    # VERDICT
    # ========================================================================
    add_colored_heading(doc, "VERDICT FINAL", level=2)
    
    verdict = doc.add_paragraph()
    verdict.alignment = WD_ALIGN_PARAGRAPH.CENTER
    verdict_run = verdict.add_run(
        "✅ ANALYSE SWOT POSITIF NET\n\n"
        "Strengths (5/5) + Opportunities (5/5) > Weaknesses (5/5) + Threats (5/5)\n\n"
        "🎯 RECOMMANDATION: GO-TO-MARKET OFFENSIVE Q1-Q2 2026"
    )
    verdict_run.font.bold = True
    verdict_run.font.size = Pt(12)
    verdict_run.font.color.rgb = RGBColor(39, 174, 96)
    
    output_file = "ANALYSE_SWOT_TABLEAU_VISUEL_WATERSENSE_2026.docx"
    doc.save(output_file)
    return output_file

if __name__ == '__main__':
    print("╔════════════════════════════════════════════════════════════╗")
    print("║  GÉNÉRATION SWOT - TABLEAU VISUEL 2x2                     ║")
    print("║  4 Quadrants Colorés + Détails Complets                  ║")
    print("╚════════════════════════════════════════════════════════════╝")
    print()
    
    output_file = create_swot_visual()
    
    print(f"✅ SWOT tableau visuel créée : {output_file}")
    print()
    print("📊 FORMAT VISUEL:")
    print("  ┌─────────────────────────┬─────────────────────────┐")
    print("  │ STRENGTHS (Bleu)        │ WEAKNESSES (Orange)     │")
    print("  │ 5 items détaillés       │ 5 items détaillés       │")
    print("  ├─────────────────────────┼─────────────────────────┤")
    print("  │ OPPORTUNITIES (Vert)    │ THREATS (Rose)          │")
    print("  │ 5 items détaillés       │ 5 items détaillés       │")
    print("  └─────────────────────────┴─────────────────────────┘")
    print()
    print("✓ Tableau 2x2 rempli avec les 20 items")
    print("✓ Couleurs distinctes par quadrant")
    print("✓ + Pages détaillées par dimension")
    print("✓ Verdict final")
    print()
    print("════════════════════════════════════════════════════════════")
    print("✅ SWOT TABLEAU VISUEL GÉNÉRÉ")
    print("════════════════════════════════════════════════════════════")
