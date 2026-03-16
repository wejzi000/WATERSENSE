#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ANALYSE SWOT WATERSENSE 2026 - TABLEAU GRAND COMPLET
Un seul tableau exhaustif avec TOUTES les infos
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_colored_heading(doc, text, level=1, color=None):
    if color is None:
        color = RGBColor(46, 125, 50)
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.size = Pt(26) if level == 1 else Pt(18) if level == 2 else Pt(14)
        run.font.bold = True
        run.font.color.rgb = color
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(12)
    return heading

def create_swot_grand_tableau():
    """Analyse SWOT dans UN seul grand tableau"""
    doc = Document()
    
    # Marges
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    
    # ========================================================================
    # COUVERTURE
    # ========================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("MATRICE SWOT COMPLÈTE")
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(46, 125, 50)
    
    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("WaterSense 2026")
    subtitle_run.font.size = Pt(40)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(13, 71, 161)
    
    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = desc.add_run("Tableau Grand Format - Tous les Détails")
    desc_run.font.size = Pt(14)
    desc_run.font.italic = True
    desc_run.font.color.rgb = RGBColor(76, 175, 80)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ========================================================================
    # GRAND TABLEAU SWOT 2x2
    # ========================================================================
    add_colored_heading(doc, "MATRICE SWOT 2x2 - TABLEAU COMPLET", level=2)
    
    # Create main 2x2 table
    main_table = doc.add_table(rows=3, cols=2)
    main_table.style = 'Table Grid'
    
    # Set column widths
    for row in main_table.rows:
        for cell in row.cells:
            cell.width = Inches(3.5)
    
    # ========================================================================
    # ROW 1: Headers (INTERNE / EXTERNE)
    # ========================================================================
    header_left = main_table.rows[0].cells[0]
    header_left_para = header_left.paragraphs[0]
    header_left_para.text = "FACTEURS INTERNES"
    for run in header_left_para.runs:
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 255, 255)
    header_left_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '1565C0')
    header_left._element.get_or_add_tcPr().append(shading)
    
    header_right = main_table.rows[0].cells[1]
    header_right_para = header_right.paragraphs[0]
    header_right_para.text = "FACTEURS EXTERNES"
    for run in header_right_para.runs:
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 255, 255)
    header_right_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '1565C0')
    header_right._element.get_or_add_tcPr().append(shading)
    
    # ========================================================================
    # ROW 2: FORCES (Vert) et OPPORTUNITÉS (Vert)
    # ========================================================================
    
    # FORCES cell
    forces_cell = main_table.rows[1].cells[0]
    forces_cell.text = ""  # Clear default
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'E8F8F5')  # Light green
    forces_cell._element.get_or_add_tcPr().append(shading)
    
    # Add forces content
    forces_heading = forces_cell.add_paragraph()
    forces_heading_run = forces_heading.add_run("🟢 FORCES (POSITIF)")
    forces_heading_run.font.bold = True
    forces_heading_run.font.size = Pt(11)
    forces_heading_run.font.color.rgb = RGBColor(39, 174, 96)
    forces_heading.paragraph_format.space_after = Pt(6)
    
    forces_content = [
        ("✅ Technologie Propriétaire IA (Brevet)",
         "Brevet INPI 2025 • Algo IA 94% précision • Propriété IP 20 ans"),
        
        ("✅ Équipe Expérimentée",
         "CTO 12 ans IoT/IA • CEO 8 ans startup agritech • Head Ag 15 ans terrain"),
        
        ("✅ LTV/CAC 56:1 (Excellent)",
         "LTV 8,400€ • CAC 150€ • Payback 2.7 mois • Churn 3%/an"),
        
        ("✅ Modèle SaaS Rentable",
         "Prix 2,800-3,200€/an • Marge brute 82% • Revenu prévisible 36+ mois"),
        
        ("✅ ISO/RGPD Compliant",
         "ISO 27001 PASSED 2025 • RGPD conforme • Données AWS eu-west-1")
    ]
    
    for title, details in forces_content:
        p = forces_cell.add_paragraph()
        p_run = p.add_run(title)
        p_run.font.bold = True
        p_run.font.size = Pt(9)
        p_run.font.color.rgb = RGBColor(39, 174, 96)
        p.paragraph_format.space_after = Pt(2)
        
        p2 = forces_cell.add_paragraph()
        p2_run = p2.add_run(f"  → {details}")
        p2_run.font.size = Pt(8)
        p2_run.font.italic = True
        p2.paragraph_format.space_after = Pt(4)
    
    # OPPORTUNITÉS cell
    opps_cell = main_table.rows[1].cells[1]
    opps_cell.text = ""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'E8F8F5')  # Light green
    opps_cell._element.get_or_add_tcPr().append(shading)
    
    opps_heading = opps_cell.add_paragraph()
    opps_heading_run = opps_heading.add_run("🟢 OPPORTUNITÉS (POSITIF)")
    opps_heading_run.font.bold = True
    opps_heading_run.font.size = Pt(11)
    opps_heading_run.font.color.rgb = RGBColor(39, 174, 96)
    opps_heading.paragraph_format.space_after = Pt(6)
    
    opps_content = [
        ("○ Pénurie Eau Accélère Adoption",
         "Allocations -40% zones méditerranéennes • Crises 70%+ probabilité • ROI eau urgent"),
        
        ("○ Subventions CAP 2023-2027",
         "Bonus +25% si irrigation raisonnée • WaterSense déjà conforme • 2.5K-6K€/exploitation"),
        
        ("○ Partenariats Coopératives (50 Cibles)",
         "280K exploitations adhérentes • Revenue share 20/80 • Distribution scale Q4 2026"),
        
        ("○ Expansion Europe 2027",
         "Espagne 3.5M ha + Italie 2.8M ha • TAM 850M€ vs France • Même régulation EU"),
        
        ("○ Acquisition Gros Players",
         "Exit 50-150M€ realistic 2027-2029 • Netafim, Valmont, John Deere acquéreurs • 5x-10x CAGR")
    ]
    
    for title, details in opps_content:
        p = opps_cell.add_paragraph()
        p_run = p.add_run(title)
        p_run.font.bold = True
        p_run.font.size = Pt(9)
        p_run.font.color.rgb = RGBColor(39, 174, 96)
        p.paragraph_format.space_after = Pt(2)
        
        p2 = opps_cell.add_paragraph()
        p2_run = p2.add_run(f"  → {details}")
        p2_run.font.size = Pt(8)
        p2_run.font.italic = True
        p2.paragraph_format.space_after = Pt(4)
    
    # ========================================================================
    # ROW 3: FAIBLESSES (Rouge) et MENACES (Rouge)
    # ========================================================================
    
    # FAIBLESSES cell
    weak_cell = main_table.rows[2].cells[0]
    weak_cell.text = ""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'FADBD8')  # Light red
    weak_cell._element.get_or_add_tcPr().append(shading)
    
    weak_heading = weak_cell.add_paragraph()
    weak_heading_run = weak_heading.add_run("🔴 FAIBLESSES (NÉGATIF)")
    weak_heading_run.font.bold = True
    weak_heading_run.font.size = Pt(11)
    weak_heading_run.font.color.rgb = RGBColor(231, 76, 60)
    weak_heading.paragraph_format.space_after = Pt(6)
    
    weak_content = [
        ("❌ Startup Jeune, Poco Brand Awareness",
         "0% brand awareness • 2.5 ans existence • 300K exploitations ne connaissent pas"),
        
        ("❌ Capital Limité vs Gros Players",
         "Levée 500K€ vs Netafim 1B€+ • Budget marketing 60K€/an vs 5M€+ concurrents"),
        
        ("❌ Réseau Partenaires Initial",
         "0 contrats coopératives signés • 0 distributeurs agritech • Dépendance early-mover direct"),
        
        ("❌ Capacité Production Démarrage",
         "Infrastructure AWS single region • Support 2 FTE pour 100+ clients • SLA gap 99.5% vs 99.9%"),
        
        ("❌ Pas Historique Commercial/Références",
         "10 pilotes court terme • 0 clients >12 mois • Pas case studies/testimonials")
    ]
    
    for title, details in weak_content:
        p = weak_cell.add_paragraph()
        p_run = p.add_run(title)
        p_run.font.bold = True
        p_run.font.size = Pt(9)
        p_run.font.color.rgb = RGBColor(231, 76, 60)
        p.paragraph_format.space_after = Pt(2)
        
        p2 = weak_cell.add_paragraph()
        p2_run = p2.add_run(f"  → {details}")
        p2_run.font.size = Pt(8)
        p2_run.font.italic = True
        p2.paragraph_format.space_after = Pt(4)
    
    # MENACES cell
    threat_cell = main_table.rows[2].cells[1]
    threat_cell.text = ""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'FADBD8')  # Light red
    threat_cell._element.get_or_add_tcPr().append(shading)
    
    threat_heading = threat_cell.add_paragraph()
    threat_heading_run = threat_heading.add_run("🔴 MENACES (NÉGATIF)")
    threat_heading_run.font.bold = True
    threat_heading_run.font.size = Pt(11)
    threat_heading_run.font.color.rgb = RGBColor(231, 76, 60)
    threat_heading.paragraph_format.space_after = Pt(6)
    
    threat_content = [
        ("● Irrigation Giants (Netafim, Valmont, Rain)",
         "80%+ marché • Revenue 1B€-3B€ • Peuvent lancer produit similaire 12-18 mois"),
        
        ("● Startups Bien Financées (3+ Levées)",
         "Ecorobotix $100M • Agworld $50M • Budget marketing/sales x5-10 vs WaterSense"),
        
        ("● Réglementation Eau Plus Stricte",
         "Taxes eau possibles x3 • Restrictions irrigation scenarios • Probabilité 15-20%"),
        
        ("● Récession Agriculture",
         "10% farm closures 2024-2025 • Budget tech -40% si margin <5% • Cycle commodity"),
        
        ("● Consolidation Marché (M&A Wave)",
         "15+ acquisitions agritech 2022-2025 • Netafim acquiert 3 startups/an • Reshuffles 2026-2027")
    ]
    
    for title, details in threat_content:
        p = threat_cell.add_paragraph()
        p_run = p.add_run(title)
        p_run.font.bold = True
        p_run.font.size = Pt(9)
        p_run.font.color.rgb = RGBColor(231, 76, 60)
        p.paragraph_format.space_after = Pt(2)
        
        p2 = threat_cell.add_paragraph()
        p2_run = p2.add_run(f"  → {details}")
        p2_run.font.size = Pt(8)
        p2_run.font.italic = True
        p2.paragraph_format.space_after = Pt(4)
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ========================================================================
    # SCORING TABLE
    # ========================================================================
    add_colored_heading(doc, "SCORING & VERDICT", level=2)
    
    scoring_table = doc.add_table(rows=8, cols=4)
    scoring_table.style = 'Light Grid Accent 1'
    
    # Header
    header_row = scoring_table.rows[0]
    headers = ["Dimension", "Score", "Détail", "Évaluation"]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1565C0')
        cell._element.get_or_add_tcPr().append(shading)
    
    # Data rows
    score_data = [
        ["🟢 FORCES", "8/10", "5 atouts majeurs (tech, équipe, economics)", "✅ Excellente base"],
        ["🟢 OPPORTUNITÉS", "9/10", "5 tailwinds majeurs (eau, CAP, coopératives)", "✅ Très favorable"],
        ["🔴 FAIBLESSES", "6/10", "5 défis opérationnels mais mitigables", "⚠️ Manageable"],
        ["🔴 MENACES", "6/10", "5 risques élevés mais pas bloquants", "⚠️ Require mitigation"],
        ["", "", "", ""],
        ["📊 BILAN NET", "7.25/10", "Opportunités > Faiblesses (9+8 vs 6+6)", "✅ POSITIF"],
        ["🎯 RECOMMANDATION", "GO", "Stratégie OFFENSIVE (SO) + mitigation (ST)", "MARKET NOW"],
    ]
    
    for row_idx in range(1, len(score_data) + 1):
        row = scoring_table.rows[row_idx]
        data = score_data[row_idx - 1]
        
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.text = text
            
            # Color rows
            if row_idx in [1, 2]:  # Green rows (Forces & Opportunités)
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'E8F8F5')
                cell._element.get_or_add_tcPr().append(shading)
            elif row_idx in [3, 4]:  # Red rows (Faiblesses & Menaces)
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'FADBD8')
                cell._element.get_or_add_tcPr().append(shading)
            elif row_idx in [6, 7]:  # Gold rows (Final)
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'FDEBD0')
                cell._element.get_or_add_tcPr().append(shading)
            
            # Format text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                if col_idx == 0 or row_idx in [6, 7]:
                    run.font.bold = True
            
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ========================================================================
    # STRATÉGIES PRIORITAIRES
    # ========================================================================
    add_colored_heading(doc, "STRATÉGIES D'ACTION", level=2)
    
    strategies_text = """
🎯 Stratégie SO (Maximiser Forces + Opportunités):
   • Leverager tech IA propriétaire vs partenaires coopératives
   • Subvention CAP = core value prop marketing (bonus +25% explicite)
   • Équipe contacte 10 top coopératives Q2 2026

🛡️ Stratégie WO (Mitiger Faiblesses via Opportunités):
   • Pénurie eau = brand awareness accelerator (PR, webinars)
   • Coopératives distribution = pas besoin capital marketing massif
   • Partenaires = références clients (mitigation 'no track record')

⚔️ Stratégie ST (Défendre contre Menaces):
   • Netafim entry: Coopératives partnerships avant concurrence
   • Startups VC-backed: Local support français = différenciation
   • Récession agri: SaaS flexible (payment plans, no lock-in)
   • Consolidation: Rester attractif acquisition 50-150M€ valuation

🚀 Stratégie OT (Positionnement Opportuniste):
   • Crises hydrique = urgency narratif marketing
   • CAP subvention = direct ROI calculator vs concurrents
   • Exit strategy: 2027-2029 acquisition realistic
   • Europe 2027 avec proven model France
"""
    
    for line in strategies_text.strip().split('\n'):
        p = doc.add_paragraph(line.strip())
        if '🎯' in line or '🛡️' in line or '⚔️' in line or '🚀' in line:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        else:
            for run in p.runs:
                run.font.size = Pt(9)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ========================================================================
    # FINAL VERDICT
    # ========================================================================
    add_colored_heading(doc, "VERDICT FINAL", level=2)
    
    verdict_table = doc.add_table(rows=1, cols=1)
    verdict_cell = verdict_table.rows[0].cells[0]
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'FDEBD0')
    verdict_cell._element.get_or_add_tcPr().append(shading)
    
    verdict_para = verdict_cell.paragraphs[0]
    verdict_run = verdict_para.add_run(
        "✅ POSITIF NET - SCORE 7.25/10\n\n"
        "Forces (8/10) + Opportunités (9/10) = 17/20 ✅\n"
        "Faiblesses (6/10) + Menaces (6/10) = 12/20 ⚠️\n\n"
        "Avantage net: +5 points = GO FOR MARKET Q1-Q2 2026\n\n"
        "Actions prioritaires:\n"
        "• Coopératives partnerships (5 signatures)\n"
        "• CAP subvention marketing campaign\n"
        "• Series A financement Q2 2026\n"
        "• Équipe scaling (support +3 FTE, sales +2 FTE)"
    )
    verdict_run.font.size = Pt(11)
    verdict_run.font.bold = True
    verdict_run.font.color.rgb = RGBColor(39, 174, 96)
    
    output_file = "ANALYSE_SWOT_TABLEAU_COMPLET_WATERSENSE_2026.docx"
    doc.save(output_file)
    return output_file

if __name__ == '__main__':
    print("╔════════════════════════════════════════════════════════════╗")
    print("║  GÉNÉRATION SWOT - TABLEAU GRAND COMPLET                  ║")
    print("║  UN seul tableau avec TOUTES les infos                    ║")
    print("╚════════════════════════════════════════════════════════════╝")
    print()
    
    output_file = create_swot_grand_tableau()
    
    print(f"✅ SWOT tableau grand créée : {output_file}")
    print()
    print("📊 TABLEAU 2x2 COMPLET :")
    print("  ├─ FORCES (5 items détaillés) ────┬─ OPPORTUNITÉS (5 items)")
    print("  │                                  │")
    print("  ├─ FAIBLESSES (5 items)           ├─ MENACES (5 items)")
    print("  └──────────────────────────────────┘")
    print()
    print("✓ Tous les détails dans les cellules")
    print("✓ Couleurs: VERT (positif) | ROUGE (négatif)")
    print("✓ Scoring & stratégies inclus")
    print()
    print("════════════════════════════════════════════════════════════")
    print("✅ SWOT TABLEAU GRAND GÉNÉRÉ")
    print("════════════════════════════════════════════════════════════")
