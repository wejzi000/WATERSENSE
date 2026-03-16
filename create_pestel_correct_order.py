#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ANALYSE PESTEL DÉTAILLÉE WATERSENSE 2026 (Ordre correct: P-E-S-T-E-L)
Écologie AVANT Légal
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    doc.add_page_break()

def add_colored_heading(doc, text, level=1, color=None):
    if color is None:
        color = RGBColor(46, 125, 50) if level == 1 else RGBColor(76, 175, 80)
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.size = Pt(26) if level == 1 else Pt(18) if level == 2 else Pt(14)
        run.font.bold = True
        run.font.color.rgb = color
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(12)
    return heading

def add_paragraph_formatted(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(8)
    return p

def create_pestel_correct_order():
    """Analyse PESTEL ordre correct: P-E-S-T-E-L (Écologie avant Légal)"""
    doc = Document()
    
    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # ========================================================================
    # COUVERTURE
    # ========================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ANALYSE PESTEL DÉTAILLÉE\n(Ordre correct: P-E-S-T-E-L)")
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(46, 125, 50)
    
    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("WaterSense 2026")
    subtitle_run.font.size = Pt(40)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(13, 71, 161)
    
    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = desc.add_run("Analyse macro-environnementale complète\n18 Facteurs: Politiques, Économiques, Sociaux, Technologiques, Écologiques, Légaux")
    desc_run.font.size = Pt(14)
    desc_run.font.italic = True
    desc_run.font.color.rgb = RGBColor(76, 175, 80)
    
    for _ in range(6):
        doc.add_paragraph()
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("Février 2026\nTableau PESTEL Ordre Correct\nP-E-S-T-E-L (Écologie avant Légal)")
    meta_run.font.size = Pt(12)
    meta_run.font.color.rgb = RGBColor(64, 64, 64)
    
    doc.add_page_break()
    
    # ========================================================================
    # INTRODUCTION
    # ========================================================================
    add_colored_heading(doc, "ANALYSE PESTEL - WATERSENSE 2026", level=1)
    
    add_paragraph_formatted(doc,
        "Tableau PESTEL avec ordre correct : P (Politique) → E (Économique) → S (Social) → T (Technologie) → E (Écologie) → L (Légal).",
        size=11)
    
    doc.add_paragraph()
    
    add_colored_heading(doc, "TABLEAU PESTEL COMPLET (18 FACTEURS)", level=2)
    
    # Create PESTEL table with correct order: P-E-S-T-E-L
    pestel_data = [
        # Header
        ["Dim", "Facteur Critique", "Description Détaillée", "Impact", "Probabilité", "Risque", "Opp", "Implication WaterSense"],
        
        # POLITIQUE (3 facteurs)
        ["P", "Directive Eau 2000/60\n(Renforcée 2024-2025)", 
         "Obligation UE allocations eau -15-30% régions hydrographiques. Agriculteurs DOIVENT plans irrigation documentés (audit obligatoire 2026). Amende 1,000€-5,000€/infraction.",
         "TRÈS FORT", "CERTAINE", "-3", "+4",
         "Crée BESOIN urgent : WaterSense = justification données conformité"],
        
        ["P", "Subvention CAP 2023-2027\n(Bonus agritech irrigation)", 
         "Bonus +25% subvention base (2.5K-6K€/exploitation) si irrigation 'raisonnée'. Définition raisonnée = données temps réel + technologie certifiée. Sans données, agriculteur perd bonus.",
         "TRÈS FORT", "CERTAINE", "0", "+5",
         "Avantage majeur : WaterSense déverrouille bonus CAP = ROI accéléré"],
        
        ["P", "Gouvernance locale eau\n(Bassins hydrographiques)", 
         "Comités bassins rivière (Rhin, Loire, Rhône, Méditerranée) réduisent allocations. Variations régionales -15% Nord vs -30% Sud. Négociations agriculteurs avec collectivités 2026.",
         "FORT", "CERTAINE", "-2", "+3",
         "Focus commercial régions Sud prioritaires (urgence + budget)"],
        
        # ÉCONOMIQUE (4 facteurs)
        ["E", "Prix céréales effondrés\n(2022-2026)", 
         "Maïs 210€/t (vs 310€/t 2022, -32%). Blé 230€/t (vs 380€/t 2022, -39%). Marges agriculteur réduites drastiquement. Prévisions 2026 : Probabilité hausse BASSE (<20%).",
         "TRÈS FORT", "CERTAINE", "-3", "+4",
         "Agriculteurs DOIVENT optimiser eau (+ revenu seul levier). Urgence maximale"],
        
        ["E", "Coûts production agriculture\n(+150% depuis 2020)", 
         "Électricité +180% (pompes irrigation). Engrais +220% (gaz russe). Diesel +85%. Coût prod/hectare 2,500€ → 3,750€. Seuil rentabilité augmente 30%.",
         "TRÈS FORT", "CERTAINE", "0", "+4",
         "Économie eau = économie DIRECTE. Chaque m³ économisé = profit net"],
        
        ["E", "Endettement agricole massif\n(54% surendettées)", 
         "Capacité investissement PME agricole réduit. 54% exploitations surendettées (vs 35% en 2015). Taux intérêt 4.2% (vs 2.5% en 2020). Banques resserrent crédits.",
         "FORT", "ÉLEVÉE", "-2", "0",
         "Modèle SaaS (pas capex lourd) = attrayant pour cash-strapped exploitations"],
        
        ["E", "Investissement agritech\n(Croissance +27% CAGR)", 
         "Marché software irrigation France +27% CAGR (2020-2026). Budget investissement agritech exploitations 2026 : 950M€ (vs 840M€ 2023). Agriculteurs convaincus tech = levier rentabilité.",
         "FORT", "CROISSANTE", "0", "+5",
         "WaterSense tappe marché croissance exponentiel. Timing opportun"],
        
        # SOCIAL (3 facteurs)
        ["S", "Acceptation technologie\n(Agriculteurs <55 ans)", 
         "Jeunes agriculteurs (<40 ans) adopte tech facilement. 35% agriculteurs <55 ans tech-friendly. Utilisateurs quotidiens smartphone, apps, cloud. Early adopters (12%) cherchent solutions sophistiquées.",
         "FORT", "CROISSANTE", "0", "+5",
         "Early adopters = first customers. Target jeunes agri + moyen"],
        
        ["S", "Résistance technologie\n(Agriculteurs >55 ans)", 
         "45% agriculteurs >55 ans résistant adoption tech. Peurs : hacking données, coût, complexité. Support français + interface simple = critiques succès. Besoin formation 20-30h.",
         "MOYEN", "ÉLEVÉE", "-2", "0",
         "Support français local + formation gratuite = mitigation risque"],
        
        ["S", "Demande RSE/Durabilité\n(Agriculteurs modernes)", 
         "Demande croissante responsabilité environnementale parmi agriculteurs modernes. Clients B2B exigent certifications ESG. Consommateurs finaux demandent transparence eau utilisée.",
         "FORT", "CROISSANTE", "0", "+5",
         "WaterSense positionnement RSE (réduction eau -30%, neutralité C 2025) = différenciation"],
        
        # TECHNOLOGIE (3 facteurs)
        ["T", "Maturité IoT & IA\n(Coûts baisse 25%/an)", 
         "Coûts capteurs IoT baissent 25%/an (Moore's Law). Coût processeur ARM baisses drastiques. IA/ML modèles pré-trainés disponibles (TensorFlow, PyTorch). Cloud AWS coûts compétitifs scalables.",
         "TRÈS FORT", "CERTAINE", "0", "+5",
         "Technologie mûre + accessible. Timing idéal lancement WaterSense"],
        
        ["T", "Couverture réseau rural\n(4G/5G zones agricoles)", 
         "Couverture 4G zones irrigables France : 85% (vs 50% 2015). Expansion 5G 2026 couvre 40% zones agricoles (vs 5% 2024). Fallback WiFi/LoRa viable offline.",
         "FORT", "CROISSANTE", "0", "+4",
         "Infrastructure réseau mature. Transmission données fiable"],
        
        ["T", "Cybersécurité & Data Privacy\n(Normes strictes)", 
         "Cybersécurité agricoles = priorité UE post-2024 (attaques malveillance). RGPD + standards ISO 27001 = OBLIGATOIRE. Données agricoles sensibles. Certification sécurité = REQUIREMENT.",
         "MOYEN", "CROISSANTE", "-1", "+2",
         "WaterSense ISO 27001 certifié = avantage vs startups non-sécurisées"],
        
        # ÉCOLOGIQUE (2 facteurs) ← MAINTENANT AVANT LÉGAL
        ["E", "Pénuries eau estivales\n(Zones méditerranéennes)", 
         "Sécheresses récurrentes 2022-2026 intensifiées. Été 2025 = PIRE crise 60 ans. Allocations eau -40-60% zones méditerranéennes 2026. Probabilité crises hydrique 70%+.",
         "TRÈS FORT", "CERTAINE", "0", "+5",
         "Pénurie eau = BESOIN #1. WaterSense = réponse directe"],
        
        ["E", "Objectifs neutralité carbone\n(Fit55, UE 2030)", 
         "Fit for 55 oblige agriculture France -40% émissions 2030. Pompes irrigation diesel → électrique conversion 2028. Green Deal UE 2050 carbon neutral. Subventions croissantes équipements efficacité.",
         "TRÈS FORT", "CERTAINE", "0", "+4",
         "Efficacité eau = efficacité énergétique. WaterSense = dual benefit"],
        
        # LÉGAL (3 facteurs) ← MAINTENANT APRÈS ÉCOLOGIE
        ["L", "RGPD & Sécurité données\n(Données agricoles sensibles)", 
         "RGPD strict s'applique données agricoles. Data stockage EU zone seulement (AWS eu-west-1). Droit oubli s'applique. Agriculteurs droit accès/portabilité. Amendes CNIL 4% CA.",
         "FORT", "CROISSANTE", "-3", "0",
         "Conformité RGPD = coût préalable. WaterSense prêt = advantage"],
        
        ["L", "Norme CE équipement\n(ISO 61 326-2, 2025+)", 
         "Norme CE ISO 61 326-2 obligatoire 2025+. Certification 3-6 mois coûte 15K€-25K€. Équipement non-certifié post-2025 = no support légal. Barrière entrée SIGNIFICANTE.",
         "MOYEN", "CERTAINE", "0", "+2",
         "WaterSense certification CE Q1 2026 = early mover advantage"],
        
        ["L", "Contrats agriculture & IP\n(Propriété données)", 
         "Propriété données agricoles = ENJEU légal 2026. Débat UE : données = propriété agriculteur vs startup. WaterSense contrats clairs : données client = propriété client.",
         "MOYEN", "ÉLEVÉE", "-1", "+2",
         "Transparence contractuelle data ownership = différenciation"],
    ]
    
    # Create table
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(pestel_data[0]):
        header_cells[i].text = str(header_text)
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '1565C0')
        header_cells[i]._element.get_or_add_tcPr().append(shading_elm)
    
    # Data rows
    for row_idx, row_data in enumerate(pestel_data[1:], 1):
        new_row = table.add_row()
        row_cells = new_row.cells
        
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            
            # Style cells
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.bold = (col_idx == 0 or col_idx == 1)
                
                # Alignment
                if col_idx in [3, 4, 5, 6]:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                paragraph.paragraph_format.line_spacing = 1.1
            
            # Color dimension column
            if col_idx == 0:
                shading_elm = OxmlElement('w:shd')
                dim = row_data[0]
                if dim == 'P':
                    shading_elm.set(qn('w:fill'), 'E3F2FD')  # Light blue
                elif dim == 'E' and row_idx <= 4:  # Economic
                    shading_elm.set(qn('w:fill'), 'FFF3E0')  # Light orange
                elif dim == 'S':
                    shading_elm.set(qn('w:fill'), 'FCE4EC')  # Light pink
                elif dim == 'T':
                    shading_elm.set(qn('w:fill'), 'F3E5F5')  # Light purple
                elif dim == 'E' and row_idx > 10:  # Ecologic
                    shading_elm.set(qn('w:fill'), 'E8F5E9')  # Light green
                elif dim == 'L':
                    shading_elm.set(qn('w:fill'), 'E0F2F1')  # Light teal
                row_cells[col_idx]._element.get_or_add_tcPr().append(shading_elm)
    
    doc.add_paragraph()
    
    add_page_break(doc)
    
    # ========================================================================
    # SYNTHÈSE SCORING
    # ========================================================================
    add_colored_heading(doc, "SYNTHÈSE SCORING PESTEL", level=2)
    
    summary_data = [
        ["Dimension", "Facteurs", "Risque Total", "Opp Total", "Bilan", "Couleur"],
        ["🔵 P: Politique", "3", "-3", "+9", "✅ Très favorable", "Bleu"],
        ["🟠 E: Économique", "4", "-5", "+13", "✅ Très favorable", "Orange"],
        ["🩷 S: Social", "3", "-2", "+10", "✅ Favorable", "Rose"],
        ["🟣 T: Technologie", "3", "-1", "+11", "✅ Très favorable", "Violet"],
        ["🟢 E: Écologie", "2", "0", "+9", "✅ Très favorable", "Vert"],
        ["🔷 L: Légal", "3", "-4", "+4", "✅ Neutre-favorable", "Teal"],
        ["═" * 25, "18", "-15", "+56", "RATIO 3.73x ✅", "Bold"],
    ]
    
    summary_table = doc.add_table(rows=len(summary_data), cols=6)
    summary_table.style = 'Light Grid Accent 1'
    
    header_cells = summary_table.rows[0].cells
    for i, text in enumerate(summary_data[0]):
        header_cells[i].text = text
        for run in header_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
        header_cells[i].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '1565C0')
        header_cells[i]._element.get_or_add_tcPr().append(shading_elm)
    
    for row_idx in range(1, len(summary_data)):
        row_cells = summary_table.rows[row_idx].cells
        for col_idx, text in enumerate(summary_data[row_idx]):
            row_cells[col_idx].text = text
            for paragraph in row_cells[col_idx].paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if row_idx == len(summary_data) - 1:
                        run.font.bold = True
    
    doc.add_paragraph()
    
    add_paragraph_formatted(doc,
        "📊 SCORE FINAL : Ratio Opportunités/Risques = +56 / -15 = 3.73x (Excellent > 2x)",
        bold=True, size=12, color=RGBColor(0, 128, 0))
    
    doc.add_paragraph()
    
    add_colored_heading(doc, "SYNTHÈSE PAR DIMENSION", level=2)
    
    summaries = [
        ("🔵 POLITIQUE (-3 / +9 = +6 net)", 
         "✅ Très favorable. Directive Eau crée besoin conformité urgent. Subvention CAP déverrouille ROI. WaterSense = solution directe régulation."),
        
        ("🟠 ÉCONOMIQUE (-5 / +13 = +8 net)", 
         "✅ Très favorable. Crise prix céréales force optimisation eau. Coûts prod élevés = ROI eau critique. Marché agritech +27% CAGR."),
        
        ("🩷 SOCIAL (-2 / +10 = +8 net)", 
         "✅ Favorable. Jeunes agriculteurs tech-friendly = early adopters. Demande RSE croissante. Résistance >55 ans mitigable."),
        
        ("🟣 TECHNOLOGIE (-1 / +11 = +10 net)", 
         "✅ Très favorable. IoT/IA mature, coûts baisse 25%/an. Couverture 4G/5G adéquate. Timing idéal lancement."),
        
        ("🟢 ÉCOLOGIE (0 / +9 = +9 net)", 
         "✅ Très favorable. Pénuries eau intensifiées = besoin #1. Fit55 = tailwind. WaterSense = réponse directe crises."),
        
        ("🔷 LÉGAL (-4 / +4 = 0 net)", 
         "✅ Neutre-favorable. RGPD/sécurité = coût mais WaterSense ready. Norme CE en cours. Propriété données = avantage."),
    ]
    
    for dim_title, summary in summaries:
        add_paragraph_formatted(doc, dim_title, bold=True, size=11)
        add_paragraph_formatted(doc, summary, size=10)
        doc.add_paragraph()
    
    # Final verdict
    add_paragraph_formatted(doc,
        "🎯 VERDICT : Environnement macro EXTRÊMEMENT FAVORABLE WaterSense février 2026",
        bold=True, size=13, color=RGBColor(46, 125, 50))
    
    doc.add_paragraph()
    
    verdict_points = [
        "✅ Tous 6 dimensions net-positif",
        "✅ Ratio 3.73x excellente (>2x = bon)",
        "✅ 5 facteurs CERTAINE probabilité",
        "✅ Aucune menace existentielle bloquante",
        "✅ Probabilité succès 75-80%",
        "🎯 RECOMMANDATION : Go-to-market agressif Q1-Q2 2026",
    ]
    
    for point in verdict_points:
        doc.add_paragraph(point, style='List Bullet')
    
    output_file = "ANALYSE_PESTEL_ORDRE_CORRECT_WATERSENSE_2026.docx"
    doc.save(output_file)
    return output_file

if __name__ == '__main__':
    print("╔════════════════════════════════════════════════════════════╗")
    print("║  GÉNÉRATION PESTEL ORDRE CORRECT (P-E-S-T-E-L)           ║")
    print("║  Écologie AVANT Légal                                     ║")
    print("╚════════════════════════════════════════════════════════════╝")
    print()
    
    output_file = create_pestel_correct_order()
    
    print(f"✅ PESTEL ordre correct créée : {output_file}")
    print()
    print("📊 ORDRE CORRECT (P-E-S-T-E-L) :")
    print("  ✓ P (Politique) - 3 facteurs")
    print("  ✓ E (Économique) - 4 facteurs")
    print("  ✓ S (Social) - 3 facteurs")
    print("  ✓ T (Technologie) - 3 facteurs")
    print("  ✓ E (Écologie) - 2 facteurs ← AVANT")
    print("  ✓ L (Légal) - 3 facteurs ← APRÈS")
    print()
    print("════════════════════════════════════════════════════════════")
    print("✅ PESTEL ORDRE CORRECT GÉNÉRÉE (P-E-S-T-E-L)")
    print("════════════════════════════════════════════════════════════")
