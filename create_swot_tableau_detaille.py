#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ANALYSE SWOT WATERSENSE 2026 - TABLEAU 2x2 AVEC TOUS LES DÉTAILS DEDANS
Tous les détails directement dans les 4 quadrants colorés
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_swot_tableau_detaille():
    """Tableau SWOT 2x2 avec TOUS les détails dans chaque cellule"""
    doc = Document()
    
    # Marges réduites pour maximiser l'espace
    for section in doc.sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        section.page_width = Inches(11.69)  # A4 landscape
        section.page_height = Inches(8.27)
    
    # ========================================================================
    # TITRE
    # ========================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ANALYSE SWOT - WATERSENSE 2026")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(46, 125, 50)
    title.paragraph_format.space_after = Pt(12)
    
    # ========================================================================
    # TABLEAU 2x2 AVEC DÉTAILS COMPLETS
    # ========================================================================
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    
    # Largeur des colonnes (page A4 landscape)
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(4.8)
    
    # Hauteur des lignes
    for row in table.rows:
        row.height = Inches(3.2)
    
    # ====================================================================
    # QUADRANT 1: STRENGTHS (Bleu) - Haut Gauche
    # ====================================================================
    s_cell = table.rows[0].cells[0]
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'B3D9E8')
    s_cell._element.get_or_add_tcPr().append(shading)
    s_cell.text = ""
    
    # Titre
    p_title = s_cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("STRENGTHS (S)")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    p_title.paragraph_format.space_after = Pt(4)
    
    strengths_details = [
        ("✓ Technologie IA Propriétaire", "Brevet INPI 2025 • Précision 94% vs 78% concurrents • 3+ ans R&D • Barrière entrée 4-5 ans"),
        ("✓ Équipe Expérimentée", "CTO 12 ans IoT • CEO exit 2022 (15M€) • Head Agri 15 ans INRA • CFO 50M€+ levées"),
        ("✓ LTV/CAC 56:1", "LTV 8,400€ • CAC 150€ • Payback 2.7 mois • Churn 3%/an vs 8-12% SaaS"),
        ("✓ Modèle SaaS Rentable", "2,800€/an/exploitation • Marge brute 82% • Contrats 36 mois • Revenue recurring 94%"),
        ("✓ ISO 27001 & RGPD", "Certification Deloitte Q1 2025 • DPO nommé • AWS eu-west-1 • Zero incidents")
    ]
    
    for titre, detail in strengths_details:
        p = s_cell.add_paragraph()
        run_titre = p.add_run(titre + "\n")
        run_titre.font.bold = True
        run_titre.font.size = Pt(8)
        run_titre.font.color.rgb = RGBColor(0, 51, 102)
        run_detail = p.add_run(detail)
        run_detail.font.size = Pt(7)
        run_detail.font.color.rgb = RGBColor(30, 30, 30)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
    
    # ====================================================================
    # QUADRANT 2: WEAKNESSES (Orange) - Haut Droit
    # ====================================================================
    w_cell = table.rows[0].cells[1]
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5D5B8')
    w_cell._element.get_or_add_tcPr().append(shading)
    w_cell.text = ""
    
    p_title = w_cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("WEAKNESSES (W)")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(153, 76, 0)
    p_title.paragraph_format.space_after = Pt(4)
    
    weaknesses_details = [
        ("✗ Startup Jeune (2.5 ans)", "Zero brand recognition 300K exploitations • Netafim 40+ ans • Valmont 100+ ans • Budget brand 60K€ vs 5M€+"),
        ("✗ Capital Limité", "500K€ levée vs Netafim 1B€+ • Ecorobotix 100M€ • Marketing 60K€ vs 5M€/an • Hiring 15 vs 100+"),
        ("✗ Réseau Partenaires Limité", "0 contrats coopératives signés • 0 distributeurs • 0/3,000 points vente • Growth 100% direct"),
        ("✗ Capacité Production", "AWS single region • Support 2 FTE pour 500+ clients • Onboarding 8h vs 3h target • SLA 99.5% vs 99.9%"),
        ("✗ Pas Références", "10 pilotes, 0 clients >12 mois • 0 case studies • 0 testimonials vidéo • Enterprise hesitant")
    ]
    
    for titre, detail in weaknesses_details:
        p = w_cell.add_paragraph()
        run_titre = p.add_run(titre + "\n")
        run_titre.font.bold = True
        run_titre.font.size = Pt(8)
        run_titre.font.color.rgb = RGBColor(153, 76, 0)
        run_detail = p.add_run(detail)
        run_detail.font.size = Pt(7)
        run_detail.font.color.rgb = RGBColor(30, 30, 30)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
    
    # ====================================================================
    # QUADRANT 3: OPPORTUNITIES (Vert) - Bas Gauche
    # ====================================================================
    o_cell = table.rows[1].cells[0]
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'D4E8D4')
    o_cell._element.get_or_add_tcPr().append(shading)
    o_cell.text = ""
    
    p_title = o_cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("OPPORTUNITIES (O)")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 102, 51)
    p_title.paragraph_format.space_after = Pt(4)
    
    opportunities_details = [
        ("○ Pénurie Eau Urgente", "Allocations -40-60% 2026 • Été 2025 = pire crise 60 ans • 70% récurrence 2026-2030 • Pertes 150-300€/jour"),
        ("○ Subventions CAP +25%", "Bonus 2.5K-6K€/exploitation • Budget 180M€ France • WaterSense déjà conforme • Fenêtre 2026-2027"),
        ("○ Partenariats Coopératives", "50+ coopératives = 280K exploitations • 1 partenariat = 3-5K clients • Revenue share 80/20 • CAC -75%"),
        ("○ Expansion Europe 2027", "Espagne 3.5M ha (3.5x France) • Italie 2.8M ha • TAM Europe 850M€ • Moins concurrence"),
        ("○ Exit 50-150M€ (2027-2029)", "Netafim, Valmont, BASF, John Deere acquéreurs • Premium +40-60% strategic • Sentera comp 50M€")
    ]
    
    for titre, detail in opportunities_details:
        p = o_cell.add_paragraph()
        run_titre = p.add_run(titre + "\n")
        run_titre.font.bold = True
        run_titre.font.size = Pt(8)
        run_titre.font.color.rgb = RGBColor(0, 102, 51)
        run_detail = p.add_run(detail)
        run_detail.font.size = Pt(7)
        run_detail.font.color.rgb = RGBColor(30, 30, 30)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
    
    # ====================================================================
    # QUADRANT 4: THREATS (Rose/Rouge) - Bas Droit
    # ====================================================================
    t_cell = table.rows[1].cells[1]
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5D5D5')
    t_cell._element.get_or_add_tcPr().append(shading)
    t_cell.text = ""
    
    p_title = t_cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("THREATS (T)")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(153, 0, 0)
    p_title.paragraph_format.space_after = Pt(4)
    
    threats_details = [
        ("● Giants Irrigation", "Netafim 1.1B€ • Valmont 3.8B€ • 80%+ market share • R&D 50M€+/an • Distribution 40+ ans"),
        ("● Startups VC-Backed", "Ecorobotix 100M€ • Agworld 50M€ • Budget 5-10x • Talent war • Runway 5-7 ans"),
        ("● Régulation Stricte", "Allocations -50-70% possible (prob 15%) • Taxe eau +300% (prob 10%) • Impact TAM -40%"),
        ("● Récession Agriculture", "Maïs 210€/t vs 310€/2022 • 10% fermetures 2024-2025 • Budget tech -40% survival • Prob 25%"),
        ("● Consolidation M&A", "15+ acquisitions 2022-2025 • Netafim ~3/an • Landscape reshuffle prob 40-50% • Exit window risk")
    ]
    
    for titre, detail in threats_details:
        p = t_cell.add_paragraph()
        run_titre = p.add_run(titre + "\n")
        run_titre.font.bold = True
        run_titre.font.size = Pt(8)
        run_titre.font.color.rgb = RGBColor(153, 0, 0)
        run_detail = p.add_run(detail)
        run_detail.font.size = Pt(7)
        run_detail.font.color.rgb = RGBColor(30, 30, 30)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # ========================================================================
    # VERDICT
    # ========================================================================
    verdict = doc.add_paragraph()
    verdict.alignment = WD_ALIGN_PARAGRAPH.CENTER
    v_run = verdict.add_run("✅ VERDICT: S(+18) + O(+21) = +39  |  W(-14) + T(-11) = -25  |  NET: +14 → GO-TO-MARKET Q1-Q2 2026")
    v_run.font.bold = True
    v_run.font.size = Pt(11)
    v_run.font.color.rgb = RGBColor(39, 174, 96)
    
    output_file = "ANALYSE_SWOT_TABLEAU_DETAILLE_WATERSENSE_2026.docx"
    doc.save(output_file)
    return output_file

if __name__ == '__main__':
    print("╔════════════════════════════════════════════════════════════╗")
    print("║  GÉNÉRATION SWOT - TABLEAU 2x2 AVEC DÉTAILS DEDANS        ║")
    print("║  Format visuel comme l'image avec tout le contenu         ║")
    print("╚════════════════════════════════════════════════════════════╝")
    print()
    
    output_file = create_swot_tableau_detaille()
    
    print(f"✅ Document créé : {output_file}")
    print()
    print("📊 FORMAT:")
    print("  ┌──────────────────────────┬──────────────────────────┐")
    print("  │ STRENGTHS (Bleu)        │ WEAKNESSES (Orange)      │")
    print("  │ 5 items + détails       │ 5 items + détails        │")
    print("  │ dans le quadrant        │ dans le quadrant         │")
    print("  ├──────────────────────────┼──────────────────────────┤")
    print("  │ OPPORTUNITIES (Vert)    │ THREATS (Rose)           │")
    print("  │ 5 items + détails       │ 5 items + détails        │")
    print("  │ dans le quadrant        │ dans le quadrant         │")
    print("  └──────────────────────────┴──────────────────────────┘")
    print()
    print("✓ TOUS les détails sont DANS le tableau")
    print("✓ Format A4 paysage pour maximiser l'espace")
    print("✓ 4 couleurs distinctes")
    print("✓ Verdict final en bas")
    print()
    print("════════════════════════════════════════════════════════════")
