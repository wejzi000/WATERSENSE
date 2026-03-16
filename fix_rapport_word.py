#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convertisseur Markdown → Word - Version corrigée
Crée des fichiers Word formatés à partir des fichiers Markdown
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def create_word_from_markdown_robust(md_file, docx_file):
    """Version robuste avec gestion d'erreurs des tableaux"""
    
    doc = Document()
    
    # Setup document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        if not line.strip():
            i += 1
            continue
        
        # Titre niveau 1
        if line.startswith('# '):
            text = line.replace('# ', '').strip()
            p = doc.add_heading(text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(24)
                run.font.bold = True
            i += 1
        
        # Titre niveau 2
        elif line.startswith('## '):
            text = line.replace('## ', '').strip()
            p = doc.add_heading(text, level=1)
            for run in p.runs:
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(46, 125, 50)
            i += 1
        
        # Titre niveau 3
        elif line.startswith('### '):
            text = line.replace('### ', '').strip()
            p = doc.add_heading(text, level=2)
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.bold = True
            i += 1
        
        # Titre niveau 4
        elif line.startswith('#### '):
            text = line.replace('#### ', '').strip()
            p = doc.add_heading(text, level=3)
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.bold = True
            i += 1
        
        # Tableau
        elif line.strip().startswith('| '):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('| '):
                table_lines.append(lines[i])
                i += 1
            
            if len(table_lines) > 1:  # Au moins 2 lignes (header + separator ou data)
                try:
                    # Parser la table
                    rows = []
                    for table_line in table_lines:
                        cells = table_line.strip().split('|')[1:-1]  # Remove first and last empty
                        if cells and not all(c.strip().replace('-', '').replace(':', '') == '' for c in cells):
                            rows.append([cell.strip() for cell in cells])
                    
                    if rows and len(rows[0]) > 0:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        table.style = 'Light Grid Accent 1'
                        
                        for j, row in enumerate(rows):
                            for k in range(min(len(row), len(rows[0]))):
                                table.rows[j].cells[k].text = row[k] if k < len(row) else ''
                                if j == 0:
                                    for paragraph in table.rows[j].cells[k].paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                except Exception as e:
                    # Fallback: afficher les lignes du tableau en tant que paragraphes
                    for table_line in table_lines:
                        doc.add_paragraph(table_line.strip())
            continue
        
        # Bullet point
        elif line.strip().startswith('- '):
            text = line.strip().replace('- ', '').strip()
            if text:
                p = doc.add_paragraph(text, style='List Bullet')
            i += 1
        
        # Checkbox
        elif line.strip().startswith('[ ] ') or line.strip().startswith('[x] '):
            text = line.strip().replace('[ ] ', '☐ ').replace('[x] ', '☑ ')
            if text:
                p = doc.add_paragraph(text, style='List Bullet')
            i += 1
        
        # Horizontal line
        elif line.strip() in ('---', '___', '***'):
            doc.add_paragraph('_' * 60)
            i += 1
        
        # Code block
        elif line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # Skip closing ```
            if code_lines:
                p = doc.add_paragraph('\n'.join(code_lines))
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
        
        # Paragraphe normal
        else:
            text = line.strip()
            if text and not text.startswith(('#', '|', '-', '[', '`', '*')):
                p = doc.add_paragraph(text)
                p.paragraph_format.line_spacing = 1.15
            i += 1
    
    doc.save(docx_file)
    return True

# ============================================================================
# CONVERTIR LE FICHIER PROBLÉMATIQUE
# ============================================================================

if __name__ == '__main__':
    md_file = 'RAPPORT_MARKETING_WATERSENSE_2026.md'
    docx_file = 'RAPPORT_MARKETING_WATERSENSE_2026.docx'
    
    print("╔════════════════════════════════════════════════════════════╗")
    print("║   Correction Markdown → Word (.docx)                      ║")
    print("╚════════════════════════════════════════════════════════════╝")
    print()
    
    if os.path.exists(md_file):
        print(f"🔄 Conversion robuste: {md_file}")
        try:
            if create_word_from_markdown_robust(md_file, docx_file):
                print(f"✅ Créé: {docx_file}")
        except Exception as e:
            print(f"❌ Erreur: {e}")
    else:
        print(f"⚠️  Fichier non trouvé: {md_file}")
    
    print()
    print("════════════════════════════════════════════════════════════")
    print("✅ CONVERSION TERMINÉE!")
    print("════════════════════════════════════════════════════════════")
