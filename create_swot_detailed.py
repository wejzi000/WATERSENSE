#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ANALYSE SWOT DÉTAILLÉE WATERSENSE 2026
Document professionnel avec couleurs: VERT (Forces/Opportunités) vs ROUGE (Faiblesses/Menaces)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    doc.add_page_break()

def add_colored_heading(doc, text, level=1, color=None):
    if color is None:
        color = RGBColor(46, 125, 50)
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.size = Pt(26) if level == 1 else Pt(18) if level == 2 else Pt(14)
        run.font.bold = True
        run.font.color.rgb = color
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(12)
    return heading

def add_paragraph_formatted(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(8)
    return p

def create_swot_detailed():
    """Analyse SWOT ultra-détaillée avec couleurs"""
    doc = Document()
    
    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # ========================================================================
    # COUVERTURE
    # ========================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ANALYSE SWOT DÉTAILLÉE")
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(46, 125, 50)
    
    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("WaterSense 2026")
    subtitle_run.font.size = Pt(40)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(13, 71, 161)
    
    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = desc.add_run("Analyse Stratégique Complète\nForces - Faiblesses - Opportunités - Menaces")
    desc_run.font.size = Pt(14)
    desc_run.font.italic = True
    desc_run.font.color.rgb = RGBColor(76, 175, 80)
    
    for _ in range(6):
        doc.add_paragraph()
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("Février 2026\nMatrice SWOT Exhaustive\n✅ Vert = Positif | ⚠️ Rouge = Négatif")
    meta_run.font.size = Pt(12)
    meta_run.font.color.rgb = RGBColor(64, 64, 64)
    
    doc.add_page_break()
    
    # ========================================================================
    # INTRODUCTION
    # ========================================================================
    add_colored_heading(doc, "ANALYSE SWOT - WATERSENSE 2026", level=1)
    
    intro_text = """L'analyse SWOT (Strengths, Weaknesses, Opportunities, Threats) identifie les facteurs clés internes (Forces/Faiblesses) et externes (Opportunités/Menaces) pour WaterSense en février 2026.

Cette matrice stratégique guide la planification marketing, les investissements prioritaires et les stratégies de mitigation des risques."""
    
    add_paragraph_formatted(doc, intro_text, size=11)
    doc.add_paragraph()
    
    # ========================================================================
    # MATRICE SWOT 2x2
    # ========================================================================
    add_colored_heading(doc, "MATRICE SWOT 2x2 - VUE D'ENSEMBLE", level=2)
    
    # Create 2x2 SWOT matrix
    swot_matrix = doc.add_table(rows=3, cols=2)
    swot_matrix.style = 'Table Grid'
    
    # Top-left: Empty cell
    swot_matrix.rows[0].cells[0].text = ""
    
    # Top-right: INTERNAL / EXTERNAL headers
    swot_matrix.rows[0].cells[1].text = ""
    
    # Row 1: FORCES (Green) and OPPORTUNITÉS (Green)
    forces_cell = swot_matrix.rows[1].cells[0]
    forces_para = forces_cell.paragraphs[0]
    forces_para.text = "🟢 FORCES (INTERNE - POSITIF)"
    for run in forces_para.runs:
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 255, 255)
    forces_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '27AE60')  # Green
    forces_cell._element.get_or_add_tcPr().append(shading_elm)
    
    opps_cell = swot_matrix.rows[1].cells[1]
    opps_para = opps_cell.paragraphs[0]
    opps_para.text = "🟢 OPPORTUNITÉS (EXTERNE - POSITIF)"
    for run in opps_para.runs:
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 255, 255)
    opps_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '27AE60')  # Green
    opps_cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Row 2: FAIBLESSES (Red) and MENACES (Red)
    weak_cell = swot_matrix.rows[2].cells[0]
    weak_para = weak_cell.paragraphs[0]
    weak_para.text = "🔴 FAIBLESSES (INTERNE - NÉGATIF)"
    for run in weak_para.runs:
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 255, 255)
    weak_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'E74C3C')  # Red
    weak_cell._element.get_or_add_tcPr().append(shading_elm)
    
    threat_cell = swot_matrix.rows[2].cells[1]
    threat_para = threat_cell.paragraphs[0]
    threat_para.text = "🔴 MENACES (EXTERNE - NÉGATIF)"
    for run in threat_para.runs:
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 255, 255)
    threat_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'E74C3C')  # Red
    threat_cell._element.get_or_add_tcPr().append(shading_elm)
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ========================================================================
    # FORCES (GREEN)
    # ========================================================================
    add_colored_heading(doc, "🟢 FORCES - FACTEURS INTERNES POSITIFS", level=2, 
                       color=RGBColor(39, 174, 96))
    
    forces_data = [
        {
            "titre": "✅ Technologie Propriétaire IA (Brevet)",
            "description": "WaterSense possède une technologie d'IA propriétaire brevetée qui crée une barrière d'entrée significative. L'algorithme de prédiction consommation eau utilise machine learning temps réel couplé à données IoT terrain.",
            "details": [
                "Brevet déposé auprès de l'INPI 2025",
                "Algorithme IA entraîné 3+ années de données agricoles",
                "Temps de prédiction <2 secondes",
                "Précision 94%+ vs 78% concurrents",
                "Propriété intellectuelle protégée 20 ans"
            ],
            "impact": "TRÈS HAUT - Crée avantage concurrentiel durable & defendable"
        },
        {
            "titre": "✅ Équipe Fondatrice Expérimentée",
            "description": "Équipe composée d'experts agricoles, data scientists et entrepreneurs ayant succès antérieurs dans secteur agritech & SaaS.",
            "details": [
                "CTO: 12 ans expérience IoT/IA industries",
                "CEO: 8 ans startup agritech (exit 2022)",
                "Head of Ag: Ingénieur agronome, 15 ans terrain",
                "CFO: 10 ans fintech, levées 50M€+",
                "Conseil: Investisseurs CA agricole & VCs"
            ],
            "impact": "HAUT - Exécution rapide & décisions stratégiques qualité"
        },
        {
            "titre": "✅ LTV/CAC Ratio 56:1 (Excellent)",
            "description": "Ratio Lifetime Value / Customer Acquisition Cost exceptionnellement élevé indique modèle économique très efficace & rentable.",
            "details": [
                "LTV: 8,400€ (3 ans rétention @ 2,800€/an)",
                "CAC: 150€ (partenariats coopératives)",
                "Payback period: 2.7 mois",
                "Churn: 3% annuel (excellent vs 8-12% SaaS)",
                "Benchmark: SaaS bon = 3:1, Excellent = 5:1+"
            ],
            "impact": "TRÈS HAUT - Viabilité économique longue terme garantie"
        },
        {
            "titre": "✅ Modèle Récurrent Rentable",
            "description": "Structure pricing SaaS garantit revenus prévisibles, scalables et marges brutes élevées (80%+).",
            "details": [
                "Prix: 2,800€-3,200€/exploitation/an",
                "Marge brute: 82% (cloud + support)",
                "Pas de COGS physique important",
                "Revenu prévisible 36+ mois contrats",
                "Expansion verticale facile (features)"
            ],
            "impact": "HAUT - Predictability investisseurs & croissance soutenable"
        },
        {
            "titre": "✅ Certification ISO/RGPD (Compliance Ready)",
            "description": "WaterSense déjà certifiée ISO 27001 & conforme RGPD complet = entrée marché sans frein réglementaire.",
            "details": [
                "ISO 27001: Audit externe 2025 PASSED",
                "RGPD: Données stockage AWS eu-west-1",
                "DPO (Data Protection Officer) nommé",
                "Contrats client clarifient propriété données",
                "Pas d'amendes/violations prévues"
            ],
            "impact": "MOYEN-HAUT - Différenciation vs startups non-conformes"
        }
    ]
    
    for force in forces_data:
        add_paragraph_formatted(doc, force["titre"], bold=True, size=12, 
                              color=RGBColor(39, 174, 96))
        add_paragraph_formatted(doc, force["description"], size=10)
        
        doc.add_paragraph("Détails clés:", style='List Bullet')
        for detail in force["details"]:
            p = doc.add_paragraph(detail, style='List Bullet 2')
            p.paragraph_format.left_indent = Cm(1.5)
        
        add_paragraph_formatted(doc, f"📊 Impact: {force['impact']}", italic=True, size=9,
                              color=RGBColor(52, 152, 219))
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========================================================================
    # FAIBLESSES (RED)
    # ========================================================================
    add_colored_heading(doc, "🔴 FAIBLESSES - FACTEURS INTERNES NÉGATIFS", level=2,
                       color=RGBColor(231, 76, 60))
    
    weaknesses_data = [
        {
            "titre": "❌ Startup Jeune, Poco Brand Awareness",
            "description": "WaterSense fondée 2023 (2.5 ans) = brand recognition très basse vs acteurs établis. Agriculteurs ne connaissent pas entreprise.",
            "details": [
                "0% brand awareness parmi 300K exploitations",
                "Aucun historique marketing vs Netafim (40 ans)",
                "Besoin gros budget marketing pour awareness",
                "Lead gen long (6-12 mois cycle vente agri)",
                "Word-of-mouth insufisant startup jeune"
            ],
            "impact": "MOYEN - Frein acquisition clients mais mitigable via marketing"
        },
        {
            "titre": "❌ Capital Limité vs Gros Players",
            "description": "Startup seed/Series A vs géants irrigation (Netafim $1B+, Valmont $1B+) = budget R&D/marketing/ventes incomparable.",
            "details": [
                "Levée WaterSense: 500K€ vs Netafim 1B€+",
                "Budget marketing: 60K€/an vs concurrents 5M€+",
                "R&D: 100K€/an vs gros players 50M€+",
                "Pas de trésorerie pour acquisitions",
                "Dépendance financement externe"
            ],
            "impact": "HAUT - Contraint stratégies expansion rapide"
        },
        {
            "titre": "❌ Réseau Partenaires Initial Limité",
            "description": "Peu de partenariats établis vs acteurs entrenched. Pas de distribution durable coopératives/distributeurs.",
            "details": [
                "0 contrats coopératives signature (50 cibles)",
                "0 distributeurs agritech partenaires",
                "Pas accès réseau 3,000 points vente agri",
                "Dépendance early-mover adoption directe",
                "Nécessité construire channel 0-to-1"
            ],
            "impact": "HAUT - Croissance dépendante évangélisation directe"
        },
        {
            "titre": "❌ Capacité Production Au Démarrage",
            "description": "Infrastructure cloud/serveurs pour support 1,000s clients pas encore 100% scalable. Risque outage croissance rapide.",
            "details": [
                "Infrastructure: AWS single region (besoin backup)",
                "Support: 2 FTE pour 100+ clients (besoin 5+ pour 500)",
                "Monitoring: Alertes manuelles vs automation",
                "SLA: 99.5% réalisé, 99.9% promis (gap)",
                "Onboarding: 8h/client, doit réduire 30%"
            ],
            "impact": "MOYEN - Scaling opérationnel challenge 2026"
        },
        {
            "titre": "❌ Pas Historique Commercial/Références",
            "description": "Zéro client de longue durée = sans références/testimonials. Agriculteurs hésitent acheter produit 'no track record'.",
            "details": [
                "10 pilotes 2024-2025 (court terme)",
                "Aucun client >12 mois rétention",
                "Pas case studies publiques publiées",
                "Absence testimonials vidéo agriculteurs",
                "Proof-of-concept demandes préalables (rallonge cycle)"
            ],
            "impact": "MOYEN-HAUT - Frein adoption mid-market conservative"
        }
    ]
    
    for weak in weaknesses_data:
        add_paragraph_formatted(doc, weak["titre"], bold=True, size=12,
                              color=RGBColor(231, 76, 60))
        add_paragraph_formatted(doc, weak["description"], size=10)
        
        doc.add_paragraph("Détails clés:", style='List Bullet')
        for detail in weak["details"]:
            p = doc.add_paragraph(detail, style='List Bullet 2')
            p.paragraph_format.left_indent = Cm(1.5)
        
        add_paragraph_formatted(doc, f"⚠️ Impact: {weak['impact']}", italic=True, size=9,
                              color=RGBColor(192, 57, 43))
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========================================================================
    # OPPORTUNITÉS (GREEN)
    # ========================================================================
    add_colored_heading(doc, "🟢 OPPORTUNITÉS - FACTEURS EXTERNES POSITIFS", level=2,
                       color=RGBColor(39, 174, 96))
    
    opportunities_data = [
        {
            "titre": "○ Pénurie Eau Accélère Adoption (Tailwind)",
            "description": "Crises hydrique intensifiées 2022-2026 = besoin solution optimisation eau devenu URGENT vs hypothétique. WaterSense = réponse directe problème existentiel.",
            "details": [
                "Allocations eau -40% zones méditerranéennes 2026",
                "Été 2025 = PIRE sécheresse 60 ans (crise réelle)",
                "70%+ probabilité crises récurrentes 2027-2030",
                "Agriculteurs DOIVENT réduire consommation eau",
                "SWOT impact: +3 ROI priorité vs +1 baseline"
            ],
            "potentiel": "TRÈS HAUT - Macro-trend accélère PMF"
        },
        {
            "titre": "○ Subventions CAP 2023-2027 Disponibles",
            "description": "Gouvernement français bonus +25% subvention CAP si irrigation 'raisonnée' documentée. WaterSense déverrouille accès 2.5K-6K€ bonus/exploitation.",
            "details": [
                "Bonus CAP: +25% base (2.5K-6K€/exploitation)",
                "Critères: Données temps réel + technologie certifiée",
                "WaterSense: Déjà conforme critères 2026",
                "Budget total: 180M€ CAP irrigation 2026",
                "Fenêtre: Disponible 2026-2027 avant révision CAP"
            ],
            "potentiel": "TRÈS HAUT - Direct ROI accelerator (payoff 14 mois)"
        },
        {
            "titre": "○ Partenariats Coopératives Agricoles (50 Cibles)",
            "description": "50+ coopératives agricoles France existantes = 280K+ exploitations adhérentes. Accord partenaire unique = distribution instantanée 3K-5K clients.",
            "details": [
                "Coopératives: Beauce, Champagne, Loire, Provence",
                "Adhérents: 280K exploitations (25% France irrigante)",
                "Modèle: Coopérative = reseller, WaterSense backend",
                "Revenue share: 20% coopérative, 80% WaterSense",
                "Timeline: Accord pilote Q2 2026, roll-out Q4"
            ],
            "potentiel": "TRÈS HAUT - Distribution scale via coopératives"
        },
        {
            "titre": "○ Expansion Europe 2027 (Espagne/Italie)",
            "description": "Agriculteurs Espagne & Italie face même crises eau. Réglementation EU Directive Eau identique. TAM Europe = 3x France (850M€ marché irrigation).",
            "details": [
                "Espagne: 3.5M ha irrigation (vs France 1M ha)",
                "Italie: 2.8M ha irrigation (intensif Po Valley)",
                "Marché: +250% vs France, moins concurrence tech",
                "Régulation: EU Directive Eau 2000/60 identique",
                "Entry: Localisation marketing Q3 2027 pilot"
            ],
            "potentiel": "HAUT - International scaling 2027-2028"
        },
        {
            "titre": "○ Acquisition par Gros Players (Exit Potential)",
            "description": "WaterSense profil acquisitions attractif (IA propriétaire, LTV/CAC 56:1, TAM croissance). Netafim, Valmont, John Deere cherchent tech irrigation innovation.",
            "details": [
                "Comparable exits: Sentera (precision ag) $50M",
                "WaterSense valuation: 50-150M€ (5x-10x CAGR)",
                "Acquéreurs: Netafim, Valmont, BASF, John Deere",
                "Timeline: Exit réaliste 2027-2029 (5x SVRR)",
                "IPO alternative: 2030+ si 50M€ ARR"
            ],
            "potentiel": "MOYEN-HAUT - Liquidity event 2027-2029"
        }
    ]
    
    for opp in opportunities_data:
        add_paragraph_formatted(doc, opp["titre"], bold=True, size=12,
                              color=RGBColor(39, 174, 96))
        add_paragraph_formatted(doc, opp["description"], size=10)
        
        doc.add_paragraph("Détails clés:", style='List Bullet')
        for detail in opp["details"]:
            p = doc.add_paragraph(detail, style='List Bullet 2')
            p.paragraph_format.left_indent = Cm(1.5)
        
        add_paragraph_formatted(doc, f"📈 Potentiel: {opp['potentiel']}", italic=True, size=9,
                              color=RGBColor(52, 152, 219))
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========================================================================
    # MENACES (RED)
    # ========================================================================
    add_colored_heading(doc, "🔴 MENACES - FACTEURS EXTERNES NÉGATIFS", level=2,
                       color=RGBColor(231, 76, 60))
    
    threats_data = [
        {
            "titre": "● Irrigation Equipment Giants (Netafim, Valmont, etc)",
            "description": "Géants irrigation établis (Netafim, Valmont, Rain) possèdent 80%+ marché hardware irrigation & relations distributeurs intrenched. Peuvent lancer produit similaire en 12-18 mois.",
            "details": [
                "Netafim: Revenue 1.1B€, R&D 80M€/an",
                "Valmont: Revenue 3.8B€, diverse business lines",
                "Avantage: Distribution existing + brand 40+ ans",
                "Entrée barrière: Channel access coopératives difficile",
                "Risque: Predatory pricing ou acquisition"
            ],
            "severite": "TRÈS HAUT - Menace existentielle si actif"
        },
        {
            "titre": "● Startups Bien Financées (3+ Levées)",
            "description": "Compétiteurs agritech bien-financés (Ecorobotix $100M, Agworld $50M, etc) entrent marché irrigation avec VC backing. Capital supérieur WaterSense x10-20.",
            "details": [
                "Ecorobotix: $100M levée, déploie tech EU",
                "Agworld: $50M levée, partenaires coopératives",
                "Encepta: $20M levée, sensor agri focus",
                "Avantage: Budget marketing/sales/R&D x5+",
                "Risk: Undercut pricing WaterSense"
            ],
            "severite": "HAUT - Compétition intensification"
        },
        {
            "titre": "● Réglementation Eau Plus Stricte (Improbable mais Possible)",
            "description": "Si UE rend obligation irrigation connectée SANS flexibilité implementation, création friction adoption. Ou si taxes eau élevées rendent investing technologie moins attractive.",
            "details": [
                "Scenario A: Obligation temps réel data (adoption forçée)",
                "Scenario B: Taxes eau >3x 2026 (budget farms réduit)",
                "Scenario C: Restrictions irrigation (besoin moins urgent)",
                "Probabilité: 15-20% scenario B/C",
                "Impact: -30% TAM potential worst-case"
            ],
            "severite": "MOYEN - Improbable mais high-impact"
        },
        {
            "titre": "● Récession Agriculture (Moins Investissement)",
            "description": "Si prix céréales restent <200€/t & sécheresses intensifient faillites exploitations = budget investissement tech réduit. Farmers mode survival vs growth.",
            "details": [
                "Récession agri: 10% farm closures 2024-2025",
                "Budget tech: -40% si rentabilité margin <5%",
                "Cycle: Dépend prix commodités (vs technologie)",
                "Risque: CycleDown 2026-2027 possible",
                "Mitigation: Modèle SaaS flexible (payment plans)"
            ],
            "severite": "MOYEN-HAUT - Risque macroéconomique"
        },
        {
            "titre": "● Consolidation Marché Possible (M&A Wave)",
            "description": "Vague M&A possible dans agritech irrigation (Netafim acquiert 3 startups dernière année). Consolidation réduit compétition mais aussi opportunities WaterSense indépendant.",
            "details": [
                "M&A trend: 2022-2025 = 15+ acquisitions agritech",
                "Acquéreurs: Netafim, Valmont, BASF, Bayer",
                "Scenario: 2-3 compétiteurs acquired 2026-2027",
                "Impact sur WaterSense: Acquisition offer vs independence",
                "Timeline: Consolidation wave Q3-Q4 2026"
            ],
            "severite": "MOYEN - Reshuffles marché 2026-2027"
        }
    ]
    
    for threat in threats_data:
        add_paragraph_formatted(doc, threat["titre"], bold=True, size=12,
                              color=RGBColor(231, 76, 60))
        add_paragraph_formatted(doc, threat["description"], size=10)
        
        doc.add_paragraph("Détails clés:", style='List Bullet')
        for detail in threat["details"]:
            p = doc.add_paragraph(detail, style='List Bullet 2')
            p.paragraph_format.left_indent = Cm(1.5)
        
        add_paragraph_formatted(doc, f"🚨 Sévérité: {threat['severite']}", italic=True, size=9,
                              color=RGBColor(192, 57, 43))
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========================================================================
    # SYNTHÈSE SCORING SWOT
    # ========================================================================
    add_colored_heading(doc, "SYNTHÈSE SCORING SWOT", level=2)
    
    add_paragraph_formatted(doc, "Scoring Impact par dimension (0-10):", bold=True, size=11)
    doc.add_paragraph()
    
    # Scoring table
    scoring_data = [
        ["Dimension", "Score", "Détail", "Bilan"],
        ["🟢 FORCES", "8/10", "5 atouts majeurs dont tech propriétaire & équipe", "✅ Excellente base"],
        ["🔴 FAIBLESSES", "6/10", "5 défis opérationnels mais mitigables", "⚠️ Manageable"],
        ["🟢 OPPORTUNITÉS", "9/10", "5 tailwinds majeurs (eau, CAP, coopératives)", "✅ Très favorable"],
        ["🔴 MENACES", "6/10", "5 risques élevés mais pas bloquants", "⚠️ Require mitigation"],
        ["", "", "", ""],
        ["BILAN GLOBAL", "7.25/10", "SWOT NET-POSITIF (Opportunités > Faiblesses)", "🎯 GO FOR IT"],
    ]
    
    swot_table = doc.add_table(rows=len(scoring_data), cols=4)
    swot_table.style = 'Light Grid Accent 1'
    
    for row_idx, row_data in enumerate(scoring_data):
        cells = swot_table.rows[row_idx].cells
        
        for col_idx, text in enumerate(row_data):
            cells[col_idx].text = text
            
            # Header row
            if row_idx == 0:
                for run in cells[col_idx].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '1565C0')
                cells[col_idx]._element.get_or_add_tcPr().append(shading_elm)
            
            # Color rows
            elif row_idx in [1, 3]:  # Forces & Opportunités (Green)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'E8F8F5')
                cells[col_idx]._element.get_or_add_tcPr().append(shading_elm)
                if col_idx == 0:  # Dimension col
                    for run in cells[col_idx].paragraphs[0].runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(39, 174, 96)
            
            elif row_idx in [2, 4]:  # Faiblesses & Menaces (Red)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FADBD8')
                cells[col_idx]._element.get_or_add_tcPr().append(shading_elm)
                if col_idx == 0:  # Dimension col
                    for run in cells[col_idx].paragraphs[0].runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(231, 76, 60)
            
            elif row_idx == 6:  # Final row (Gold)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FDEBD0')
                cells[col_idx]._element.get_or_add_tcPr().append(shading_elm)
                for run in cells[col_idx].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(183, 149, 11)
            
            cells[col_idx].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ========================================================================
    # STRATÉGIES PRIORITAIRES
    # ========================================================================
    add_colored_heading(doc, "STRATÉGIES PRIORITAIRES - ACTION PLAN SWOT", level=2)
    
    strategies = [
        ("🎯 Stratégie 1: Maximiser Forces + Opportunités (SO)",
         [
             "Leverager tech propriétaire IA comme core value prop pour partenaires coopératives",
             "Utiliser subventions CAP comme 'sweetener' vente (bonus +25% explicite dans pitch)",
             "Équipe expérimentée contacte directly 10 top coopératives Q2 2026",
             "Showcase PESTEL favorable (9/10 Opportunités) à investisseurs pour Series A Q3"
         ]),
        
        ("🛡️ Stratégie 2: Mitiger Faiblesses via Opportunités (WO)",
         [
             "Utiliser tailwind pénurie eau pour accélérer brand awareness (press, webinars)",
             "Modèle coopératives réduit besoin de capital marketing direct",
             "Partenaires coopératives = références clients (mitigation 'no track record')",
             "SaaS cloud = pas capex, scalable avec peu capital"
         ]),
        
        ("⚔️ Stratégie 3: Défendre contre Menaces (ST)",
         [
             "Netafim/Valmont entry: Accélérer coopératives partnerships avant concurrence",
             "Startups VC-backed: Différencier sur 'local support français' vs outsourced",
             "Récession agri: Modèle SaaS flexible (payment plans, no lock-in)",
             "Consolidation: Rester attractif acquisition target via growth 3x+ CAGR"
         ]),
        
        ("🚀 Stratégie 4: Positionnement Opportuniste (OT)",
         [
             "Crises hydrique = urgency narratif marketing (lead generation)",
             "CAP subvention bonus = direct ROI calculator vs concurrents",
             "Exit strategy: 2027-2029 acquisition realistic (50-150M€ valuation)",
             "Europe expansion Q3 2027 avec proven model France"
         ])
    ]
    
    for strat_title, strat_points in strategies:
        add_paragraph_formatted(doc, strat_title, bold=True, size=12,
                              color=RGBColor(13, 71, 161))
        for point in strat_points:
            doc.add_paragraph(point, style='List Bullet')
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========================================================================
    # VERDICT FINAL
    # ========================================================================
    add_colored_heading(doc, "VERDICT FINAL SWOT", level=2)
    
    # Create table for verdict box (better formatting)
    verdict_table = doc.add_table(rows=1, cols=1)
    verdict_cell = verdict_table.rows[0].cells[0]
    verdict_para = verdict_cell.paragraphs[0]
    verdict_run = verdict_para.add_run(
        "✅ POSITIF NET\n\n"
        "Forces (8/10) + Opportunités (9/10) >> Faiblesses (6/10) + Menaces (6/10)\n\n"
        "Score Global: 7.25/10 - GO FOR MARKET\n\n"
        "Recommandation: Stratégie OFFENSIVE (SO) avec mitigation proactive menaces (ST)"
    )
    verdict_run.font.size = Pt(12)
    verdict_run.font.bold = True
    verdict_run.font.color.rgb = RGBColor(39, 174, 96)
    
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'E8F8F5')
    verdict_cell._element.get_or_add_tcPr().append(shading_elm)
    
    doc.add_paragraph()
    
    final_rec = doc.add_paragraph()
    final_rec_run = final_rec.add_run(
        "Actions Q1-Q2 2026:\n"
        "• Coopératives partnerships (5 signatures prioritaires)\n"
        "• Subvention CAP marketing campaign (testimonials, ROI calculator)\n"
        "• Series A financement (leverager favorable SWOT)\n"
        "• Équipe scaling (support, sales) pour croissance 3x 2026"
    )
    final_rec_run.font.size = Pt(10)
    final_rec_run.font.italic = True
    
    output_file = "ANALYSE_SWOT_DETAILLEE_WATERSENSE_2026.docx"
    doc.save(output_file)
    return output_file

if __name__ == '__main__':
    print("╔════════════════════════════════════════════════════════════╗")
    print("║  GÉNÉRATION ANALYSE SWOT DÉTAILLÉE                        ║")
    print("║  Couleurs: VERT (Positif) | ROUGE (Négatif)              ║")
    print("╚════════════════════════════════════════════════════════════╝")
    print()
    
    output_file = create_swot_detailed()
    
    print(f"✅ SWOT détaillée créée : {output_file}")
    print()
    print("📊 CONTENU :")
    print("  ✓ Matrice SWOT 2x2 (Forces/Faiblesses/Opportunités/Menaces)")
    print("  ✓ 5 Forces détaillées (Vert)")
    print("  ✓ 5 Faiblesses détaillées (Rouge)")
    print("  ✓ 5 Opportunités détaillées (Vert)")
    print("  ✓ 5 Menaces détaillées (Rouge)")
    print("  ✓ Synthèse scoring (0-10 par dimension)")
    print("  ✓ 4 Stratégies prioritaires (SO/WO/ST/OT)")
    print("  ✓ Verdict final + Action plan")
    print()
    print("🎨 COULEURS :")
    print("  ✓ VERT (#27AE60) = Forces & Opportunités")
    print("  ✓ ROUGE (#E74C3C) = Faiblesses & Menaces")
    print("  ✓ Tables avec fond clair (bleu header)")
    print()
    print("════════════════════════════════════════════════════════════")
    print("✅ SWOT ULTRA-DÉTAILLÉE GÉNÉRÉE")
    print("════════════════════════════════════════════════════════════")
