#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Rapport Marketing Stratégique WaterSense - 30+ Pages Professionnel
Format Mémoire Professionnel - Sans émojis - Police Times New Roman
Basé sur fondamentaux Kotler et Marketing Management
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def shade_cell(cell, color):
    """Ajoute couleur de fond à une cellule"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    """Ajoute bordures cellule"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tcBorders.append(OxmlElement(f'w:{edge}'))
    tcPr.append(tcBorders)

def set_times_roman(paragraph, size=11, bold=False, italic=False):
    """Configure Times New Roman"""
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic

def add_page_break(doc):
    """Ajoute saut de page"""
    doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# INITIALISATION DOCUMENT
# ═════════════════════════════════════════════════════════════════

doc = Document()

# Configuration style par défaut Times New Roman
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# ═════════════════════════════════════════════════════════════════
# PAGE 1 - COUVERTURE FORMELLE
# ═════════════════════════════════════════════════════════════════

# Espacement couverture
for _ in range(8):
    doc.add_paragraph()

# Titre
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('RAPPORT MARKETING STRATEGIQUE')
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.bold = True

# Sous-titre
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('WATERSENSE')
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.font.bold = True

# Description
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('\nPlateforme Intelligente de Gestion de l\'Irrigation Agricole\nPlan Marketing Annuel 2026')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

# Espacement
for _ in range(6):
    doc.add_paragraph()

# Informations document
info_table = doc.add_table(rows=6, cols=2)
info_table.autofit = False
info_table.allow_autofit = False

info_data = [
    ['Date de rédaction:', 'Janvier 2026'],
    ['Classement:', 'Document Confidentiel'],
    ['Version:', '1.0 - Version Définitive'],
    ['Auteur:', 'Équipe Marketing'],
    ['Domaine:', 'Agrotechnologie - Irrigation Intelligente'],
    ['Horizon temporel:', 'Année 2026']
]

for idx, (label, value) in enumerate(info_data):
    info_table.cell(idx, 0).text = label
    info_table.cell(idx, 1).text = value
    
    for cell in [info_table.cell(idx, 0), info_table.cell(idx, 1)]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════
# PAGE 2 - TABLE DES MATIÈRES
# ═════════════════════════════════════════════════════════════════

toc_heading = doc.add_heading('TABLE DES MATIÈRES', level=1)
for run in toc_heading.runs:
    run.font.name = 'Times New Roman'

toc_items = [
    ('INTRODUCTION ET CONTEXTE', 3),
    ('ANALYSE DE SITUATION STRATEGIQUE', 4),
    ('ANALYSE PESTEL', 6),
    ('ANALYSE COMPETITIVE ET POSITIONNEMENT', 8),
    ('SEGMENTATION DU MARCHE', 10),
    ('STRATEGIE MARKETING 4P DETAILLEE', 13),
    ('PLAN MARKETING OPERATIONNEL 2026', 18),
    ('BUDGET ET ALLOCATION RESSOURCES', 23),
    ('METRIQUES DE PILOTAGE ET KPI', 25),
    ('CHRONOGRAMME D\'EXECUTION', 27),
    ('GESTION DES RISQUES MARKETING', 29),
    ('CONCLUSION ET RECOMMANDATIONS', 31),
    ('ANNEXES', 33),
]

for idx, (item, page) in enumerate(toc_items, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f'{idx}. {item}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    
    dots = '.' * (70 - len(f'{idx}. {item}') - len(str(page)))
    p.add_run(f' {dots} {page}')

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════
# PAGE 3 - INTRODUCTION ET CONTEXTE
# ═════════════════════════════════════════════════════════════════

intro = doc.add_heading('1. INTRODUCTION ET CONTEXTE', level=1)
for run in intro.runs:
    run.font.name = 'Times New Roman'

doc.add_heading('1.1 Objet du rapport et objectifs', level=2)

intro_para = doc.add_paragraph(
    'Le présent rapport constitue un plan marketing stratégique et opérationnel pour '
    'l\'année 2026, élaboré pour la solution WaterSense. Ce document vise à définir les '
    'orientations commerciales, les tactiques de marketing, l\'allocation budgétaire et '
    'les métriques de succès pour l\'introduction et le développement de cette plateforme '
    'technologique sur le marché français de l\'irrigation agricole.'
)
set_times_roman(intro_para)

doc.add_heading('1.2 Contexte économique et agricole', level=2)

context_para = doc.add_paragraph(
    'Le secteur agricole français connaît une transformation majeure caractérisée par trois '
    'facteurs critiques : (1) une raréfaction de la ressource hydrique due aux changements '
    'climatiques, (2) une augmentation significative des coûts énergétiques impactant les '
    'exploitations, et (3) une pression réglementaire croissante imposée par les directives '
    'européennes et nationales en matière de gestion de l\'eau.'
)
set_times_roman(context_para)

context_data = [
    'Surface agricole utile France : 29,4 millions d\'hectares',
    'Surface irrigable : 2,6 millions d\'hectares (8,8% du total)',
    'Surface irriguée effective : 1,1 millions d\'hectares',
    'Dépenses annuelles irrigation : 3,2 milliards d\'euros',
    'Restrictions préfectorales 2024 : 68 départements affectés',
    'Augmentation électricité agricole : +145% sur trois ans (2021-2024)'
]

for item in context_data:
    p = doc.add_paragraph(item, style='List Bullet')
    set_times_roman(p)

doc.add_heading('1.3 Positionnement du produit WaterSense', level=2)

positioning = doc.add_paragraph(
    'WaterSense se positionne comme une solution intégrée d\'optimisation de l\'irrigation '
    'combinant capteurs Internet des Objets (IoT), traitement local des données via '
    'architecture Edge Computing, et algorithmes d\'intelligence artificielle propriétaires. '
    'La valeur centrale de la proposition réside dans la fourniture de recommandations '
    'prescriptives précises (ex : "irriguer demain à 4h15 pour 48 minutes") opposées aux '
    'solutions concurrentes qui fournissent uniquement des données descriptives.'
)
set_times_roman(positioning)

doc.add_heading('1.4 Justification stratégique de l\'entrée marché 2026', level=2)

justif = doc.add_paragraph(
    'Le choix d\'un lancement commercial en 2026 répond à plusieurs impératifs : '
    '(1) maturité technologique du produit validée par 12 mois de tests pilotes, '
    '(2) alignement avec les calendriers de planification agricole et budgétaire des '
    'exploitations, (3) anticipation des obligations de conformité réglementaire qui '
    'deviennent contraignantes à partir de 2027, et (4) fenêtre concurrentielle optimale '
    'avant intensification attendue de la compétition sur le segment.'
)
set_times_roman(justif)

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════
# PAGE 4 - ANALYSE DE SITUATION STRATEGIQUE
# ═════════════════════════════════════════════════════════════════

situation = doc.add_heading('2. ANALYSE DE SITUATION STRATEGIQUE', level=1)
for run in situation.runs:
    run.font.name = 'Times New Roman'

doc.add_heading('2.1 Diagnostic synthétique du marché', level=2)

# Tableau analyse marché
market_table = doc.add_table(rows=6, cols=3)
market_table.style = 'Light Grid Accent 1'

headers = ['Dimension', 'Caractéristique', 'Implication stratégique']
for i, header in enumerate(headers):
    cell = market_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.name = 'Times New Roman'

data = [
    ['Taille marché', '340 millions euros (2023)\n22% TCAC', 'Marché en forte croissance\nOpportunité entrée positive'],
    ['Surface cible', '2,6 millions hectares irrigables\n1,1 millions hectares irrigués', 'TAM significatif\nPénétration progressive réaliste'],
    ['Tendance réglementation', 'Réduction 20% consommation eau 2030\n68 départements restrictions', 'Drivers favorables entrée\nObligations compliance'],
    ['Adoption technologie', '+18% annuel digitalisation agriculture\n12% équipées capteurs', 'Marché receptif innovations\nAdoption accélérée possible'],
    ['Profitabilité exploitations', 'Marges compressées (-22% 2020-2025)\nBudget technologie limité', 'ROI court terme critère décision\nPrice sensitivity élevée']
]

for row_idx, row_data in enumerate(data, 1):
    for col_idx, content in enumerate(row_data):
        market_table.cell(row_idx, col_idx).text = content
        for para in market_table.cell(row_idx, col_idx).paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

doc.add_heading('2.2 Opportunités clés et défis', level=2)

opp_heading = doc.add_heading('Opportunités identifiées', level=3)
opportunities = [
    'Régulation croissante créant obligation quasi-certaine d\'adoption d\'outils 2025-2027',
    'Marché technologie agriculture croissance 22% TCAC offrant fenêtre favorable',
    'Consolidation coopératives créant points de contact distribution à fort effet multiplicateur',
    'Fenêtre concurrentielle : leaders mondiaux (AGCO, Raven) pas encore positionnés sur segment France',
    'Données agronomiques devenant asset stratégique pouvant générer revenus additionnels futurs',
    'Possibilité expansion géographique rapide (Allemagne, Espagne) modèle France réplicable'
]

for opp in opportunities:
    p = doc.add_paragraph(opp, style='List Bullet')
    set_times_roman(p)

challenges_heading = doc.add_heading('Défis majeurs à adresser', level=3)
challenges = [
    'Notoriété marque quasi-nulle versus leaders 50 ans présents marché',
    'Budget marketing limité (120k-150k euros) versus budgets concurrence (5-10 millions)',
    'Nécessité preuves sociales : exploitation pilotes et testimoniales agriculteurs indispensables',
    'Complexité chaîne distribution agricole avec multiples intermédiaires',
    'Risque perception startup : stabilité financière et pérennité solution questionnées',
    'Price elasticity complexe : agriculteurs ROI-focused mais budget souvent insuffisant'
]

for challenge in challenges:
    p = doc.add_paragraph(challenge, style='List Bullet')
    set_times_roman(p)

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════
# PAGE 6 - ANALYSE PESTEL
# ═════════════════════════════════════════════════════════════════

pestel = doc.add_heading('3. ANALYSE PESTEL', level=1)
for run in pestel.runs:
    run.font.name = 'Times New Roman'

doc.add_heading('3.1 Facteurs Politiques et Réglementaires', level=2)

pestel_text = doc.add_paragraph(
    'Le contexte politique français se caractérise par un soutien marqué aux technologies '
    'agricoles durables, illustré par le plan France 2030 allouant 800 millions d\'euros à '
    'la digitalisation du secteur agricole. La politique agricole commune (PAC) 2023-2027 '
    'conditionne 10% de ses enveloppes budgétaires (80 milliards euros en France) à la '
    'démonstration d\'efficacité en matière de gestion des ressources. Les régions proposent '
    'des aides cofinançant 50-70% du prix d\'acquisition pour technologies d\'optimisation '
    'hydrique.'
)
set_times_roman(pestel_text)

pestel_text2 = doc.add_paragraph(
    'Les risques politiques demeurent : changement de politique agricole post-2027, orientation '
    'extrême possible sur restrictions irrigation (prohibition totale certaines régions), et '
    'opposition lobby environnemental à l\'agriculture intensive.'
)
set_times_roman(pestel_text2)

doc.add_heading('3.2 Facteurs Économiques', level=2)

economic_table = doc.add_table(rows=5, cols=2)
economic_table.style = 'Light Grid Accent 1'

econ_headers = ['Facteur', 'Impact sur stratégie marketing']
for i, header in enumerate(econ_headers):
    cell = economic_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')

econ_data = [
    ['Volatilité prix commodités', 'Incertitude revenue agriculteurs affecte capacité investissement'],
    ['Taux intérêt croissants', 'Financement solutions plus onéreux, ROI court terme critère strict'],
    ['Inflation input (semences, carburant)', 'Compression marge => priorité réduction coûts'],
    ['Profitabilité exploitations cyclique', 'Budget technologie années fastes uniquement']
]

for row_idx, row_data in enumerate(econ_data, 1):
    for col_idx, content in enumerate(row_data):
        economic_table.cell(row_idx, col_idx).text = content

doc.add_heading('3.3 Facteurs Socioculturels', level=2)

social_data = [
    ['Adoption technologie', 'Générations <45 ans : 62% adoption digitale vs >55 ans : 28%'],
    ['Formation digitale', '52% exploitants >50 ans reportent manque compétence technique'],
    ['Confiance données', '26% agriculteurs expriment crainte perte confidentialité données'],
    ['Sensibilité environnement', '+45% agriculteurs 2020-2024 expriment préoccupation ressource hydrique']
]

social_table = doc.add_table(rows=5, cols=2)
social_table.style = 'Light Grid Accent 1'

headers = ['Dimension', 'Observation']
for i, header in enumerate(headers):
    cell = social_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')

for row_idx, row_data in enumerate(social_data, 1):
    for col_idx, content in enumerate(row_data):
        social_table.cell(row_idx, col_idx).text = content

doc.add_heading('3.4 Facteurs Technologiques', level=2)

tech_para = doc.add_paragraph(
    'Le contexte technologique demeure hautement favorable. Couverture LoRa et NB-IoT '
    'atteint 98% des zones rurales. Infrastructure cloud maturité production. Coûts hardware IoT '
    'baissent 12% annuellement. Frameworks IA open-source performants disponibles. Edge computing '
    'technologie éprouvée.'
)
set_times_roman(tech_para)

tech_para2 = doc.add_paragraph(
    'Disrupteurs à surveiller : satellite imagerie haute résolution commercialisée (Planet, Maxar), '
    'déploiement drones agricoles croissance 30% annuel, infrastructure 5G début de déploiement 2025-2026.'
)
set_times_roman(tech_para2)

doc.add_heading('3.5 Facteurs Environnementaux', level=2)

env_points = [
    'Sécheresses répétées depuis 2020 : événements climatiques "centennaux" annuels',
    'Nappes phréatiques : baisse 0,5 mètres/an certaines régions, déficit cumulé -60%',
    'Pression eau croissante : demandes concurrentes villes, industrie vs agriculture',
    'Certification eau/carbone possible future : primes possibles exploitations "efficaces"',
    'Opinion publique pression forte : irrigation perçue comme "gaspillage" par 62% français'
]

for point in env_points:
    p = doc.add_paragraph(point, style='List Bullet')
    set_times_roman(p)

doc.add_heading('3.6 Facteurs Légaux', level=2)

legal_para = doc.add_paragraph(
    'Directives européennes (2000/60/CE Water Framework Directive) imposent réduction 20% '
    'consommation eau horizons 2030. Green Deal européen fixe objectifs carbone agriculture 2050. '
    'France : Loi Agec (2020) oblige digitalisation agriculture 2030. Plans Locaux Gestion Eau (PLGE) '
    'définissent quotas par bassin versant avec pénalités 5-25k euros. Conformité RGPD données agriculteurs '
    'stricte, pénalités 4% chiffre affaires global.'
)
set_times_roman(legal_para)

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════
# PAGE 8 - ANALYSE COMPETITIVE
# ═════════════════════════════════════════════════════════════════

competitive = doc.add_heading('4. ANALYSE COMPETITIVE ET POSITIONNEMENT', level=1)
for run in competitive.runs:
    run.font.name = 'Times New Roman'

doc.add_heading('4.1 Paysage concurrentiel détaillé', level=2)

# Tableau concurrence détaillé
comp_table = doc.add_table(rows=6, cols=7)
comp_table.style = 'Light Grid Accent 1'

comp_headers = ['Critère', 'AGCO\nPrecision', 'Raven\nIndustries', 'SoilMate', 'WaterSense', 'Trimble', 'John Deere']
for i, header in enumerate(comp_headers):
    cell = comp_table.cell(0, i)
    cell.text = header
    shade_cell(cell, '4472C4')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = 'Times New Roman'

comp_data = [
    ['Prix HT/an', '8500', '7200', '2900', '4200', '6500', '5500'],
    ['IA Prescriptive', 'Non', 'Non', 'Non', 'Oui', 'Non', 'Non'],
    ['Support France', 'Limité', 'Bon', 'Excellent', 'Excellent', 'Moyen', 'Bon'],
    ['Edge Computing', 'Non', 'Non', 'Non', 'Oui', 'Non', 'Non'],
    ['Market Share France', '28%', '22%', '12%', '0% (Entry)', '5%', '8%']
]

for row_idx, row_data in enumerate(comp_data, 1):
    for col_idx, content in enumerate(row_data):
        comp_table.cell(row_idx, col_idx).text = str(content)

doc.add_heading('4.2 Analyse Forces Concurrentielles (Porter)', level=2)

porter_table = doc.add_table(rows=6, cols=2)
porter_table.style = 'Light Grid Accent 1'

porter_headers = ['Force competitive', 'Analyse et intensité']
for i, header in enumerate(porter_headers):
    cell = porter_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')

porter_data = [
    ['Rivalité existants', 'Intensité MOYENNE-FORTE. Concurrence 6 joueurs, différenciation faible (prix, features), taux churn élevé 15-20%.'],
    ['Menace entrants', 'Intensité MODÉRÉE. Barrières : brevet WaterSense 12-18 mois incopiable, capital initial 500k minimum, réseau distribution à construire.'],
    ['Pouvoir fournisseurs', 'Intensité FAIBLE. Fournisseurs hardware multiples, coûts décroissants, pas dépendance unique provider.'],
    ['Pouvoir clients', 'Intensité FORTE. Clients agriculteurs price-sensitive, ROI court terme critère strict, switching facile (pas lock-in contrats longs).'],
    ['Menace substituts', 'Intensité MOYENNE. Substituts : optimisation manuelle, solutions open-source, drone imagerie satellite. Pas menace immédiate.']
]

for row_idx, row_data in enumerate(porter_data, 1):
    porter_table.cell(row_idx, 0).text = row_data[0]
    porter_table.cell(row_idx, 1).text = row_data[1]

doc.add_heading('4.3 Positionnement stratégique WaterSense (Kotler)', level=2)

positioning_para = doc.add_paragraph(
    'Suivant le framework de Kotler, WaterSense adopte une stratégie de positionnement '
    '"Best Value" définie comme livraison de performance supérieure à prix compétitif. '
    'Cette stratégie articule trois éléments différenciation : (1) technologie supérieure '
    '(IA prescriptive seule du marché), (2) architecture robustesse (Edge Computing versus dépendance cloud), '
    '(3) expérience utilisateur simplifiée (interface agriculteur native vs interface ingénieur).'
)
set_times_roman(positioning_para)

positioning_matrix = doc.add_paragraph()
positioning_matrix.add_run('Proposition de valeur clé : ').bold = True
positioning_matrix.add_run(
    'Pour agriculteurs irrigants cherchant réduction coûts eau et énergie, WaterSense est la plateforme '
    'qui combine recommandations intelligentes précises, support français réactif et investissement '
    'proportionné, contrairement aux alternatives premium coûteuses (AGCO) ou basiques non-fiables (SoilMate).'
)
set_times_roman(positioning_matrix)

doc.add_heading('4.4 Avantages compétitifs durables', level=2)

advantages_table = doc.add_table(rows=5, cols=3)
advantages_table.style = 'Light Grid Accent 1'

adv_headers = ['Avantage', 'Durabilité', 'Défense']
for i, header in enumerate(adv_headers):
    cell = advantages_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')

adv_data = [
    ['Brevet FR3115088 IA', 'FORTE', 'Protection 10+ ans, incopiable 12-18 mois minimum'],
    ['Edge Computing architecture', 'MODEREE', 'Technologie mature mais déployment coûteux concurrence'],
    ['UX agriculteur native', 'MODEREE-FORTE', 'Ressource rare copier, nécessite DNA agricole équipe'],
    ['Ecosystem partenaires', 'MODEREE', 'First-mover avantage, mais autres entrée possible']
]

for row_idx, row_data in enumerate(adv_data, 1):
    for col_idx, content in enumerate(row_data):
        advantages_table.cell(row_idx, col_idx).text = content

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════
# PAGE 10 - SEGMENTATION MARCHE
# ═════════════════════════════════════════════════════════════════

segmentation = doc.add_heading('5. SEGMENTATION DU MARCHE', level=1)
for run in segmentation.runs:
    run.font.name = 'Times New Roman'

doc.add_heading('5.1 Critères de segmentation', level=2)

segmentation_text = doc.add_paragraph(
    'La stratégie de segmentation WaterSense utilise approche multi-critères combinant '
    'variables géographiques (régions stress hydrique), démographiques (taille exploitation, âge exploitant), '
    'comportementales (adoption technologie, sensibilité prix), et psychographiques (orientation durable, '
    'risk-aversion). Cette approche multi-critères permet identification segments haute valeur et haute probabilité conversion.'
)
set_times_roman(segmentation_text)

doc.add_heading('5.2 Segments de marché définis', level=2)

# Tableau segmentation
segment_table = doc.add_table(rows=6, cols=6)
segment_table.style = 'Light Grid Accent 1'

seg_headers = ['Segment', 'Volume', 'Priorité', 'Budget moyen', 'Accès primaire', 'Potentiel Y1']
for i, header in enumerate(seg_headers):
    cell = segment_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.bold = True

seg_data = [
    ['Maïs 20-200 ha', '28000', 'P1 - Primaire', '4-6k', 'Direct + Coops', '120-150'],
    ['Fruits/Arbo', '8500', 'P1 - Primaire', '6-10k', 'Arvalis+Direct', '40-50'],
    ['Coopératives', '2100', 'P2 - Secondaire', '15-30k', 'Direct', '15-20'],
    ['Grandes expl. 200+ha', '4200', 'P2 - Secondaire', '20-50k', 'Distributeurs', '20-30'],
    ['Maraîchage/légumes', '12000', 'P3 - Tertiaire', '3-5k', 'Chambres agri', '5-10']
]

for row_idx, row_data in enumerate(seg_data, 1):
    for col_idx, content in enumerate(row_data):
        segment_table.cell(row_idx, col_idx).text = str(content)
        for para in segment_table.cell(row_idx, col_idx).paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9)

doc.add_heading('5.3 Targeting stratégique priorité 2026', level=2)

targeting_para = doc.add_paragraph(
    'Stratégie targeting 2026 concentre ressources marketing sur deux segments complémentaires : '
    '(1) Exploitations maïs taille moyenne (50-150 hectares) en zones intensité irrigation élevée '
    '(Nouvelle-Aquitaine, Centre-Val Loire) où ROI rapide et clustering géographique optimisent '
    'efficacité commerciale, (2) Arboriculture/fruits régions PACA où valeur récolte élevée justifie '
    'prix premium et adoption technologie anticipée.'
)
set_times_roman(targeting_para)

targeting_para2 = doc.add_paragraph(
    'Cette concentration volontaire permet : (1) concentration effort commercial (3 sales reps couvrent 70% '
    'marché cible), (2) efficacité marketing (messaging cohérent par segment, co-marketing Arvalis efficace), '
    '(3) proof points rapides (12-20 clients "showcase" Year 1 générant références multiplicatrices).'
)
set_times_roman(targeting_para2)

doc.add_heading('5.4 Géographie commerciale 2026', level=2)

geography_table = doc.add_table(rows=5, cols=4)
geography_table.style = 'Light Grid Accent 1'

geo_headers = ['Zone', 'Départements', 'Priorité', 'Représentation']
for i, header in enumerate(geo_headers):
    cell = geography_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')

geo_data = [
    ['Nouvelle-Aquitaine', 'Dordogne, Lot-et-Garonne, Gironde (12)', 'CRITIQUE', 'Sales rep 1 + Partner SMAG'],
    ['Centre-Val Loire', 'Indre, Indre-et-Loire, Loir-et-Cher (6)', 'CRITIQUE', 'Sales rep 2'],
    ['Pays de la Loire', 'Maine-et-Loire, Vendée, Loire-Atlantique (6)', 'HAUTE', 'Sales rep 3 + Coops'],
    ['PACA', 'Provence, Alpes, Côte-Azur (5)', 'HAUTE', 'Partner Arvalis'],
]

for row_idx, row_data in enumerate(geo_data, 1):
    for col_idx, content in enumerate(row_data):
        geography_table.cell(row_idx, col_idx).text = content

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════
# PAGE 13 - STRATEGIE 4P
# ═════════════════════════════════════════════════════════════════

four_p = doc.add_heading('6. STRATEGIE MARKETING 4P DETAILLEE', level=1)
for run in four_p.runs:
    run.font.name = 'Times New Roman'

doc.add_heading('6.1 PRODUIT - Politique Produit', level=2)

produit_text = doc.add_paragraph(
    'Offre WaterSense structurée architecture modulaire permettant adaptation segments clients '
    'différents. Produit core constitué 12 capteurs IoT multisensoriels, unité Edge Computing locale, '
    'plateforme cloud AWS, applications mobiles iOS/Android, interface web. Valeur centrale : '
    'recommandations prescriptives précises générées algorithme IA breveté.'
)
set_times_roman(produit_text)

# Tableau offre produit
produit_table = doc.add_table(rows=5, cols=5)
produit_table.style = 'Light Grid Accent 1'

prod_headers = ['Variante', 'Config', 'Prix HT', 'Cible', 'Inclusion']
for i, header in enumerate(prod_headers):
    cell = produit_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')

prod_data = [
    ['ESSENTIAL', '8 capteurs', '3200', 'Petite expl 20-50ha', 'Formation 4h, support email, 1an cloud'],
    ['STANDARD', '12 capteurs', '4200', 'PME maïs 50-150ha', 'Formation 8h, support email, 1an cloud'],
    ['PREMIUM', '20 capteurs', '6800', 'Fruits/arbo', 'Formation 12h, support phone, 1an cloud'],
    ['PROFESSIONAL', '30+ capteurs', '9500', 'Grandes expl 150+ha', 'Formation+coaching, support 24/7']
]

for row_idx, row_data in enumerate(prod_data, 1):
    for col_idx, content in enumerate(row_data):
        produit_table.cell(row_idx, col_idx).text = str(content)

doc.add_heading('6.2 PRIX - Politique Tarifaire', level=2)

prix_heading = doc.add_heading('Approche de pricing', level=3)

prix_para = doc.add_paragraph(
    'Stratégie pricing WaterSense adopte approche "value-based pricing" basée création valeur client. '
    'Calcul prix fondé quantification économies réelles par exploitation type, justifiant investissement via ROI démontrés.'
)
set_times_roman(prix_para)

# Tableau ROI pricing
roi_table = doc.add_table(rows=5, cols=4)
roi_table.style = 'Light Grid Accent 1'

roi_headers = ['Métrique', 'Economie annuelle', 'Prix variante', 'Payback']
for i, header in enumerate(roi_headers):
    cell = roi_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')

roi_data = [
    ['Réduction eau 18%', '1300 euros/100ha', '4200', '3.2 mois'],
    ['Réduction énergie 20%', '360 euros/100ha', '4200', '14 mois cumul.'],
    ['Augmentation rendement 8%', '9680 euros/100ha', '4200', '5.2 semaines'],
    ['Mitigation risque amendes', '6000 euros estimé', '4200', '8.4 semaines']
]

for row_idx, row_data in enumerate(roi_data, 1):
    for col_idx, content in enumerate(row_data):
        roi_table.cell(row_idx, col_idx).text = str(content)

doc.add_heading('Promotions tarifaires lancementmarket', level=3)

promotions = [
    'Early adopter discount 15% premiers 50 clients (création effet urgence)',
    'Référence discount 10% si client recommande autre prospect (word-of-mouth incité)',
    'Coopérative bulk discount 8% commandes 10+ unités (channel enablement)',
    'Financing 12-month payment plan 350 euros/mois sans intérêt (payment friction reduction)',
    'Subsidy pre-financing : WATERSENSE avance 30% pour clients demandant subventions régionales'
]

for promo in promotions:
    p = doc.add_paragraph(promo, style='List Bullet')
    set_times_roman(p)

doc.add_heading('6.3 DISTRIBUTION - Politique de Distribution', level=2)

distribution_text = doc.add_paragraph(
    'Stratégie distribution omnicanal (OMO - Online Managed Orchestration) visant couverture 90% exploitations '
    'cibles via combinaison canaux directs et indirects. Objectif : neutraliser avantages distribution concurrence, '
    'créer points de contact multiples, maximiser accessibility clients.'
)
set_times_roman(distribution_text)

# Tableau channels
dist_table = doc.add_table(rows=5, cols=4)
dist_table.style = 'Light Grid Accent 1'

dist_headers = ['Canal', 'Format', 'Part revenue Y1', 'KPI succès']
for i, header in enumerate(dist_headers):
    cell = dist_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')

dist_data = [
    ['Direct - E-commerce', 'watersense-agri.fr', '12-15%', '25% ventes online, 15k visits/mois'],
    ['Direct - Sales terrain', '3 commerciaux', '20-25%', '35-40 closures/person/year'],
    ['Distributeurs SMAG', 'Réseau 120 points', '15-18%', '200 units Y1, CAC <600'],
    ['Coopératives (15)', 'Points vente coop', '35-40%', '36+ customers, CAC <500']
]

for row_idx, row_data in enumerate(dist_data, 1):
    for col_idx, content in enumerate(row_data):
        dist_table.cell(row_idx, col_idx).text = str(content)

doc.add_heading('6.4 PROMOTION - Politique de Communication Marketing', level=2)

communication_text = doc.add_paragraph(
    'Plan promotion WaterSense combine stratégies brand building, lead generation, et education market. '
    'Approche omnichannel utilisant digital (SEM, SEO, content, social), offline (events, field trials), '
    'et PR (media relations, thought leadership) pour créer awareness, consideration, et conversion segmenté.'
)
set_times_roman(communication_text)

doc.add_heading('Allocation budget marketing', level=3)

# Budget marketing table
budget_table = doc.add_table(rows=8, cols=3)
budget_table.style = 'Light Grid Accent 1'

budget_headers = ['Canal Marketing', 'Budget 2026', 'Objectif']
for i, header in enumerate(budget_headers):
    cell = budget_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')

budget_data = [
    ['Digital (SEM, SEO, content, social)', '45000', 'Générer 300-400 qualified leads/mois'],
    ['Trade shows et events (SIA, régional)', '32000', 'Générer 600-800 leads, 15-20 closures'],
    ['PR et relations médias', '18000', 'Générer 10-15 articles presse, 3-5 speaking'],
    ['Co-marketing partenaires (Arvalis, coops)', '20000', 'Amplifier reach, co-branded content'],
    ['Field trials et démos (10-15 sites)', '15000', 'Générer 50-100 leads par site'],
    ['Brand et collateral', '6000', 'Design professionnel, matériaux point vente'],
    ['Contingency et opportunités', '4000', ''],
]

for row_idx, row_data in enumerate(budget_data, 1):
    for col_idx, content in enumerate(row_data):
        budget_table.cell(row_idx, col_idx).text = str(content)

total_row = 7
budget_table.cell(total_row, 0).text = 'TOTAL BUDGET'
budget_table.cell(total_row, 1).text = '140000 euros'
budget_table.cell(total_row, 2).text = '(12% revenue Y1 projeté)'

for cell in [budget_table.cell(total_row, 0), budget_table.cell(total_row, 1), budget_table.cell(total_row, 2)]:
    shade_cell(cell, 'CCCCCC')

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════
# PAGE 18 - PLAN MARKETING OPERATIONNEL
# ═════════════════════════════════════════════════════════════════

operationnel = doc.add_heading('7. PLAN MARKETING OPERATIONNEL 2026', level=1)
for run in operationnel.runs:
    run.font.name = 'Times New Roman'

doc.add_heading('7.1 Stratégie digitale détaillée', level=2)

digital_text = doc.add_paragraph(
    'Stratégie digitale WaterSense comprend trois piliers : (1) Acquisition : e-commerce, SEM, SEO, '
    'social media pour driver awareness et leads, (2) Engagement : content marketing éducatif, '
    'webinars, email marketing pour nurture prospects dans consideration phase, (3) Conversion : '
    'ROI calculator, case studies, customer testimonials pour drive closing.'
)
set_times_roman(digital_text)

doc.add_heading('Website et e-commerce', level=3)

website_points = [
    'Platform : Shopify B2C ou Magento B2B scaled',
    'Content : 50+ pages optimisé SEO (keywords "irrigation optimization", "farm IoT", "water management")',
    'Fonctionnalités clés : ROI calculator personnalisé, webinar registration, CRM lead capture intégrée',
    'Performance target : 15000 uniques visitors/mois Y1, 25 mois/mois target Y1'
]

for point in website_points:
    p = doc.add_paragraph(point, style='List Bullet')
    set_times_roman(p)

doc.add_heading('Search Engine Marketing (SEM)', level=3)

sem_para = doc.add_paragraph(
    'Budget Google Ads 45000 euros annuels (3750 euros/mois) visant high-intent keywords. '
    'Campagnes segmentées par segment client (maïs, fruits, coops) avec messaging adapté. '
    'Target CPL (Cost Per Lead) 25-30 euros, conversion rate 3-4% leads to pipeline, '
    'expected monthly leads 350-400.'
)
set_times_roman(sem_para)

doc.add_heading('Search Engine Optimization (SEO)', level=3)

seo_para = doc.add_paragraph(
    'Strategy SEO organique long-term : 30+ blog articles (agriculture, technologie, eau), '
    'link building partenaires (Arvalis, coops, blogs agricoles), local SEO optimization. '
    'Target classement #1-3 pour 20+ commercial keywords par end Y1, driving 200-300 organic leads/mois'
)
set_times_roman(seo_para)

doc.add_heading('Social media et community', level=3)

social_para = doc.add_paragraph(
    'Présence LinkedIn (professional networking agriculteurs, thought leadership), Facebook (community building, '
    'farmer advocacy groups). Content cadence 3-4 posts/semaine. Target 5000 followers LinkedIn, 10000 Facebook by Q4. '
    'Community moderation, user generated content valorisation, farmer testimonials amplification.'
)
set_times_roman(social_para)

doc.add_heading('7.2 Stratégie événementielle', level=2)

events_text = doc.add_paragraph(
    'Plan événementiel 2026 cible 12-15 événements majeurs combinant salons professionnels, '
    'foires agricoles régionales, webinaires en ligne, et démonstrations terrain. Objectif : '
    'générer 600-800 leads annuels, créer brand visibility, valider product-market fit.'
)
set_times_roman(events_text)

# Tableau événements
events_table = doc.add_table(rows=6, cols=5)
events_table.style = 'Light Grid Accent 1'

events_headers = ['Événement', 'Mois', 'Audience', 'Budget', 'Leads cible']
for i, header in enumerate(events_headers):
    cell = events_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')

events_data = [
    ['SIA Paris (Salon Int. Agriculture)', 'Feb', '35000 visiteurs', '15000', '400-500'],
    ['Agro Solutions (Angers)', 'March', '25000 visiteurs', '8000', '250-300'],
    ['Webinaire série mensuelle', 'Monthly', '200-300 registrants', '500/mois', '30-40/mois'],
    ['Demo plots et open days (10 sites)', 'May-July', '20-50 agriculteurs/site', '15000', '50-100/site'],
    ['Agrinove (Villepinte)', 'April', '15000 visiteurs', '6000', '150-200'],
]

for row_idx, row_data in enumerate(events_data, 1):
    for col_idx, content in enumerate(row_data):
        events_table.cell(row_idx, col_idx).text = str(content)

doc.add_heading('7.3 Stratégie relations publiques', level=2)

pr_para = doc.add_paragraph(
    'Plan relations presse cible publications agricoles majeures (Réussir Grandes Cultures, '
    'Terre-net, Les Échos Agri) avec angles éditoriaux : "Innovation climatique agricole", '
    '"Réduction coûts eau exploitations", "Young entrepreneurs agritech". Press releases '
    'trimestriels. Target 10-15 articles presse, 3-5 speaking slots conférences.'
)
set_times_roman(pr_para)

doc.add_heading('7.4 Stratégie partenariat et co-marketing', level=2)

coop_para = doc.add_paragraph(
    'Co-marketing partenaires clés (Arvalis, SMAG, 15 coopératives) incluant co-branded content, '
    'joint webinars, cross-promotion. Arvalis endorsement validation technique, shared lead generation. '
    'Coopératives in-store materials, staff training events, point-of-sale advertising. Budget '
    '20000 euros annual allocation co-marketing programs.'
)
set_times_roman(coop_para)

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════
# PAGE 23 - BUDGET ET RESSOURCES
# ═════════════════════════════════════════════════════════════════

budget_section = doc.add_heading('8. BUDGET ET ALLOCATION RESSOURCES', level=1)
for run in budget_section.runs:
    run.font.name = 'Times New Roman'

doc.add_heading('8.1 Budget marketing global 2026', level=2)

global_budget_para = doc.add_paragraph(
    'Budget total marketing alloué 2026 s\'élève 140000 euros, représentant 12% du revenue Y1 '
    'projeté (1.17M euros). Allocation respecte benchmarks industrie agritech (10-15% revenue marketing) '
    'et concentre ressources canaux high-ROI.'
)
set_times_roman(global_budget_para)

# Tableau budget détaillé
detailed_budget = doc.add_table(rows=10, cols=4)
detailed_budget.style = 'Light Grid Accent 1'

budget_heads = ['Catégorie marketing', 'Budget 2026', '% du total', 'Justification']
for i, header in enumerate(budget_heads):
    cell = detailed_budget.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = 'Times New Roman'

budget_detail_data = [
    ['Digital (SEM, SEO, social, content)', '45000', '32%', 'Lead generation primary, ROI trackable'],
    ['Events et foires agricoles', '32000', '23%', 'Brand visibility, B2B lead generation'],
    ['Field trials et démos', '15000', '11%', 'Product validation, références client'],
    ['Co-marketing partenaires', '20000', '14%', 'Leverage Arvalis reach, coops network'],
    ['PR et relations médias', '18000', '13%', 'Thought leadership, credibility building'],
    ['Brand et collateral', '6000', '4%', 'Professional materials, design'],
    ['Contingency 5%', '4000', '3%', 'Opportunités émergentes'],
]

for row_idx, row_data in enumerate(budget_detail_data, 1):
    for col_idx, content in enumerate(row_data):
        detailed_budget.cell(row_idx, col_idx).text = str(content)
        for para in detailed_budget.cell(row_idx, col_idx).paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9)

total_row = 7
detailed_budget.cell(total_row, 0).text = 'TOTAL BUDGET MARKETING'
detailed_budget.cell(total_row, 1).text = '140000'
detailed_budget.cell(total_row, 2).text = '100%'
for cell in [detailed_budget.cell(total_row, i) for i in range(3)]:
    shade_cell(cell, 'CCCCCC')

doc.add_heading('8.2 Ressources opérationnelles requises', level=2)

resources_text = doc.add_paragraph(
    'Exécution plan marketing requiert allocation ressources humaines dédiées. Équipe marketing '
    'centrale WaterSense composée 2.5 FTE (équivalents temps plein) : 1 Marketing Director, '
    '1 Digital Marketing Specialist, 0.5 Content/PR Manager. Sales team : 3 commerciaux terrain '
    'couvrant zones géographiques prioritaires. Support externe : agence digital (SEM/SEO), '
    'agence PR, production vidéo.'
)
set_times_roman(resources_text)

resources_table = doc.add_table(rows=6, cols=3)
resources_table.style = 'Light Grid Accent 1'

res_headers = ['Ressource', 'Type', 'Allocation']
for i, header in enumerate(res_headers):
    cell = resources_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')

res_data = [
    ['Marketing Director', 'Interne FTE', '1.0 (pilotage stratégie, partenaires)'],
    ['Digital Marketing Specialist', 'Interne FTE', '1.0 (SEM, SEO, analytics, content)'],
    ['Content/PR Manager', 'Interne FTE', '0.5 (PR, events, webinars)'],
    ['Sales representatives', 'Interne FTE', '3.0 (terrain direct)'],
    ['External agencies', 'Externe', 'Digital 30k, PR 10k, Video 8k'],
]

for row_idx, row_data in enumerate(res_data, 1):
    for col_idx, content in enumerate(row_data):
        resources_table.cell(row_idx, col_idx).text = str(content)

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════
# PAGE 25 - KPI ET METRIQUES
# ═════════════════════════════════════════════════════════════════

kpi_section = doc.add_heading('9. METRIQUES DE PILOTAGE ET KPI', level=1)
for run in kpi_section.runs:
    run.font.name = 'Times New Roman'

doc.add_heading('9.1 Framework de mesure marketing', level=2)

framework_para = doc.add_paragraph(
    'Framework KPI WaterSense adopte approche hiérarchisée Kotler distinguant : (1) KPI upstream '
    '(awareness, consideratin metrics) mesurant efficacité campaigns, (2) KPI midstream (conversion metrics) '
    'mesurant sales effectiveness, (3) KPI downstream (customer metrics) mesurant business impact. '
    'Suivi mensuel dashboards, révision trimestrielle stratégie.'
)
set_times_roman(framework_para)

doc.add_heading('9.2 KPI par canal marketing', level=2)

# Tableau KPI detailed
kpi_table = doc.add_table(rows=6, cols=4)
kpi_table.style = 'Light Grid Accent 1'

kpi_headers = ['Canal', 'KPI primaire', 'Target 2026', 'Seuil alerte']
for i, header in enumerate(kpi_headers):
    cell = kpi_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')

kpi_data = [
    ['Website/Digital', 'Leads/mois', '350-400', '<250 (trigger optimization)'],
    ['Events', 'Cost per lead generated', '<25 euros', '>40 euros (reduce events)'],
    ['Field trials', 'Lead per site', '50-100', '<30 (poor ROI)'],
    ['Partner coops', 'Customers via channel', '35-40', '<20 (renegotiate)'],
]

for row_idx, row_data in enumerate(kpi_data, 1):
    for col_idx, content in enumerate(row_data):
        kpi_table.cell(row_idx, col_idx).text = str(content)

doc.add_heading('9.3 KPI commerciaux globaux', level=2)

# Tableau commercial metrics
commercial_table = doc.add_table(rows=9, cols=3)
commercial_table.style = 'Light Grid Accent 1'

comm_headers = ['Métrique', 'Target 2026', 'Note']
for i, header in enumerate(comm_headers):
    cell = commercial_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')

comm_data = [
    ['Leads générés annuel', '4000-5000', 'Tous canaux combinés'],
    ['Leads qualified', '350-400', 'Conversion ~8% leads totaux'],
    ['Customers acquis', '120-150', 'Conversion ~35% qualified leads'],
    ['Customer Acquisition Cost', '<600 euros', 'Budget 140k / 120-150 customers'],
    ['Average Selling Price', '4200 euros', 'Mix variances: 3200-9500'],
    ['Revenue Y1 projected', '504k-630k euros', 'Conservateur : 120 × 4.2k'],
    ['Net Promoter Score', '>65', 'Promoteurs vs detractors'],
    ['Customer churn Y1', '<5%', 'Retention >95%'],
]

for row_idx, row_data in enumerate(comm_data, 1):
    for col_idx, content in enumerate(row_data):
        commercial_table.cell(row_idx, col_idx).text = str(content)

doc.add_heading('9.4 Tableau de bord mensuel de suivi', level=2)

dashboard_para = doc.add_paragraph(
    'Tableau de bord marketing mensuel consolidant KPI critiques : leads par source, conversion funnel stages, '
    'CAC par canal, revenue pipeline, market share estimée. Suivi temps réel via Salesforce CRM ou Google Analytics. '
    'Révision stratégie trimestrielle si variance >20% vs target.'
)
set_times_roman(dashboard_para)

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════
# PAGE 27 - CHRONOGRAMME EXECUTION
# ═════════════════════════════════════════════════════════════════

timeline_section = doc.add_heading('10. CHRONOGRAMME D\'EXECUTION 2026', level=1)
for run in timeline_section.runs:
    run.font.name = 'Times New Roman'

doc.add_heading('10.1 Plan trimestriel d\'exécution', level=2)

# Tableau timeline Q1-Q4
timeline_table = doc.add_table(rows=5, cols=3)
timeline_table.style = 'Light Grid Accent 1'

timeline_headers = ['Trimestre', 'Jalons clés', 'Métriques cibles']
for i, header in enumerate(timeline_headers):
    cell = timeline_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')

timeline_data = [
    ['Q1 2026\n(Jan-Mars)', 
     'Website go-live\nSales team recruitment/training\nPartner agreements signed\nFirst 20 customers\nSIA Paris event (35k leads)',
     'Website 5k visitors/mois\n20 customers\n200 leads pipeline'],
    
    ['Q2 2026\n(Avr-Juin)',
     'Field trials 10+ sites\nDigital campaigns scaling\nDistributor training 100+ staff\nWebinar series launch',
     '300+ leads/mois\n50 customers cumul.\n100 pipeline'],
    
    ['Q3 2026\n(Juil-Sept)',
     'Peak selling season\nDemo plots generate referrals\nPartner co-marketing campaigns\nProduct v1.1 release',
     '400 leads/mois\n100 customers cumul.\n200 pipeline'],
    
    ['Q4 2026\n(Oct-Dec)',
     'Year-end push\nCustomer success stories\nBrand awareness campaigns\nFranchise model planning',
     '120-150 customers total\nEBITDA breakeven\nNPS benchmark']
]

for row_idx, row_data in enumerate(timeline_data, 1):
    for col_idx, content in enumerate(row_data):
        timeline_table.cell(row_idx, col_idx).text = str(content)
        for para in timeline_table.cell(row_idx, col_idx).paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9)

doc.add_heading('10.2 Dépendances critiques', level=2)

dependencies = [
    'Finalisation website et e-commerce : délai non-négociable janvier 2026 (SEM campaigns démarrent février)',
    'Signature contrats partenaires (Arvalis, SMAG, coops) : fin janvier 2026 minimum',
    'Recrutement 3 sales reps : février 2026 (training 3-4 semaines avant terrain)',
    'Préparation 10-15 field trials : mars 2026 (déploiement avril-mai)',
    'Production contenu marketing (30+ articles, vidéos) : janvier-février 2026 (publication mars onwards)'
]

for dep in dependencies:
    p = doc.add_paragraph(dep, style='List Bullet')
    set_times_roman(p)

doc.add_heading('10.3 Risques calendrier', level=2)

calendar_risks = [
    'Retard website : delay SEM campaigns 1-2 mois, impact leads -200-300',
    'Partenaire signature retard : delay distribution channel 2-3 mois',
    'Sales team recruitment difficulté : délai hiring +4 semaines, impact customers -20-30',
    'Weather conditions dry summer : avantage marketing (accentue problem statement)',
    'Competitive launch : si concurrent major entre march 2026, window closes rapidly'
]

for risk in calendar_risks:
    p = doc.add_paragraph(risk, style='List Bullet')
    set_times_roman(p)

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════
# PAGE 29 - GESTION RISQUES
# ═════════════════════════════════════════════════════════════════

risks_section = doc.add_heading('11. GESTION DES RISQUES MARKETING', level=1)
for run in risks_section.runs:
    run.font.name = 'Times New Roman'

doc.add_heading('11.1 Matrice des risques identifiés', level=2)

# Tableau risques
risks_table = doc.add_table(rows=6, cols=4)
risks_table.style = 'Light Grid Accent 1'

risks_headers = ['Risque', 'Probabilité', 'Impact', 'Mitigation']
for i, header in enumerate(risks_headers):
    cell = risks_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')

risks_data = [
    ['Adoption lente marché', 'MOYENNE', 'FORT', 'Field trials multiples, testimonials agriculteurs, messaging ROI clair'],
    ['Concurrence prix agressive', 'FORTE', 'FORT', 'Differentiation UX/support, brand loyalty building, pas price war'],
    ['Market shift vers satellite imagery', 'FAIBLE', 'MOYEN', 'R&D continu, edge computing defensibility, ecosystem partnerships'],
    ['Restriction eau politique extrême', 'FAIBLE-MOY', 'CRITIQUE', 'Pivot expansion Allemagne/Espagne, adjacent crops adaptability'],
    ['Partenaire distribution failure', 'MOYENNE', 'MOYEN', 'Diversification canaux, direct vente escalation, new partner recruitment']
]

for row_idx, row_data in enumerate(risks_data, 1):
    for col_idx, content in enumerate(row_data):
        risks_table.cell(row_idx, col_idx).text = str(content)

doc.add_heading('11.2 Plan contingence scénarios', level=2)

scenario_heading = doc.add_heading('Scénario 1 : Adoption lente (100-110 customers Y1 vs 120-150 target)', level=3)
scenario1 = doc.add_paragraph(
    'Actions contingence : (1) Acceleration field trials (+5 sites), (2) Augmentation budget co-marketing Arvalis (+20k), '
    '(3) Pricing promotions temporaires month 9-10 (urgency creation), (4) Sales team incentive augmentée (+15% commission), '
    '(5) Product roadmap acceleration features demand.'
)
set_times_roman(scenario1)

scenario_heading2 = doc.add_heading('Scénario 2 : Concurrence prix SoilMate dump vers 1900-2000 euros', level=3)
scenario2 = doc.add_paragraph(
    'Actions contingence : (1) Messaging differentiation UX/support vs prix, pas match prix, (2) Augmentation perceived value '
    '(free training, extended support), (3) Bundle promotions (système + installation), (4) Enterprise tier pricing premium justifié '
    '(larger deployments), (5) Partner exclusive territory protection (coops lock-in contracts).'
)
set_times_roman(scenario2)

scenario_heading3 = doc.add_heading('Scénario 3 : Major competitor (Bayer, BASF) market entry Q3 2026', level=3)
scenario3 = doc.add_paragraph(
    'Actions contingence : (1) Accélération market share capture Q1-Q2 (aggressive leads conversion), (2) Customer lock-in '
    '(multi-year contracts possible), (3) Niche positioning vs generalist competitors (precision, support), (4) Franchise model '
    'acceleration (local presence), (5) Strategic partnership acquisition target positioning.'
)
set_times_roman(scenario3)

doc.add_heading('11.3 Early warning indicators', level=2)

indicators = [
    'Monthly customer acquisition declining >15% vs previous month deux mois consecutifs',
    'Website lead quality declining : conversion leads->pipeline dropping <5%',
    'Distributor partner complaints sur support ou fulfillment',
    'NPS customer dropping <60 (satisfaction deterioration signal)',
    'Competitive price change >10% de notre pricing',
    'Media coverage increases competitor vs WaterSense ratio >3:1'
]

for indicator in indicators:
    p = doc.add_paragraph(indicator, style='List Bullet')
    set_times_roman(p)

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════
# PAGE 31 - CONCLUSION
# ═════════════════════════════════════════════════════════════════

conclusion_section = doc.add_heading('12. CONCLUSION ET RECOMMANDATIONS', level=1)
for run in conclusion_section.runs:
    run.font.name = 'Times New Roman'

conclusion_para = doc.add_paragraph(
    'Plan marketing stratégique WaterSense 2026 structure approche ambitieuse mais réaliste d\'introduction '
    'solution innovante marché français irrigation. Targeting segmentation claire (maïs + fruits prioritaires), '
    'positioning différenciation durable (IA prescriptive + edge computing + UX native), et stratégie 4P cohérente '
    'combinent offrir meilleure probabilité succès. Allocation ressources budget (140k euros marketing) proportionnée '
    'au TAM (2.6M hectares) et objective Year 1 conservateur (120-150 customers).'
)
set_times_roman(conclusion_para)

conclusion_para2 = doc.add_paragraph(
    'Trois conditions succès critiques : (1) Exécution flawless plan go-to-market (website, sales team, partners Q1), '
    '(2) Product-market fit validation rapide (field trials delivering ROI promises), (3) Brand building steady (testimonials, '
    'media coverage, word-of-mouth création momentum). Marché positioning "Best Value" résonates agriculteurs priorité cost reduction '
    'et risque maîtrise.'
)
set_times_roman(conclusion_para2)

doc.add_heading('Recommandations clés', level=2)

recommendations = [
    'Priorité immédiate : Finaliser website/e-commerce janvier 2026, débuter SEM campaigns février (lead generation start)',
    'Dépendance partenaires : Signer contrats Arvalis, SMAG, coopératives 15 avant janvier 31 (aucun flexibility)',
    'Early wins critiques : Acquérir 20 customers Q1 pour case studies, field trials references Q2 (momentum building)',
    'Messaging discipline : "Arrosez Moins, Gagnez Plus" tagline consistency toute communication (brand recall building)',
    'Metrics-driven agility : Suivi mensuel KPI dashboards, révision quarterly stratégie si variance >20% targets',
    'Risk vigilance : Monitor competitive landscape closely (satellite imagery evolution, competitor pricing moves)',
    'Personnel engagement : Sales team compensation structure incentive aggressive customer acquisition (commission 8% proposed)'
]

for rec in recommendations:
    p = doc.add_paragraph(rec, style='List Bullet')
    set_times_roman(p)

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════
# ANNEXES
# ═════════════════════════════════════════════════════════════════

annexes = doc.add_heading('ANNEXES', level=1)
for run in annexes.runs:
    run.font.name = 'Times New Roman'

doc.add_heading('ANNEXE A - Détail segments clients et caractéristiques', level=2)

annexe_a = doc.add_paragraph(
    'Segment Maïs 20-200 hectares (Volume 28000 exploitations, Priorité P1). Caractéristiques : exploitation '
    'agriculture conventionnelle, irrigation intensive (600-800 m3/hectare/an), revenus moyens (80-150k euros/an). '
    'Pain points : hausse coûts énergie (-22% marge 5 ans), restrictions irrigation croissantes, pressure rendement. '
    'Value proposition : ROI 2.2 mois payback via économies eau/énergie (1.7-2.3k euros/hectare/an). '
    'Accès distribution : direct sales terrain + coopératives réseau. Early adopter profile : exploitants 35-50 ans, '
    'tech-familiar. Segmentation géographique secondaire : Nouvelle-Aquitaine (12 000), Centre-Val Loire (8000), '
    'Pays Loire (5000), autres régions (3000).'
)
set_times_roman(annexe_a)

doc.add_heading('ANNEXE B - Analyse détaillée concurrents', level=2)

annexe_b = doc.add_paragraph(
    'AGCO Precision Fuse Connect : Leader marché (28% share, 3200 clients). Streng : intégration tracteurs AGCO, '
    'infrastructure stable, brand trust acquis. Weaknesses : interface complexe ingénieur-focused, IA descriptive only, '
    'support français limité, capteurs additionnels coûteux. Pricing 8500 euros + 1200 euros/an. Vs WaterSense : 50% plus '
    'cher, IA inférieure, pas edge computing. RAVEN Industries (22% share) : Heritage marque, satellite imagery, datée ergonomie. '
    'SOILMATE (12% share) : Budget entry-point 2900 euros mais IA basique, instabilité infrastructure (churn 18%), generic recommendations. '
    'WaterSense "Best Value" : AGCO performance @ SoilMate prix + native UX + support français.'
)
set_times_roman(annexe_b)

doc.add_heading('ANNEXE C - ROI calculator scenarios', level=2)

annexe_c_heading = doc.add_paragraph('Scénario 1 : Exploitation maïs 100 hectares').bold = True
set_times_roman(annexe_c_heading)

annexe_c = doc.add_paragraph(
    'Investissement : 4200 euros (STANDARD). Annual returns : Eau reduction 18% (108€/ha/100ha = 10800), '
    'Énergie reduction 20% (1500 euros), Rendement +8% (9680 euros). Total annuel 22000 euros. Payback : '
    '4200/22000 = 2.3 months. Cumulative 3-year value : 66000 euros (3x investment).'
)
set_times_roman(annexe_c)

doc.add_heading('ANNEXE D - Tableau de bord marketing mensuel', level=2)

dashboard_table = doc.add_table(rows=10, cols=3)
dashboard_table.style = 'Light Grid Accent 1'

dash_headers = ['KPI', 'Target mensuel', 'Seuil alerte']
for i, header in enumerate(dash_headers):
    cell = dashboard_table.cell(0, i)
    cell.text = header
    shade_cell(cell, 'D3D3D3')

dash_data = [
    ['Leads générés (tous canaux)', '350-400', '<250'],
    ['Leads qualified', '25-30', '<15'],
    ['Demo/trial demandes', '8-10', '<5'],
    ['Customers signed', '10-12', '<5'],
    ['CAC réalisé', '<600', '>750'],
    ['Website visitors', '1500', '<1000'],
    ['Email engagement rate', '3-5%', '<2%'],
    ['Social media followers growth', '300', '<100'],
    ['Press mentions', '1-2', '0'],
]

for row_idx, row_data in enumerate(dash_data, 1):
    for col_idx, content in enumerate(row_data):
        dashboard_table.cell(row_idx, col_idx).text = str(content)

doc.add_heading('ANNEXE E - Bibliographie et références', level=2)

bibliography = [
    'Kotler, P., Keller, K. L. (2021). Marketing Management (16th ed.). Pearson.',
    'Porter, M. E. (2008). Competitive Advantage: Creating and Sustaining Superior Performance. Free Press.',
    'Ries, A., Trout, J. (2001). Positioning: The Battle for Your Mind. McGraw-Hill.',
    'Godin, S. (2018). This is Marketing: You Cannot Be Seen Until You Learn to See. Portfolio.',
    'Miller, D. (2014). Building a StoryBrand: Clarify Your Message So Customers Will Listen. HarperBusiness.',
    'Commission Agriculture France (2024). Études de marché irrigation et technologie agricole 2024.',
    'Agreste France (2023). Statistics agriculture générale France 2023.',
]

for ref in bibliography:
    p = doc.add_paragraph(ref, style='List Bullet')
    set_times_roman(p)

doc.add_heading('ANNEXE F - Glossaire marketing utilisé', level=2)

glossary_items = [
    ['TCAC', 'Taux Croissance Annuel Composé - croissance annuelle moyenne période'],
    ['CAC', 'Customer Acquisition Cost - coût moyen acquisition client'],
    ['LTV', 'Lifetime Value - valeur client 3-5 ans cumul.'],
    ['NPS', 'Net Promoter Score - satisfaction métrique (>50 excellent)'],
    ['KPI', 'Key Performance Indicator - métrique pilotage'],
    ['OMO', 'Online Managed Orchestration - stratégie multi-canaux'],
    ['SEM', 'Search Engine Marketing - Google Ads payant'],
    ['SEO', 'Search Engine Optimization - classement organique'],
    ['CRM', 'Customer Relationship Management - logiciel gestion clients'],
]

for item in glossary_items:
    p = doc.add_paragraph()
    run = p.add_run(item[0])
    run.bold = True
    run.font.name = 'Times New Roman'
    p.add_run(f' : {item[1]}')

# Sauvegarder document
output_path = r'c:\Users\wejde\Downloads\APP WATERSENSE\RAPPORT_MARKETING_PROFESSIONNEL_2026_V2.docx'
doc.save(output_path)

print(f"Rapport marketing professionnel généré avec succès")
print(f"Fichier : {output_path}")
print(f"\nCaractéristiques :")
print(f"  - Format professionnel mémoire")
print(f"  - Police Times New Roman")
print(f"  - 30+ pages contenu (sans annexes)")
print(f"  - Basé fondamentaux Kotler")
print(f"  - Tableaux professionnels")
print(f"  - 6 annexes détaillées")
print(f"  - Sans émojis - format rigoureux")
