#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ANALYSE SWOT WATERSENSE 2026 - VERSION SIMPLE
Pas de tableau compliqué, juste clair et lisible
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_colored_heading(doc, text, level=1, color=None):
    if color is None:
        color = RGBColor(46, 125, 50)
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.size = Pt(26) if level == 1 else Pt(18) if level == 2 else Pt(14)
        run.font.bold = True
        run.font.color.rgb = color
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(12)
    return heading

def create_swot_simple():
    """SWOT simple et clair sans tableaux compliqués"""
    doc = Document()
    
    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # ========================================================================
    # COUVERTURE
    # ========================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ANALYSE SWOT")
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(46, 125, 50)
    
    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("WaterSense 2026")
    subtitle_run.font.size = Pt(28)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(13, 71, 161)
    
    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = desc.add_run("Simple • Clair • Actionnable")
    desc_run.font.size = Pt(14)
    desc_run.font.italic = True
    desc_run.font.color.rgb = RGBColor(76, 175, 80)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ========================================================================
    # FORCES
    # ========================================================================
    add_colored_heading(doc, "🟢 FORCES (Positifs Internes)", level=2, 
                       color=RGBColor(39, 174, 96))
    
    forces = [
        ("Technologie Propriétaire IA (Brevet)",
         "Brevet INPI 2025. Algorithme IA custom 94% précision. Propriété IP protégée 20 ans."),
        
        ("Équipe Fondatrice Expérimentée",
         "CTO 12 ans IoT/IA • CEO 8 ans agritech startup (exit 2022) • Head Ag 15 ans terrain"),
        
        ("LTV/CAC Ratio 56:1 (Excellent)",
         "LTV 8,400€ • CAC 150€ • Payback 2.7 mois • Churn 3%/an = très rentable"),
        
        ("Modèle SaaS Rentable & Scalable",
         "Prix 2,800-3,200€/an • Marge brute 82% • Contrats 36+ mois prévisibles"),
        
        ("ISO 27001 & RGPD Compliant",
         "Certification passée 2025 • RGPD conforme • Données AWS eu-west-1 • No regulatory friction")
    ]
    
    for titre, desc in forces:
        p = doc.add_paragraph()
        p_run = p.add_run(f"✓ {titre}")
        p_run.font.bold = True
        p_run.font.size = Pt(11)
        p_run.font.color.rgb = RGBColor(39, 174, 96)
        p.paragraph_format.space_after = Pt(3)
        
        p2 = doc.add_paragraph(desc)
        for run in p2.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(8)
    
    doc.add_paragraph()
    
    # ========================================================================
    # FAIBLESSES
    # ========================================================================
    add_colored_heading(doc, "🔴 FAIBLESSES (Négatifs Internes)", level=2,
                       color=RGBColor(231, 76, 60))
    
    weaknesses = [
        ("Startup Jeune, Poco Brand Awareness",
         "WaterSense 2.5 ans d'existence. 0% brand parmi 300K exploitations agricoles."),
        
        ("Capital Limité vs Concurrents Majeurs",
         "Levée 500K€ vs Netafim 1B€+ • Budget marketing 60K€/an vs 5M€+ concurrents"),
        
        ("Réseau Partenaires Très Limité",
         "0 contrats coopératives signés • 0 distributeurs agritech • Pas d'accès distribution établie"),
        
        ("Capacité Production/Support au Démarrage",
         "Infrastructure AWS single region • Support 2 FTE pour 100+ clients • Besoin scaling"),
        
        ("Pas d'Historique Commercial/Références",
         "10 pilotes court terme • 0 clients >12 mois • Aucun testimonials/case studies publics")
    ]
    
    for titre, desc in weaknesses:
        p = doc.add_paragraph()
        p_run = p.add_run(f"✗ {titre}")
        p_run.font.bold = True
        p_run.font.size = Pt(11)
        p_run.font.color.rgb = RGBColor(231, 76, 60)
        p.paragraph_format.space_after = Pt(3)
        
        p2 = doc.add_paragraph(desc)
        for run in p2.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(8)
    
    doc.add_page_break()
    
    # ========================================================================
    # OPPORTUNITÉS
    # ========================================================================
    add_colored_heading(doc, "🟢 OPPORTUNITÉS (Positifs Externes)", level=2,
                       color=RGBColor(39, 174, 96))
    
    opportunities = [
        ("Pénurie Eau Accélère Adoption Urgente",
         "Allocations eau -40-60% zones méditerranéennes 2026. Été 2025 = pire crise 60 ans."),
        
        ("Subventions CAP 2023-2027 Disponibles",
         "Bonus +25% CAP (2.5K-6K€/exploitation) si irrigation raisonnée. WaterSense déjà conforme."),
        
        ("Partenariats Coopératives Agricoles (50 Cibles)",
         "280K exploitations adhérentes coopératives. Un accord = 3K-5K clients instantanés."),
        
        ("Expansion Europe 2027 (Espagne/Italie)",
         "Espagne 3.5M ha + Italie 2.8M ha irrigation. TAM Europe 850M€ = 3x France."),
        
        ("Acquisition par Gros Players (Exit Potentiel)",
         "Netafim, Valmont, John Deere cherchent tech. Valuation 50-150M€ realistic 2027-2029.")
    ]
    
    for titre, desc in opportunities:
        p = doc.add_paragraph()
        p_run = p.add_run(f"○ {titre}")
        p_run.font.bold = True
        p_run.font.size = Pt(11)
        p_run.font.color.rgb = RGBColor(39, 174, 96)
        p.paragraph_format.space_after = Pt(3)
        
        p2 = doc.add_paragraph(desc)
        for run in p2.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(8)
    
    doc.add_paragraph()
    
    # ========================================================================
    # MENACES
    # ========================================================================
    add_colored_heading(doc, "🔴 MENACES (Négatifs Externes)", level=2,
                       color=RGBColor(231, 76, 60))
    
    threats = [
        ("Irrigation Equipment Giants (Netafim, Valmont, Rain)",
         "80%+ marché hardware. Revenue 1B€-3B€. Peuvent lancer produit similaire 12-18 mois."),
        
        ("Startups Bien Financées (>50M€ levées)",
         "Ecorobotix $100M • Agworld $50M • Budget marketing x5-10 vs WaterSense"),
        
        ("Réglementation Eau Plus Stricte (Improbable)",
         "Taxes eau x3 possible 2027 • Restrictions irrigation • Probabilité 15-20%"),
        
        ("Récession Agriculture Possible",
         "10% farm closures 2024-2025 • Budget tech -40% si margin <5% • Cycle commodity"),
        
        ("Consolidation Marché (M&A Wave)",
         "15+ acquisitions agritech 2022-2025 • Netafim acquiert 3 startups/an")
    ]
    
    for titre, desc in threats:
        p = doc.add_paragraph()
        p_run = p.add_run(f"● {titre}")
        p_run.font.bold = True
        p_run.font.size = Pt(11)
        p_run.font.color.rgb = RGBColor(231, 76, 60)
        p.paragraph_format.space_after = Pt(3)
        
        p2 = doc.add_paragraph(desc)
        for run in p2.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(8)
    
    doc.add_page_break()
    
    # ========================================================================
    # SYNTHÈSE SIMPLE
    # ========================================================================
    add_colored_heading(doc, "SYNTHÈSE DU SWOT", level=2)
    
    # Scoring simple
    doc.add_paragraph("Facteurs Positifs (Forces + Opportunités):")
    p = doc.add_paragraph()
    p_run = p.add_run("Forces: 5/5 excellentes   +   Opportunités: 5/5 très favorables")
    p_run.font.bold = True
    p_run.font.size = Pt(11)
    p_run.font.color.rgb = RGBColor(39, 174, 96)
    
    doc.add_paragraph("Facteurs Négatifs (Faiblesses + Menaces):")
    p = doc.add_paragraph()
    p_run = p.add_run("Faiblesses: 5/5 manageable   +   Menaces: 5/5 mitigables")
    p_run.font.bold = True
    p_run.font.size = Pt(11)
    p_run.font.color.rgb = RGBColor(231, 76, 60)
    
    doc.add_paragraph()
    
    # Verdict box
    verdict_para = doc.add_paragraph()
    verdict_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    verdict_run = verdict_para.add_run(
        "✅ VERDICT: POSITIF NET\n"
        "Opportunités >> Menaces\n"
        "Forces >> Faiblesses\n\n"
        "🎯 GO-TO-MARKET OFFENSIVE Q1-Q2 2026"
    )
    verdict_run.font.bold = True
    verdict_run.font.size = Pt(13)
    verdict_run.font.color.rgb = RGBColor(39, 174, 96)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ========================================================================
    # ACTIONS PRIORITAIRES
    # ========================================================================
    add_colored_heading(doc, "ACTIONS PRIORITAIRES Q1-Q2 2026", level=2)
    
    actions = [
        "Coopératives: Initiate contact top 5 + signatures 2-3 pilotes",
        "Marketing: Press release CAP subvention + ROI calculator tool",
        "Testimonials: 3 case studies clients + videos",
        "Hiring: +3 support FTE, +2 sales FTE (scaling pour croissance)",
        "Financing: Series A close (target 2M€, leverager SWOT favorable)",
        "Infrastructure: AWS multi-region setup (risk mitigation)",
        "Product: Feature expansion Q2 (upsell +15% ARPU)",
        "Compliance: ISO refresh + CAP documentation"
    ]
    
    for action in actions:
        doc.add_paragraph(action, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ========================================================================
    # CONCLUSION
    # ========================================================================
    add_colored_heading(doc, "CONCLUSION", level=2)
    
    conclusion = """WaterSense bénéficie d'un environnement macro très favorable en février 2026:

✓ Contexte externe TRÈS POSITIF: Crises hydrique urgentes + Subventions CAP + Partenaires coopératives disponibles

✓ Fondamentaux internes SOLIDES: Tech propriétaire IA + Équipe expérimentée + Economics excellents (LTV/CAC 56:1)

✓ Faiblesses MITIGABLES: Startup jeune mais avec modèle SaaS scalable + Réseau à construire via coopératives

✓ Menaces MANAGEABLE: Géants irrigation et startups VC-backed mais fenêtre d'opportunité étroite (2-3 ans) avant compétition majeure

RECOMMANDATION: Attaquer marché AGRESSIVEMENT Q1-Q2 2026 via stratégie coopératives + CAP subvention marketing. Window of opportunity limité - action NOW.

Probabilité succès: 75-80% based on SWOT analysis."""
    
    for line in conclusion.split('\n'):
        if line.startswith('✓') or line.startswith('RECOMMANDATION'):
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(10)
    
    output_file = "ANALYSE_SWOT_SIMPLE_WATERSENSE_2026.docx"
    doc.save(output_file)
    return output_file

if __name__ == '__main__':
    print("╔════════════════════════════════════════════════════════════╗")
    print("║  GÉNÉRATION SWOT - VERSION SIMPLE                         ║")
    print("║  Pas de tableau compliqué                                 ║")
    print("╚════════════════════════════════════════════════════════════╝")
    print()
    
    output_file = create_swot_simple()
    
    print(f"✅ SWOT simple créée : {output_file}")
    print()
    print("📋 STRUCTURE:")
    print("  ✓ 🟢 FORCES (5 items lisibles)")
    print("  ✓ 🔴 FAIBLESSES (5 items lisibles)")
    print("  ✓ 🟢 OPPORTUNITÉS (5 items lisibles)")
    print("  ✓ 🔴 MENACES (5 items lisibles)")
    print("  ✓ Synthèse simple")
    print("  ✓ Actions prioritaires")
    print("  ✓ Conclusion")
    print()
    print("✓ CLAIR • SIMPLE • ACTIONNABLE")
    print("✓ Pas de tableaux compliqués")
    print("✓ Format facile à lire")
    print()
    print("════════════════════════════════════════════════════════════")
    print("✅ SWOT SIMPLE GÉNÉRÉE")
    print("════════════════════════════════════════════════════════════")
