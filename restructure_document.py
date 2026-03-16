from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color):
    """Ajoute couleur de fond à une cellule"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

print("📖 Chargement du document existant...")
doc = Document(r'c:\Users\wejde\Downloads\APP WATERSENSE\DOSSIER_MARKETING_WATERSENSE_DETAILLE.docx')

print("✏️ Restructuration du document...")

# Nettoyer et réorganiser
for para in doc.paragraphs:
    # Améliorer les titres
    if para.style.name.startswith('Heading'):
        for run in para.runs:
            run.font.bold = True
            if para.style.name == 'Heading 1':
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0, 102, 204)
            elif para.style.name == 'Heading 2':
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 102, 204)
    
    # Améliorer l'espacement
    if para.text.strip():
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.15

# Améliorer les tableaux
print("📊 Amélioration des tableaux...")
colors = ['0066CC', '00AA44', '1976D2', 'F57C00', 'C62828', '6A1B9A']
for table_idx, table in enumerate(doc.tables):
    if len(table.rows) > 0:
        header_row = table.rows[0]
        color = colors[table_idx % len(colors)]
        
        # Colorer l'en-tête
        for cell in header_row.cells:
            shade_cell(cell, color)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Alternance de couleur
        for row_idx, row in enumerate(table.rows[1:], 1):
            if row_idx % 2 == 0:
                for cell in row.cells:
                    shade_cell(cell, 'F0F0F0')

# Ajouter une page de titre moderne
title_para = doc.paragraphs[0]._element
new_title = doc.add_paragraph()
new_title_run = new_title.add_run('WaterSense')
new_title_run.font.size = Pt(28)
new_title_run.font.bold = True
new_title_run.font.color.rgb = RGBColor(0, 102, 204)
new_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Dossier Marketing - Irrigation Intelligente pour Agriculteurs')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()  # Saut

# Sauvegarder
output_path = r'c:\Users\wejde\Downloads\APP WATERSENSE\DOSSIER_MARKETING_WATERSENSE_DETAILLE.docx'
doc.save(output_path)

print(f"✅ Document restructuré et sauvegardé!")
print(f"📄 Fichier : {output_path}")
print(f"✓ Mise en forme appliquée:")
print(f"  - En-têtes colorés et structurés")
print(f"  - Tableaux améliorés avec couleurs")
print(f"  - Espacement optimisé")
print(f"  - Police professionnelle")
