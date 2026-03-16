#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convertisseur Markdown → Word (.docx)
Crée des fichiers Word formatés à partir des fichiers Markdown
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

def add_shading(paragraph, color):
    """Ajoute une couleur de fond au paragraphe"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    paragraph._element.get_or_add_pPr().append(shading_elm)

def markdown_to_word(md_file, docx_file):
    """Convertit un fichier Markdown en Word (.docx)"""
    
    doc = Document()
    
    # Lire le fichier Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Paragraphe vide
            doc.add_paragraph()
        
        elif line.startswith('# '):
            # Titre niveau 1
            title = line.replace('# ', '')
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif line.startswith('## '):
            # Titre niveau 2
            title = line.replace('## ', '')
            heading = doc.add_heading(title, level=2)
            add_shading(heading, 'E8F5E9')  # Vert clair
            
        elif line.startswith('### '):
            # Titre niveau 3
            title = line.replace('### ', '')
            doc.add_heading(title, level=3)
            
        elif line.startswith('#### '):
            # Titre niveau 4
            title = line.replace('#### ', '')
            doc.add_heading(title, level=4)
            
        elif line.startswith('- '):
            # Bullet point
            item = line.replace('- ', '')
            doc.add_paragraph(item, style='List Bullet')
            
        elif line.startswith('✓ ') or line.startswith('✅ '):
            # Checkbox
            item = line.replace('✓ ', '').replace('✅ ', '')
            p = doc.add_paragraph(item, style='List Bullet')
            
        elif line.startswith('| '):
            # Tableau - à traiter spécialement
            pass  # Sera traité après
            
        else:
            # Paragraphe normal
            if line and not line.startswith(('---', '___')):
                p = doc.add_paragraph(line)
                p.paragraph_format.line_spacing = 1.5
    
    # Sauvegarder le document Word
    doc.save(docx_file)
    print(f"Created: {docx_file}")

def create_word_from_markdown_advanced(md_file, docx_file):
    """Version avancée avec meilleur formatage"""
    
    doc = Document()
    
    # Setup document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Traiter les sections
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        if not line.strip():
            i += 1
            continue
        
        # Titre niveau 1
        if line.startswith('# '):
            text = line.replace('# ', '').strip()
            p = doc.add_heading(text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(24)
                run.font.bold = True
            i += 1
        
        # Titre niveau 2
        elif line.startswith('## '):
            text = line.replace('## ', '').strip()
            p = doc.add_heading(text, level=1)
            p_format = p.paragraph_format
            p_format.space_before = Pt(12)
            p_format.space_after = Pt(12)
            for run in p.runs:
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(46, 125, 50)  # Vert foncé
            i += 1
        
        # Titre niveau 3
        elif line.startswith('### '):
            text = line.replace('### ', '').strip()
            p = doc.add_heading(text, level=2)
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.bold = True
            i += 1
        
        # Titre niveau 4
        elif line.startswith('#### '):
            text = line.replace('#### ', '').strip()
            p = doc.add_heading(text, level=3)
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.bold = True
            i += 1
        
        # Tableau
        elif line.strip().startswith('| '):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('| '):
                table_lines.append(lines[i])
                i += 1
            
            if table_lines:
                # Parser la table
                rows = [line.strip().split('|')[1:-1] for line in table_lines]
                rows = [[cell.strip() for cell in row] for row in rows]

                # Supprimer les lignes séparatrices (---)
                filtered_rows = []
                for row in rows:
                    if not row:
                        continue
                    if all(set(cell.replace(':', '').replace('-', '').strip()) == set() for cell in row):
                        continue
                    filtered_rows.append(row)

                if filtered_rows:
                    max_cols = max(len(r) for r in filtered_rows)
                    normalized_rows = [r + [''] * (max_cols - len(r)) for r in filtered_rows]

                    table = doc.add_table(rows=len(normalized_rows), cols=max_cols)
                    table.style = 'Light Grid Accent 1'

                    for j, row in enumerate(normalized_rows):
                        for k, cell in enumerate(row):
                            table.rows[j].cells[k].text = cell
                            if j == 0:  # Header
                                for paragraph in table.rows[j].cells[k].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
            continue
        
        # Bullet point
        elif line.strip().startswith('- '):
            text = line.strip().replace('- ', '').strip()
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
        
        # Checkbox
        elif line.strip().startswith('[ ] ') or line.strip().startswith('[x] '):
            text = line.strip().replace('[ ] ', '☐ ').replace('[x] ', '☑ ')
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
        
        # Horizontal line
        elif line.strip() in ('---', '___', '***'):
            doc.add_paragraph('_' * 60)
            i += 1
        
        # Paragraphe normal
        else:
            text = line.strip()
            if text and not text.startswith(('#', '|', '-', '[', '`')):
                p = doc.add_paragraph(text)
                p.paragraph_format.line_spacing = 1.15
            i += 1
    
    doc.save(docx_file)
    print(f"Created: {docx_file}")

# ============================================================================
# CONVERTIR TOUS LES FICHIERS
# ============================================================================

if __name__ == '__main__':
    files_to_convert = [
        ('PITCH_WATERSENSE_ONE_PAGE.md', 'PITCH_WATERSENSE_ONE_PAGE.docx'),
        ('DOSSIER_MARKETING_COMPLET_WATERSENSE_2026.md', 'DOSSIER_MARKETING_COMPLET_WATERSENSE_2026.docx'),
        ('RAPPORT_MARKETING_WATERSENSE_2026.md', 'RAPPORT_MARKETING_WATERSENSE_2026.docx'),
        ('ACTIONPLAN_2026_CHECKLIST.md', 'ACTIONPLAN_2026_CHECKLIST.docx'),
        ('00_INDEX_COLLECTION_COMPLETE.md', '00_INDEX_COLLECTION_COMPLETE.docx'),
        ('DELIVERABLES_SUMMARY.md', 'DELIVERABLES_SUMMARY.docx'),
    ]
    
    print("╔════════════════════════════════════════════════════════════╗")
    print("║   Conversion Markdown → Word (.docx)                      ║")
    print("╚════════════════════════════════════════════════════════════╝")
    print()
    
    for md_file, docx_file in files_to_convert:
        if os.path.exists(md_file):
            print(f"🔄 Conversion: {md_file} → {docx_file}")
            try:
                create_word_from_markdown_advanced(md_file, docx_file)
            except Exception as e:
                print(f"❌ Erreur: {e}")
        else:
            print(f"⚠️  Fichier non trouvé: {md_file}")
        print()
    
    print("════════════════════════════════════════════════════════════")
    print("✅ CONVERSION TERMINÉE - Tous les fichiers Word sont prêts!")
    print("════════════════════════════════════════════════════════════")
