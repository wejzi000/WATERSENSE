#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ANALYSE SWOT WATERSENSE 2026 - TABLEAU 2x2 STRUCTURÉ (4 QUADRANTS)
4 dimensions clairement séparées avec détails complets
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_colored_heading(doc, text, level=1, color=None):
    if color is None:
        color = RGBColor(46, 125, 50)
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.size = Pt(26) if level == 1 else Pt(18) if level == 2 else Pt(14)
        run.font.bold = True
        run.font.color.rgb = color
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(12)
    return heading

def add_quadrant_content(cell, title, color, factors):
    """Add content to a SWOT quadrant cell"""
    cell.text = ""
    
    # Color the cell
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading)
    
    # Title
    title_para = cell.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.bold = True
    title_run.font.size = Pt(12)
    title_run.font.color.rgb = RGBColor(255, 255, 255) if color in ['27AE60', 'E74C3C'] else RGBColor(255, 255, 255)
    title_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(8)
    
    # Add each factor
    for factor in factors:
        # Factor title
        f_para = cell.add_paragraph()
        f_run = f_para.add_run(f"✓ {factor['titre']}")
        f_run.font.bold = True
        f_run.font.size = Pt(9)
        f_run.font.color.rgb = RGBColor(39, 174, 96) if 'O' not in title and 'E' not in title else RGBColor(39, 174, 96) if 'O' in title else RGBColor(231, 76, 60)
        f_para.paragraph_format.space_after = Pt(3)
        f_para.paragraph_format.left_indent = Cm(0.2)
        
        # Description
        d_para = cell.add_paragraph()
        d_run = d_para.add_run(factor['description'])
        d_run.font.size = Pt(8)
        d_run.font.italic = True
        d_para.paragraph_format.space_after = Pt(2)
        d_para.paragraph_format.left_indent = Cm(0.3)
        
        # Details
        details_para = cell.add_paragraph()
        details_run = details_para.add_run(f"→ {factor['details']}")
        details_run.font.size = Pt(7.5)
        details_run.font.color.rgb = RGBColor(70, 70, 70)
        details_para.paragraph_format.space_after = Pt(2)
        details_para.paragraph_format.left_indent = Cm(0.4)
        
        # Impact/Score
        impact_para = cell.add_paragraph()
        impact_run = impact_para.add_run(f"• Impact: {factor['impact']} | Score: {factor['score']}")
        impact_run.font.size = Pt(7)
        impact_run.font.bold = True
        impact_para.paragraph_format.space_after = Pt(6)
        impact_para.paragraph_format.left_indent = Cm(0.3)

def create_swot_4dim():
    """Tableau SWOT 2x2 avec 4 dimensions"""
    doc = Document()
    
    # Marges
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    
    # ========================================================================
    # COUVERTURE
    # ========================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("MATRICE SWOT 2x2")
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(46, 125, 50)
    
    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("WaterSense 2026")
    subtitle_run.font.size = Pt(40)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(13, 71, 161)
    
    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = desc.add_run("4 Quadrants Distincts - Analyse Complète")
    desc_run.font.size = Pt(14)
    desc_run.font.italic = True
    desc_run.font.color.rgb = RGBColor(76, 175, 80)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ========================================================================
    # MAIN 2x2 TABLE
    # ========================================================================
    add_colored_heading(doc, "TABLEAU SWOT 2x2 - 4 QUADRANTS DÉTAILLÉS", level=2)
    
    # Create 2x2 table
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(3.8)
        row.cells[1].width = Inches(3.8)
    
    # ====== DATA ======
    forces = [
        {
            "titre": "Technologie Propriétaire IA (Brevet)",
            "description": "Brevet INPI 2025. Algorithme IA 94% précision, propriété IP 20 ans.",
            "details": "Brevet déposé | IA propriétaire | Barrier to entry",
            "impact": "TRÈS HAUT", "score": "+4"
        },
        {
            "titre": "Équipe Expérimentée",
            "description": "CTO 12 ans IoT/IA, CEO 8 ans agritech startup (exit 2022), Head Ag 15 ans terrain.",
            "details": "Exécution rapide | Décisions quality | Network établi",
            "impact": "TRÈS HAUT", "score": "+4"
        },
        {
            "titre": "LTV/CAC 56:1 (Excellent)",
            "description": "LTV 8,400€, CAC 150€, payback 2.7 mois, churn 3%/an.",
            "details": "Unit economics excellent | Viable long-term | Investisseur attractif",
            "impact": "TRÈS HAUT", "score": "+4"
        },
        {
            "titre": "Modèle SaaS Rentable",
            "description": "Prix 2,800-3,200€/an, marge brute 82%, contrats 36+ mois.",
            "details": "Revenue prévisible | Marges élevées | Scaling simple",
            "impact": "HAUT", "score": "+3"
        },
        {
            "titre": "ISO/RGPD Compliant",
            "description": "ISO 27001 PASSED 2025, RGPD complet, données AWS eu-west-1.",
            "details": "Compliance ready | No regulatory friction | Advantage vs startups",
            "impact": "HAUT", "score": "+3"
        }
    ]
    
    opportunities = [
        {
            "titre": "Pénurie Eau Accélère Adoption",
            "description": "Allocations -40-60% zones méditerranéennes. Été 2025 = pire crise 60 ans. Urgence maximale.",
            "details": "BESOIN #1 agriculteurs | Crises long-term | ROI eau critique",
            "impact": "TRÈS FORT", "score": "+5"
        },
        {
            "titre": "Subventions CAP 2023-2027",
            "description": "Bonus +25% CAP (2.5K-6K€/exploitation) si irrigation raisonnée. WaterSense conforme.",
            "details": "Budget 180M€ | ROI accelerator | Fenêtre 2026-2027",
            "impact": "TRÈS FORT", "score": "+5"
        },
        {
            "titre": "Partenariats Coopératives (50 Cibles)",
            "description": "280K exploitations adhérentes. Accord unique = 3K-5K clients instantanés.",
            "details": "Distribution scale | Revenue share 20/80 | Croissance 10x",
            "impact": "TRÈS FORT", "score": "+4"
        },
        {
            "titre": "Expansion Europe 2027",
            "description": "Espagne 3.5M ha + Italie 2.8M ha. TAM 850M€ (3x France). Même régulation UE.",
            "details": "International scaling | 2027-2028 launch | Revenue 3x",
            "impact": "FORT", "score": "+4"
        },
        {
            "titre": "Acquisition Gros Players (Exit)",
            "description": "Netafim, Valmont, John Deere cherchent tech. Valuation 50-150M€ 2027-2029.",
            "details": "Exit realistic 2027-2029 | Comparable exits 50M€+ | ROI shareholder significatif",
            "impact": "MOYEN-HAUT", "score": "+3"
        }
    ]
    
    weaknesses = [
        {
            "titre": "Startup Jeune, Poco Brand Awareness",
            "description": "WaterSense 2.5 ans. 0% brand parmi 300K exploitations. Besoin marketing budget.",
            "details": "0 historique | 0 références | Budget limité | Long sales cycle",
            "impact": "MOYEN", "score": "-2"
        },
        {
            "titre": "Capital Limité vs Concurrents",
            "description": "Levée 500K€ vs Netafim 1B€+, Ecorobotix 100M€. Budget marketing 60K€/an.",
            "details": "R&D réduit | Budget sales limité | Dépendance financement",
            "impact": "HAUT", "score": "-3"
        },
        {
            "titre": "Réseau Partenaires Limité",
            "description": "0 contrats coopératives signés. 0 distributeurs agritech. Pas accès 3,000 points vente.",
            "details": "Channel to-build | Distribution to-establish | No established partners",
            "impact": "HAUT", "score": "-3"
        },
        {
            "titre": "Capacité Production/Support Démarrage",
            "description": "Infrastructure AWS single region. Support 2 FTE pour 100+ clients. SLA 99.5% vs 99.9%.",
            "details": "Scaling challenge | Onboarding slow (8h vs 3h target) | Risk outage",
            "impact": "MOYEN", "score": "-2"
        },
        {
            "titre": "Pas Historique Commercial/Références",
            "description": "10 pilotes court terme. 0 clients >12 mois. Aucun case studies publics.",
            "details": "No testimonials | Short track record | Mid-market hesitant",
            "impact": "MOYEN-HAUT", "score": "-2"
        }
    ]
    
    threats = [
        {
            "titre": "Irrigation Giants (Netafim, Valmont, Rain)",
            "description": "80%+ marché hardware. Revenue 1.1B€-3.8B€. Can launch product 12-18 mois.",
            "details": "Lanch produit similaire | Distribution entrenched | Predatory pricing",
            "impact": "TRÈS FORT", "score": "-3"
        },
        {
            "titre": "Startups Bien Financées (>50M€)",
            "description": "Ecorobotix $100M, Agworld $50M, Encepta $20M. VC backing massif.",
            "details": "Budget x5-10 vs WS | Undercut pricing | Talent poaching",
            "impact": "HAUT", "score": "-3"
        },
        {
            "titre": "Réglementation Eau Plus Stricte",
            "description": "Taxes eau x3 possible 2027. Restrictions irrigation. Probabilité 15-20%.",
            "details": "Budget farm réduit | Moins urgent tech invest | TAM -30% worst-case",
            "impact": "MOYEN", "score": "-1"
        },
        {
            "titre": "Récession Agriculture",
            "description": "10% farm closures 2024-2025. Budget tech -40% si margin <5%. Commodity cycle.",
            "details": "Survival vs growth mode | De-prioritize tech | Cycle risk",
            "impact": "MOYEN-HAUT", "score": "-2"
        },
        {
            "titre": "Consolidation Marché (M&A Wave)",
            "description": "15+ acquisitions agritech 2022-2025. Netafim acquiert 3 startups/an.",
            "details": "2-3 compétiteurs acquired | Landscape reshuffles | Acquisition offer",
            "impact": "MOYEN", "score": "-2"
        }
    ]
    
    # Top-left: FORCES (Green)
    forces_cell = table.rows[0].cells[0]
    add_quadrant_content(forces_cell, "🟢 FORCES (INTERNE - POSITIF)", "27AE60", forces)
    
    # Top-right: OPPORTUNITÉS (Green)
    opps_cell = table.rows[0].cells[1]
    add_quadrant_content(opps_cell, "🟢 OPPORTUNITÉS (EXTERNE - POSITIF)", "27AE60", opportunities)
    
    # Bottom-left: FAIBLESSES (Red)
    weak_cell = table.rows[1].cells[0]
    add_quadrant_content(weak_cell, "🔴 FAIBLESSES (INTERNE - NÉGATIF)", "E74C3C", weaknesses)
    
    # Bottom-right: MENACES (Red)
    threat_cell = table.rows[1].cells[1]
    add_quadrant_content(threat_cell, "🔴 MENACES (EXTERNE - NÉGATIF)", "E74C3C", threats)
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ========================================================================
    # SYNTHÈSE SCORING
    # ========================================================================
    add_colored_heading(doc, "SYNTHÈSE SCORING & VERDICT", level=2)
    
    # Scoring summary
    summary_text = """
┌─────────────────────────────────────────────────────────────────┐
│ DIMENSIÓN          SCORE    ÉVALUATION              DÉTAIL       │
├─────────────────────────────────────────────────────────────────┤
│ 🟢 FORCES          +18/20   ✅ EXCELLENTE           5 facteurs   │
│ 🔴 FAIBLESSES      -14/20   ⚠️ MANAGEABLE            5 facteurs   │
│                    ─────────────────────────────────────────   │
│ BILAN INTERNE      +4/40    ✅ POSITIF              Net +4      │
├─────────────────────────────────────────────────────────────────┤
│ 🟢 OPPORTUNITÉS    +21/20   ✅ TRÈS FAVORABLE       5 facteurs   │
│ 🔴 MENACES         -11/20   ⚠️ MITIGABLE            5 facteurs   │
│                    ─────────────────────────────────────────   │
│ BILAN EXTERNE      +10/40   ✅ TRÈS POSITIF         Net +10     │
├─────────────────────────────────────────────────────────────────┤
│ TOTAL NET          +14/80   ✅ POSITIF              Score 7.25   │
│ RATIO Opp/Menaces  21 / 11  ✅ 1.91x FAVORABLE      > 1.5x ✓    │
└─────────────────────────────────────────────────────────────────┘
"""
    
    for line in summary_text.split('\n'):
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
    
    doc.add_paragraph()
    
    # ========================================================================
    # VERDICT FINAL
    # ========================================================================
    add_colored_heading(doc, "VERDICT & RECOMMANDATIONS STRATÉGIQUES", level=2)
    
    verdict_para = doc.add_paragraph()
    verdict_run = verdict_para.add_run(
        "✅ VERDICT FINAL: POSITIF NET - GO-TO-MARKET OFFENSIVE Q1-Q2 2026"
    )
    verdict_run.font.bold = True
    verdict_run.font.size = Pt(13)
    verdict_run.font.color.rgb = RGBColor(39, 174, 96)
    
    doc.add_paragraph()
    
    analysis = """
Analyse SWOT Complète:
✓ Forces internes solides (18/20) compensent faiblesses (14/20)
✓ Contexte externe très favorable (21/20 vs 11/20 menaces)
✓ Ratio Opportunités/Menaces = 1.91x (excellent, >1.5x = bon)
✓ Aucune menace existentielle bloquante identifiée
✓ Timing fenêtre d'opportunité critique (CAP 2026-2027, eau urgence)

Stratégies Prioritaires par Quadrant:

🔝 STRATÉGIE SO (Maximiser Forces + Opportunités):
   → Leverager tech IA propriétaire + équipe expérimentée
   → Coopératives partnerships comme channel distribution
   → CAP subvention bonus comme core value proposition (marketing)
   → Équipe de fondateurs contacte top 10 coopératives Q2 2026

🛡️ STRATÉGIE WO (Mitiger Faiblesses via Opportunités):
   → Pénurie eau = brand awareness accelerator (urgency narrative)
   → Coopératives partnerships réduit besoin capital marketing massif
   → Partenaires = références clients (mitigation 'no track record')

⚔️ STRATÉGIE ST (Défendre contre Menaces):
   → Netafim/giants entry: Coopératives partnerships rapide (before competition)
   → Startups VC-backed: Local support français + interface simple = différenciation
   → Récession agri: SaaS flexible (payment plans, no lock-in)
   → Consolidation: Rester attractif acquisition (50-150M€ valuation target 2027-2029)

🚀 STRATÉGIE OT (Positionnement Opportuniste):
   → Crises hydrique = urgency narratif marketing principal
   → CAP subvention = ROI calculator direct vs concurrents
   → Exit strategy 2027-2029: acquisition realistic (5x-10x CAGR)
   → Europe 2027: Espagne/Italie expansion avec proven model France
"""
    
    for line in analysis.split('\n'):
        if line.startswith('   →'):
            p = doc.add_paragraph(line[4:], style='List Bullet 2')
            p.paragraph_format.left_indent = Cm(1.2)
        elif any(x in line for x in ['STRATÉGIE', '🔝', '🛡️', '⚔️', '🚀', 'Stratégies']):
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # Action plan
    add_colored_heading(doc, "PLAN D'ACTION Q1-Q2 2026", level=3)
    
    actions = [
        ("Q1 2026", [
            "Coopératives: Initiate contact top 5 (Beauce, Loire, Provence)",
            "Marketing: Press release CAP subvention + ROI calculator tool",
            "Testimonials: Collect 3 pilote case studies + videos",
            "Financing: Series A prep (leverager SWOT favorable)",
            "Infrastructure: AWS multi-region setup (risk mitigation)",
        ]),
        ("Q2 2026", [
            "Coopératives: 2-3 pilot agreements signed (revenue share)",
            "Hiring: +3 support FTE, +2 sales FTE (scaling team)",
            "Series A: Fundraising close (target 2M€)",
            "Product: Feature expansion (upsell +15% ARPU)",
            "Compliance: ISO 27001 audit refresh, CAP documentation",
        ])
    ]
    
    for quarter, items in actions:
        q_para = doc.add_paragraph(f"{quarter}:")
        for run in q_para.runs:
            run.font.bold = True
            run.font.size = Pt(10)
        
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    output_file = "ANALYSE_SWOT_2x2_QUADRANTS_WATERSENSE_2026.docx"
    doc.save(output_file)
    return output_file

if __name__ == '__main__':
    print("╔════════════════════════════════════════════════════════════╗")
    print("║  GÉNÉRATION SWOT - TABLEAU 2x2 (4 QUADRANTS)             ║")
    print("║  Analyse Complète & Plan d'Action                        ║")
    print("╚════════════════════════════════════════════════════════════╝")
    print()
    
    output_file = create_swot_4dim()
    
    print(f"✅ SWOT 4 quadrants créée : {output_file}")
    print()
    print("📊 STRUCTURE 2x2 :")
    print("  ┌─────────────────────────┬─────────────────────────┐")
    print("  │ 🟢 FORCES (18/20)       │ 🟢 OPPORTUNITÉS (21/20) │")
    print("  ├─────────────────────────┼─────────────────────────┤")
    print("  │ 🔴 FAIBLESSES (-14/20)  │ 🔴 MENACES (-11/20)     │")
    print("  └─────────────────────────┴─────────────────────────┘")
    print()
    print("✓ 5 facteurs par quadrant (20 total)")
    print("✓ Synthèse scoring complète")
    print("✓ 4 stratégies (SO/WO/ST/OT)")
    print("✓ Plan action Q1-Q2 2026")
    print("✓ Verdict: +14 net POSITIF (7.25/10)")
    print()
    print("════════════════════════════════════════════════════════════")
    print("✅ SWOT 4 QUADRANTS GÉNÉRÉE")
    print("════════════════════════════════════════════════════════════")
